import re
import io
from typing import Dict, Any
from docx import Document
import pdfplumber


class ResumeParser:
    """
    简历解析工具类
    支持解析 Word(.docx) 和 PDF 格式的简历文件
    """

    @classmethod
    def parse_word_resume(cls, file_path: str) -> Dict[str, Any]:
        """
        解析 Word 格式简历
        
        :param file_path: Word 文件路径
        :return: 解析后的简历信息字典
        """
        try:
            doc = Document(file_path)
            text = '\n'.join([para.text for para in doc.paragraphs])
            
            # 解析表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' '.join([cell.text for cell in row.cells])
                    text += '\n' + row_text
            
            return cls._extract_info(text)
        except Exception as e:
            raise Exception(f'Word 简历解析失败: {str(e)}')

    @classmethod
    def parse_pdf_resume(cls, file_path: str) -> Dict[str, Any]:
        """
        解析 PDF 格式简历
        
        :param file_path: PDF 文件路径
        :return: 解析后的简历信息字典
        """
        try:
            text = ''
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + '\n'
            
            return cls._extract_info(text)
        except Exception as e:
            raise Exception(f'PDF 简历解析失败: {str(e)}')

    @classmethod
    def _extract_info(cls, text: str) -> Dict[str, Any]:
        """
        从文本中提取简历信息
        
        :param text: 简历文本内容
        :return: 提取的信息字典
        """
        info = {
            'name': None,
            'sex': None,
            'age': None,
            'graduate_school': None,
            'phone': None,
            'email': None,
            'work_years': None,
        }
        
        # 提取姓名（通常在文档开头或"姓名："后面）
        name_patterns = [
            r'姓名[：:]\s*([^\s\n]+)',
            r'([^\s\n]{2,4})\s*(?:男|女)',  # 姓名后跟性别
            r'^([^\s\n]{2,4})\s*$',  # 文档开头可能是姓名
        ]
        for pattern in name_patterns:
            match = re.search(pattern, text, re.MULTILINE)
            if match:
                info['name'] = match.group(1).strip()
                break
        
        # 提取性别
        sex_patterns = [
            r'性别[：:]\s*([男女])',
            r'(男|女)',
        ]
        for pattern in sex_patterns:
            match = re.search(pattern, text)
            if match:
                sex_text = match.group(1)
                info['sex'] = '1' if sex_text == '女' else '0'
                break
        
        # 提取年龄
        age_patterns = [
            r'年龄[：:]\s*(\d{1,2})',
            r'(\d{1,2})\s*岁',
        ]
        for pattern in age_patterns:
            match = re.search(pattern, text)
            if match:
                info['age'] = int(match.group(1))
                break
        
        # 提取毕业院校
        school_patterns = [
            r'毕业院校[：:]\s*([^\s\n]+)',
            r'学校[：:]\s*([^\s\n]+)',
            r'([^\s\n]{2,20}(?:大学|学院|学校))',
        ]
        for pattern in school_patterns:
            match = re.search(pattern, text)
            if match:
                info['graduate_school'] = match.group(1).strip()
                break
        
        # 提取手机号码
        phone_patterns = [
            r'(?:电话|手机|手机?号|联系?电话|移动?电话)[：:]\s*(\d{11})',
            r'(1[3-9]\d{9})',
        ]
        for pattern in phone_patterns:
            match = re.search(pattern, text)
            if match:
                info['phone'] = match.group(1).strip()
                break
        
        # 提取邮箱
        email_patterns = [
            r'(?:邮箱|电子?邮箱|email|e-mail)[：:]\s*([\w\.-]+@[\w\.-]+\.\w+)',
            r'([\w\.-]+@[\w\.-]+\.\w+)',
        ]
        for pattern in email_patterns:
            match = re.search(pattern, text)
            if match:
                info['email'] = match.group(1).strip()
                break
        
        # 提取工作年限
        work_years_patterns = [
            r'(?:工作?年限|工作?经验|从业?年限|行业?经验)[：:]\s*(\d+)',
            r'(\d+)\s*年(?:工作|从业|行业)?经验',
            r'(?:工作|从业|行业)?经验\s*(\d+)\s*年',
        ]
        for pattern in work_years_patterns:
            match = re.search(pattern, text)
            if match:
                info['work_years'] = int(match.group(1))
                break
        
        return info