"""
简历知识库模块服务层
"""
import json
import re
import time
import uuid
import asyncio
from collections.abc import AsyncGenerator
from typing import Any

from fastapi import Request
from sqlalchemy import or_
from sqlalchemy.ext.asyncio import AsyncSession

from common.vo import CrudResponseModel, PageModel
from exceptions.exception import ServiceException
from module_ai.dao.ai_model_dao import AiModelDao
from module_ai.entity.vo.ai_model_vo import AiModelModel
from module_resume_kb.dao.resume_dao import ResumeDao
from module_resume_kb.entity.do.resume_do import OaResume, OaResumeWork, OaResumeProject
from module_resume_kb.entity.vo.resume_vo import (
    ResumeModel,
    ResumePageQueryModel,
    ResumeUploadResultModel,
    ResumeWorkExperienceModel,
    ResumeProjectExperienceModel,
    ResumeSearchResultModel,
    ResumeChatRequestModel,
    ResumeChatRelatedResumeModel,
)
from utils.ai_util import AiUtil
from config.env import AiResumeParserConfig
from utils.common_util import CamelCaseUtil
from utils.crypto_util import CryptoUtil
from utils.log_util import logger
from utils.resume_parser import ResumeParser
import litellm


class ResumeKbService:
    """
    简历知识库模块服务层
    """

    # ========== 简历标签关键词配置 ==========
    EDUCATION_KEYWORDS = ['博士', '硕士', '研究生', '本科', '学士', '大专', '专科', '高中', '中专', '职高']
    RESUME_KEYWORDS = [
        '工作经历', '项目经验', '教育背景', '专业技能', '自我评价',
        '工作描述', '项目描述', '职责', '任职', '毕业', '院校'
    ]
    TECH_SKILLS_KEYWORDS = [
        # 编程语言
        'Java', 'Python', 'C++', 'C#', 'Go', 'Golang', 'JavaScript', 'TypeScript', 'Kotlin', 'Scala', 'Rust',
        'PHP', 'Ruby', 'Swift', 'Objective-C', 'Shell', 'Bash', 'SQL', 'R', 'MATLAB',
        # 前端框架/库
        'React', 'Vue', 'Vue3', 'Angular', 'Next.js', 'Nuxt.js', 'Element Plus', 'Element-Plus',
        'Ant Design', 'TailwindCSS', 'Webpack', 'Vite', 'jQuery', 'Bootstrap',
        # 后端框架
        'Spring', 'SpringBoot', 'SpringCloud', 'SpringMVC', 'SpringAI', 'Spring AI', 'Spring AI Alibaba',
        'SpringCache', 'MyBatis', 'MyBatis-Plus', 'MybatisPlus', 'Mybatis-Plus', 'Hibernate', 'JPA',
        'Django', 'Flask', 'FastAPI', 'Tornado', 'Express', 'Koa', 'NestJS',
        # 微服务/分布式
        'Nacos', 'OpenFeign', 'Feign', 'Gateway', 'Sentinel', 'Seata', 'Dubbo', 'Grpc', 'gRPC',
        'Consul', 'Eureka', 'Ribbon', 'Hystrix', 'Loadbalancer', 'SpringCloud Alibaba',
        # 数据库
        'MySQL', 'Redis', 'MongoDB', 'Elasticsearch', 'ElasticSearch', 'PostgreSQL', 'Oracle',
        'SQL Server', 'ClickHouse', 'Neo4j', 'Milvus', 'Qdrant', 'TiDB', 'Cassandra',
        # 消息队列
        'Kafka', 'RabbitMQ', 'RocketMQ', 'ActiveMQ', 'Pulsar',
        # 容器/部署/运维
        'Docker', 'Kubernetes', 'K8s', 'Istio', 'Jenkins', 'GitLab CI', 'Nginx', 'Tomcat',
        'Ansible', 'SaltStack', 'Terraform', 'Prometheus', 'Grafana', 'Zabbix',
        # 大数据/AI
        'Hadoop', 'Spark', 'Flink', 'Hive', 'HBase', 'Zookeeper',
        'TensorFlow', 'PyTorch', 'Scikit-learn', 'LangChain', 'LlamaIndex', 'LangGraph',
        'RAG', 'FunctionCalling', 'Function Calling', 'MCP', 'Prompt工程',
        'Pandas', 'NumPy', 'OpenCV', 'PaddleOCR', 'Celery', 'Streamlit',
        # 工具/其他
        'Git', 'Maven', 'Gradle', 'Svn', 'Xxl-Job', 'XXL-JOB', 'XXL-Job',
        'Redisson', 'ShardingSphere', 'Canal', 'MyCat', 'EasyExcel', 'HuTool', 'Hutool',
        'WebSocket', 'JFreeChart', 'RestTemplate', 'WebClient', 'OkHttp', 'HttpClient',
        'OSS', 'S3', 'MinIO', 'FastDFS',
        'JWT', 'OAuth2', 'Spring Security', 'Shiro',
        # 技术概念
        '微服务', '分布式', '大数据', '人工智能', '机器学习', '深度学习',
        '自然语言处理', '计算机视觉', '大模型', '智能体', 'Agent',
        '容器化', '服务网格', '服务治理', '中间件', '负载均衡',
    ]

    @classmethod
    async def upload_resume_services(
            cls, request: Request, query_db: AsyncSession,
            file_path: str, file_name: str, current_user_id: int
    ) -> ResumeUploadResultModel:
        """
        上传简历并解析入库 service

        :param request: Request 对象
        :param query_db: orm 对象
        :param file_path: 文件路径
        :param file_name: 文件名
        :param current_user_id: 当前登录用户 ID
        :return: 简历上传结果
        """
        try:
            start_time = int(time.time() * 1000)
            current_time = int(time.time())

            # 1. 解析文档
            full_text = cls._parse_document(file_path, file_name)
            if not full_text:
                raise ServiceException(message='文档解析失败，无法提取文本内容')

            # 2. 提取结构化信息
            structured_info = await cls._extract_structured_info(full_text)

            # 3. 规则校验：自动检测异常字段
            abnormal_fields = cls._validate_parsed_fields(structured_info)

            # 4. 如果存在异常字段，调用LLM复核
            if abnormal_fields:
                logger.info(f'检测到{len(abnormal_fields)}个异常字段，启动LLM复核: {list(abnormal_fields.keys())}')
                await cls._llm_recheck_fields(full_text, abnormal_fields, structured_info)

            # 5. 生成标签
            tags = cls._generate_tags(structured_info)

            # 6. 生成简历UUID
            resume_uuid = str(uuid.uuid4())

            # 7. 构建DO对象
            resume_do = OaResume(
                resume_uuid=resume_uuid,
                file_name=file_name,
                name=structured_info.get('name', ''),
                gender=structured_info.get('gender', ''),
                age=structured_info.get('age', 0) or 0,
                birth_date=structured_info.get('birth_date', ''),
                phone=structured_info.get('phone', ''),
                email=structured_info.get('email', ''),
                education=structured_info.get('education', ''),
                major=structured_info.get('major', ''),
                school=structured_info.get('school', ''),
                graduation_date=structured_info.get('graduation_date', ''),
                work_years=structured_info.get('work_years', 0) or 0,
                current_company=structured_info.get('current_company', ''),
                current_position=structured_info.get('current_position', ''),
                degree=structured_info.get('degree', '') or '',
                school_system=structured_info.get('school_system', '') or '',
                study_form=structured_info.get('study_form', '') or '',
                id_card_number=structured_info.get('id_card_number', '') or '',
                id_card_address=structured_info.get('id_card_address', '') or '',
                technical_skills=json.dumps(structured_info.get('technical_skills', []), ensure_ascii=False),
                certifications=json.dumps(structured_info.get('certifications', []), ensure_ascii=False),
                tags=','.join(tags),
                full_text=full_text,
                parse_time=int(time.time() * 1000) - start_time,
                status=1,
                admin_id=current_user_id,
                create_time=current_time,
                update_time=current_time,
            )

            # 6. 入库
            await ResumeDao.add_resume_dao(query_db, resume_do)
            await query_db.flush()
            resume_id = resume_do.id

            # 7. 入库工作经历
            work_experiences = structured_info.get('work_experiences', [])
            for idx, work in enumerate(work_experiences):
                work_do = OaResumeWork(
                    resume_id=resume_id,
                    company=work.get('company', ''),
                    position=work.get('position', ''),
                    start_date=work.get('start_date', ''),
                    end_date=work.get('end_date', ''),
                    description=work.get('description', ''),
                    sort=idx,
                    create_time=current_time,
                )
                await ResumeDao.add_resume_work_dao(query_db, work_do)

            # 8. 入库项目经验
            project_experiences = structured_info.get('project_experiences', [])
            for idx, project in enumerate(project_experiences):
                project_do = OaResumeProject(
                    resume_id=resume_id,
                    project_name=project.get('project_name', ''),
                    role=project.get('role', ''),
                    start_date=project.get('start_date', ''),
                    end_date=project.get('end_date', ''),
                    description=project.get('description', ''),
                    technologies=json.dumps(project.get('technologies', []), ensure_ascii=False),
                    sort=idx,
                    create_time=current_time,
                )
                await ResumeDao.add_resume_project_dao(query_db, project_do)

            await query_db.commit()

            processing_time = int(time.time() * 1000) - start_time
            logger.info(f'简历解析入库成功: resume_id={resume_id}, name={structured_info.get("name")}, processing_time={processing_time}ms')

            return ResumeUploadResultModel(
                success=True,
                message='简历解析入库成功',
                resume_id=resume_id,
                resume_uuid=resume_uuid,
                file_name=file_name,
                processing_time=processing_time,
                name=structured_info.get('name', ''),
                education=structured_info.get('education', ''),
                major=structured_info.get('major', ''),
                age=structured_info.get('age', 0),
                work_years=structured_info.get('work_years', 0),
                technical_skills=structured_info.get('technical_skills', []),
                tags=tags,
            )

        except ServiceException:
            await query_db.rollback()
            raise
        except Exception as e:
            await query_db.rollback()
            logger.error(f'简历解析入库失败: {str(e)}')
            raise ServiceException(message=f'简历解析入库失败: {str(e)}')

    @classmethod
    async def get_resume_list_services(
            cls, query_db: AsyncSession, query_object: ResumePageQueryModel,
            user_id: int, where_conditions: list, is_page: bool = False
    ) -> PageModel | list[dict[str, Any]]:
        """
        获取简历列表信息 service

        :param query_db: orm 对象
        :param query_object: 查询参数对象
        :param user_id: 当前用户ID
        :param where_conditions: 查询条件列表
        :param is_page: 是否开启分页
        :return: 简历列表信息对象
        """
        resume_list_result = await ResumeDao.get_resume_list(
            query_db, query_object, where_conditions, is_page
        )

        return resume_list_result

    @classmethod
    async def build_query_conditions(cls, query_object: ResumePageQueryModel, user_id: int) -> list:
        """
        构建简历查询条件

        :param query_object: 查询参数对象
        :param user_id: 当前用户ID
        :return: 查询条件列表
        """
        return await ResumeDao.build_query_conditions(query_object, user_id)

    @classmethod
    async def get_resume_detail_services(cls, query_db: AsyncSession, resume_id: int) -> ResumeModel:
        """
        获取简历详细信息 service（含工作经历和项目经验）

        :param query_db: orm 对象
        :param resume_id: 简历 id
        :return: 简历详细信息
        """
        resume = await ResumeDao.get_resume_detail_by_id(query_db, resume_id)

        if not resume:
            return ResumeModel()

        # 获取工作经历
        work_list = await ResumeDao.get_work_experiences(query_db, resume_id)
        work_models = []
        for w in work_list:
            work_models.append(ResumeWorkExperienceModel(
                company=w.company,
                position=w.position,
                start_date=w.start_date,
                end_date=w.end_date,
                description=w.description,
            ))

        # 获取项目经验
        project_list = await ResumeDao.get_project_experiences(query_db, resume_id)
        project_models = []
        for p in project_list:
            try:
                technologies = json.loads(p.technologies) if p.technologies else []
            except:
                technologies = []
            project_models.append(ResumeProjectExperienceModel(
                project_name=p.project_name,
                role=p.role,
                start_date=p.start_date,
                end_date=p.end_date,
                description=p.description,
                technologies=technologies,
            ))

        # 解析JSON字段
        try:
            technical_skills = json.loads(resume.technical_skills) if resume.technical_skills else []
        except:
            technical_skills = []

        try:
            certifications = json.loads(resume.certifications) if resume.certifications else []
        except:
            certifications = []

        tags = resume.tags.split(',') if resume.tags else []

        result = ResumeModel(
            id=resume.id,
            resume_uuid=resume.resume_uuid,
            file_name=resume.file_name,
            name=resume.name,
            gender=resume.gender,
            age=resume.age,
            birth_date=resume.birth_date,
            phone=resume.phone,
            email=resume.email,
            education=resume.education,
            major=resume.major,
            school=resume.school,
            graduation_date=resume.graduation_date,
            work_years=resume.work_years,
            current_company=resume.current_company,
            current_position=resume.current_position,
            technical_skills=technical_skills,
            certifications=certifications,
            tags=tags,
            full_text=resume.full_text,
            parse_time=resume.parse_time,
            work_experiences=work_models,
            project_experiences=project_models,
            admin_id=resume.admin_id,
            create_time=resume.create_time,
            update_time=resume.update_time,
        )

        return result

    @classmethod
    async def delete_resume_services(
            cls, query_db: AsyncSession, resume_id: int, current_user_id: int
    ) -> CrudResponseModel:
        """
        删除简历 service（逻辑删除）

        :param query_db: orm 对象
        :param resume_id: 简历ID
        :param current_user_id: 当前登录用户 ID
        :return: 删除结果
        """
        resume = await ResumeDao.get_resume_detail_by_id(query_db, resume_id)
        if not resume:
            raise ServiceException(message='简历不存在')

        if resume.admin_id != current_user_id and current_user_id != 1:
            raise ServiceException(message='只有超级管理员和创建人才有权限操作')

        try:
            await ResumeDao.delete_resume_dao(query_db, resume_id)
            await query_db.commit()
            logger.info(f'删除简历成功，ID: {resume_id}')
            return CrudResponseModel(is_success=True, message='删除成功')
        except Exception as e:
            await query_db.rollback()
            logger.error(f'删除简历失败: {str(e)}')
            raise ServiceException(message=f'删除简历失败: {str(e)}')

    # ========== 简历智能问答（RAG） ==========

    @classmethod
    async def chat_with_resume_services(
        cls,
        query_db: AsyncSession,
        chat_req: ResumeChatRequestModel,
        user_id: int,
    ) -> AsyncGenerator[str, None]:
        """
        简历智能问答（RAG模式）

        :param query_db: orm对象
        :param chat_req: 对话请求对象
        :param user_id: 用户ID
        :return: SSE消息生成器
        """
        # 1. 获取模型配置
        ai_model = await AiModelDao.get_ai_model_detail_by_id(query_db, chat_req.model_id)
        model_config = AiModelModel(**CamelCaseUtil.transform_result(ai_model)) if ai_model else AiModelModel()
        if not model_config:
            raise ServiceException(message='模型不存在')

        # 2. 检索相关简历
        related_resumes = await cls._retrieve_related_resumes(
            query_db, chat_req.message, chat_req.top_k or 5
        )

        # 3. 构建简历上下文
        resume_context = cls._build_resume_context(related_resumes)

        # 4. 构建聊天配置
        session_id = chat_req.session_id or str(uuid.uuid4())
        chat_config = cls._build_resume_chat_config(
            model_config=model_config,
            resume_context=resume_context,
            user_id=user_id,
            session_id=session_id,
        )

        # 5. 流式返回
        async for chunk in cls._stream_resume_chat(
            chat_config=chat_config,
            message=chat_req.message,
            related_resumes=related_resumes,
            session_id=session_id,
        ):
            yield chunk

    # ========== 【第3步】MySQL检索 + 【第4步】上下文构建 ==========

    @classmethod
    async def _retrieve_related_resumes(
        cls, query_db: AsyncSession, query: str, top_k: int, force_all: bool = False
    ) -> list[ResumeModel]:
        """
        检索相关简历 - 基于关键词提取的分词检索

        :param query_db: orm对象
        :param query: 查询关键词
        :param top_k: 返回数量
        :param force_all: 是否强制返回所有简历（用于统计类查询）
        :return: 相关简历列表
        """
        # 1. 提取查询中的有效关键词
        keywords = cls._extract_query_keywords(query)
        
        # 2. 构建查询条件
        where_conditions = [OaResume.status == 1]
        
        if keywords and not force_all:
            # 有有效关键词时，用OR匹配
            or_conditions = []
            for kw in keywords:
                kw_pattern = f'%{kw}%'
                or_conditions.append(OaResume.name.like(kw_pattern))
                or_conditions.append(OaResume.full_text.like(kw_pattern))
                or_conditions.append(OaResume.current_company.like(kw_pattern))
                or_conditions.append(OaResume.current_position.like(kw_pattern))
                or_conditions.append(OaResume.major.like(kw_pattern))
                or_conditions.append(OaResume.school.like(kw_pattern))
                or_conditions.append(OaResume.education.like(kw_pattern))
                or_conditions.append(OaResume.technical_skills.like(kw_pattern))
                or_conditions.append(OaResume.tags.like(kw_pattern))
            if or_conditions:
                where_conditions.append(or_(*or_conditions))
        else:
            # 无有效关键词或force_all=True时，返回所有有效简历
            pass  # 只保留 status == 1 的条件
        
        # 3. 查询简历列表
        query_object = ResumePageQueryModel(
            page_num=1,
            page_size=top_k,
        )
        result = await ResumeDao.get_resume_list(
            query_db, query_object, where_conditions, is_page=True
        )

        resumes = []
        if result and hasattr(result, 'rows'):
            for row in result.rows:
                # 先把数据库行转为字典并预处理JSON字段，避免Pydantic验证失败
                row_dict = dict(row) if hasattr(row, 'keys') else row
                
                # 手动解析JSON字段为Python列表
                try:
                    ts = row_dict.get('technical_skills', '[]')
                    row_dict['technical_skills'] = json.loads(ts) if ts else []
                except (json.JSONDecodeError, TypeError):
                    row_dict['technical_skills'] = []
                
                try:
                    cert = row_dict.get('certifications', '[]')
                    row_dict['certifications'] = json.loads(cert) if cert else []
                except (json.JSONDecodeError, TypeError):
                    row_dict['certifications'] = []
                
                row_dict['tags'] = row_dict.get('tags', '').split(',') if row_dict.get('tags') else []
                
                # 转成ResumeModel（此时字段已经是正确的Python类型）
                resume = ResumeModel(**CamelCaseUtil.transform_result(row_dict))
                resumes.append(resume)

        return resumes

    # 有意义的简历相关词汇（用于识别有效关键词）
    _MEANINGFUL_WORDS = {
        # 学历
        '博士', '硕士', '研究生', '本科', '学士', '大专', '专科', '高中', '中专', '职高', '初中',
        # 职位/角色
        '工程师', '开发', '经理', '总监', '主管', '架构师', '负责人', '专员', '顾问',
        '前端', '后端', '全栈', '运维', '测试', '产品', '设计', '运营', '销售', '市场',
        'Java', 'Python', 'C++', 'Go', 'JavaScript', 'TypeScript', 'React', 'Vue',
        'Spring', 'SpringBoot', 'Docker', 'Kubernetes', 'MySQL', 'Redis', 'MongoDB',
        'Elasticsearch', 'Kafka', 'RabbitMQ', 'Hadoop', 'Spark', 'Flink', 'Linux',
        'Git', 'Maven', 'Gradle', 'Nginx', 'Tomcat', 'MyBatis', 'Hibernate',
        'Oracle', 'PostgreSQL', 'SQL Server', 'AWS', 'Azure',
        '微服务', '分布式', '大数据', '人工智能', '机器学习', '深度学习',
        '自然语言处理', '计算机视觉', 'HTML', 'CSS', 'Node.js', 'Django', 'Flask',
        'TensorFlow', 'PyTorch', 'Pandas', 'NumPy', 'Scikit-learn',
        # 公司相关
        '百度', '阿里', '腾讯', '字节', '美团', '京东', '滴滴', '小米', '华为',
        '网易', '快手', '拼多多', '携程', '去哪儿', '饿了么',
        # 通用有意义的词
        '简历', '候选人', '人才', '员工', '工作', '经验', '项目', '技能',
        '姓名', '年龄', '性别', '学校', '专业', '毕业', '学历',
        '年限', '年', '月', '薪资', '期望', '城市', '北京', '上海',
        '广州', '深圳', '杭州', '成都', '武汉', '西安', '南京',
        # 技术方向
        '前端开发', '后端开发', '移动端', 'iOS', 'Android', '小程序',
        'Web', 'API', '接口', '数据库', '缓存', '消息队列',
        '云原生', 'DevOps', 'CI/CD', '容器', '虚拟化',
        # 行业
        '互联网', '金融', '电商', '教育', '医疗', '游戏', '汽车',
        '人工智能', 'AI', '区块链', '物联网', '云计算',
    }

    # 停用词表 - 过滤无意义虚词
    _STOP_WORDS = {
        '系统', '现在', '有', '多少', '份', '的', '了', '在', '是', '我', '你', '他', '它', '们',
        '这', '那', '哪', '什么', '怎么', '为什么', '吗', '呢', '吧', '啊', '哦', '嗯',
        '个', '些', '过', '上', '下', '中', '里', '外', '前', '后', '大', '小', '好', '坏',
        '高', '低', '多', '少', '新', '老', '长', '短', '远', '近', '很', '非常', '比较',
        '都', '就', '也', '还', '但', '而', '或者', '以及', '及', '与', '和', '跟', '同',
        '为', '给', '对', '将', '被', '把', '让', '向', '到', '从', '关于', '对于', '由',
        '可以', '能', '会', '要', '想', '看', '知道', '请问', '请', '帮忙', '帮',
        '告诉', '说说', '讲', '描述', '介绍', '有没有', '有无', '是否',
        '一下', '一些', '一点', '看看', '查', '查询', '查一下',
        '搜索', '搜', '找', '找一下', '找出', '列出', '显示', '展示', '给我', '告诉我',
        '请给我', '麻烦', '谁', '哪位', '哪里',
        '哪些', '那种', '之类',
    }

    @classmethod
    def _extract_query_keywords(cls, query: str) -> list[str]:
        """
        从用户查询中提取有效关键词（无需外部依赖）
        
        策略：
        1. 先提取英文单词（技术词汇等）
        2. 清理中文，用有意义词表匹配
        3. 过滤停用词和单字
        4. 当查询过于宽泛时返回空列表（由调用方决定返回所有简历）
        """
        if not query:
            return []
        
        keywords = set()
        
        # 提取英文单词（技术词汇、公司名称等）
        english_words = re.findall(r'[a-zA-Z][a-zA-Z0-9\+\.\#\-]*', query)
        for word in english_words:
            if len(word) >= 2 and word.lower() not in {'or', 'and', 'the', 'is', 'are', 'was', 'were', 'be', 'been', 'being', 'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'shall', 'should', 'can', 'could', 'may', 'might', 'must'}:
                keywords.add(word)
        
        # 清理查询：去掉英文和标点，保留中文和数字
        chinese_query = re.sub(r'[a-zA-Z]', '', query)
        chinese_query = re.sub(r'[^\u4e00-\u9fa5\d]', '', chinese_query)
        
        # 策略1：用有意义词表匹配（优先匹配长的）
        if chinese_query:
            # 从长到短匹配有意义词
            meaningful_sorted = sorted(cls._MEANINGFUL_WORDS, key=len, reverse=True)
            remaining = chinese_query
            for word in meaningful_sorted:
                if word in remaining:
                    keywords.add(word)
                    remaining = remaining.replace(word, ' ')
            
            # 策略2：提取剩余部分中的2-4字组合，但要过滤无意义组合
            remaining = re.sub(r'[^\u4e00-\u9fa5]', '', remaining)
            if len(remaining) >= 2:
                # 只保留包含"实义词"的组合（有名词、动词倾向的字）
                # 简单策略：检查是否包含简历核心字段相关的字
                core_chars = set('简历候选人人才员工工作项目技能姓名年龄性别学校专业毕业学历年限年月薪资期望城市工程师开发经理')
                for length in range(4, 1, -1):
                    for i in range(len(remaining) - length + 1):
                        candidate = remaining[i:i+length]
                        # 过滤条件：不能全是停用词字、不能是纯数字、至少包含一个核心字
                        if (candidate not in cls._STOP_WORDS and 
                            not candidate.isdigit() and
                            any(c in core_chars for c in candidate)):
                            keywords.add(candidate)
        
        # 最终过滤
        result = []
        for kw in keywords:
            if len(kw) > 15:
                continue
            if kw.isdigit():
                continue
            # 如果词在停用词表里，去掉
            if kw in cls._STOP_WORDS:
                continue
            result.append(kw)
        
        # 按长度排序，长的优先
        result.sort(key=len, reverse=True)
        
        # 限制关键词数量
        return result[:10]

    @classmethod
    def _is_all_stop_chars(cls, text: str) -> bool:
        """检查文本是否全部由停用词中的字符组成"""
        all_stop_chars = set(''.join(cls._STOP_WORDS))
        return all(c in all_stop_chars for c in text)


    @classmethod
    def _build_resume_context(cls, resumes: list[ResumeModel]) -> str:
        """
        构建简历上下文文本

        :param resumes: 简历列表
        :return: 上下文文本
        """
        if not resumes:
            return '当前简历库中没有找到相关简历。'

        context_parts = ['以下是从简历库中检索到的相关简历信息：\n']

        for idx, resume in enumerate(resumes, 1):
            part = f'【简历 {idx}】\n'
            part += f'姓名：{resume.name or "未知"}\n'
            part += f'性别：{resume.gender or "未知"}\n'
            part += f'年龄：{resume.age or "未知"}岁\n'
            part += f'学历：{resume.education or "未知"}\n'
            part += f'专业：{resume.major or "未知"}\n'
            part += f'毕业院校：{resume.school or "未知"}\n'
            part += f'毕业时间：{resume.graduation_date or "未知"}\n'
            part += f'工作年限：{resume.work_years or 0}年\n'
            part += f'当前公司：{resume.current_company or "未知"}\n'
            part += f'当前职位：{resume.current_position or "未知"}\n'
            # 投标文件来源信息
            if resume.source_type == 2 and resume.source_name:
                part += f'来源：投标文件《{resume.source_name}》\n'
            part += f'技术技能：{", ".join(resume.technical_skills or []) if resume.technical_skills else "未知"}\n'
            if resume.certifications:
                part += f'证书：{", ".join(resume.certifications)}\n'

            # 工作经历
            if resume.work_experiences:
                part += '\n工作经历：\n'
                for work in resume.work_experiences:
                    part += f'  - {work.company or "未知"} | {work.position or "未知"} | {work.start_date or ""} ~ {work.end_date or "至今"}\n'
                    if work.description:
                        part += f'    描述：{work.description[:200]}\n'

            # 项目经验
            if resume.project_experiences:
                part += '\n项目经验：\n'
                for proj in resume.project_experiences:
                    part += f'  - {proj.project_name or "未知"} | 角色：{proj.role or "未知"}\n'
                    if proj.technologies:
                        part += f'    技术：{", ".join(proj.technologies)}\n'
                    if proj.description:
                        part += f'    描述：{proj.description[:200]}\n'

            context_parts.append(part)

        context_parts.append('\n请根据以上简历信息回答用户的问题。如果简历信息不足以回答问题，请如实说明。')
        return '\n'.join(context_parts)

    @classmethod
    def _build_resume_chat_config(
        cls,
        model_config: AiModelModel,
        resume_context: str,
        user_id: int,
        session_id: str,
    ) -> dict:
        """
        构建简历问答配置（litellm 直接调用）

        :param model_config: 模型配置
        :param resume_context: 简历上下文
        :param user_id: 用户ID
        :param session_id: 会话ID
        :return: 配置字典
        """
        real_api_key = CryptoUtil.decrypt(model_config.api_key)

        system_prompt = f"""你是一个专业的简历分析助手，擅长根据简历信息回答招聘相关问题。

{resume_context}

回答要求：
1. 只基于提供的简历信息回答，不要编造信息
2. 如果简历信息不足以回答问题，请如实说明
3. 回答要简洁、准确、有条理
4. 涉及多人对比时，要清晰列出每个人的情况
5. 可以适当总结和提炼关键信息"""

        return {
            'api_key': real_api_key,
            'base_url': model_config.base_url,
            'model': model_config.model_code,
            'temperature': model_config.temperature or 0.7,
            'max_tokens': model_config.max_tokens or 4096,
            'system_prompt': system_prompt,
            'user_id': user_id,
            'session_id': session_id,
        }

    @classmethod
    async def _stream_resume_chat(
        cls,
        chat_config: dict,
        message: str,
        related_resumes: list[ResumeModel],
        session_id: str,
    ) -> AsyncGenerator[str, None]:
        """
        流式输出简历问答结果（litellm 直接调用）

        :param chat_config: 聊天配置字典
        :param message: 用户问题
        :param related_resumes: 关联简历列表
        :param session_id: 会话ID
        :return: SSE消息生成器
        """
        try:
            # 先发送会话信息
            yield json.dumps({'session_id': session_id, 'type': 'meta'}) + '\n'

            # 发送关联简历信息
            related_resume_data = []
            for r in related_resumes:
                related_resume_data.append(
                    ResumeChatRelatedResumeModel(
                        id=r.id,
                        name=r.name,
                        position=r.current_position,
                        education=r.education,
                        work_years=r.work_years,
                        school=r.school,
                        skills=r.technical_skills,
                    ).model_dump(by_alias=True)
                )
            yield json.dumps({'related_resumes': related_resume_data, 'type': 'related'}) + '\n'

            # 使用 litellm 直接流式调用，避免 agno 的 developer role 问题
            messages = [
                {'role': 'system', 'content': chat_config['system_prompt']},
                {'role': 'user', 'content': message},
            ]

            def _call_litellm():
                return litellm.completion(
                    model=chat_config['model'],
                    messages=messages,
                    api_key=chat_config['api_key'],
                    base_url=chat_config['base_url'],
                    temperature=chat_config['temperature'],
                    max_tokens=chat_config['max_tokens'],
                    stream=True,
                )

            # 使用线程池执行同步调用，避免阻塞async事件循环
            loop = asyncio.get_event_loop()
            response = await asyncio.wait_for(
                loop.run_in_executor(None, _call_litellm),
                timeout=35,
            )

            full_response = ''
            for chunk in response:
                content = chunk.choices[0].delta.content if chunk and chunk.choices else None
                if content:
                    full_response += content
                    yield json.dumps({'content': content, 'type': 'content'}) + '\n'

            # 完成标记
            yield json.dumps({'type': 'done', 'full_response': full_response}) + '\n'

        except Exception as e:
            logger.error(f'简历问答流式调用失败: {str(e)}')
            yield json.dumps({'error': str(e), 'type': 'error'}) + '\n'

    # ========== 文档解析工具方法 ==========

    @classmethod
    def _is_garbled_text(cls, text: str) -> bool:
        """
        检测文本是否为乱码或不可读
        多层检测：
        1. 基础统计检测（字符比例）
        2. Unicode替换字符检测（�字符过多说明解码失败）
        3. 语义关键词检测（文本中必须包含至少一个简历核心词）
        任一检测不通过即判定为乱码
        """
        if not text or len(text.strip()) < 10:
            logger.warning(f'_is_garbled_text: 文本为空或太短（{len(text.strip())}字符），判定为乱码')
            return True
        
        total_chars = len(text)
        
        # === 检测1: Unicode替换字符 ===
        # U+FFFD � 是解码失败时产生的替换字符，出现说明有严重编码问题
        replacement_chars = text.count('\ufffd')
        replacement_ratio = replacement_chars / total_chars if total_chars > 0 else 0
        if replacement_chars >= 3 and replacement_ratio > 0.01:
            logger.warning(f'Unicode替换字符{replacement_chars}个(占比{replacement_ratio:.2%})，编码解码失败，判定为乱码')
            return True
        
        # === 检测2: 统计指标 ===
        readable_chars = sum(1 for c in text if c.isprintable() and c not in '\ufffd\u0000-\u001f\u007f-\u009f')
        chinese_chars = sum(1 for c in text if '\u4e00' <= c <= '\u9fff')
        english_chars = sum(1 for c in text if c.isalpha() and c.isascii())
        
        readable_ratio = readable_chars / total_chars if total_chars > 0 else 0
        meaningful_ratio = (chinese_chars + english_chars) / total_chars if total_chars > 0 else 0
        
        if readable_ratio < 0.5:
            logger.warning(f'可读字符比例{readable_ratio:.2f} < 0.5，判定为乱码')
            return True
        if meaningful_ratio < 0.3:
            logger.warning(f'有意义字符比例{meaningful_ratio:.2f} < 0.3，判定为乱码')
            return True
        
        # === 检测3: 语义关键词检测（最关键）===
        # 正常简历中应该包含这些核心词之一，完全没有说明内容不可读
        # 用正则匹配，避免大小写/全半角问题
        core_keywords = [
            '姓名', '性别', '男', '女', '出生', '年龄', '学历', '本科', '硕士', '博士',
            '大专', '毕业', '学校', '院校', '专业', '工作', '经历', '经验', '职位',
            '岗位', '公司', '项目', '技能', '电话', '手机', '邮箱', '邮箱地址',
            '求职', '意向', '自我介绍', '个人优势', '自我评价',
        ]
        # 同时匹配常见英文简历词
        en_keywords = ['name', 'gender', 'age', 'education', 'experience', 'skill', 'phone', 'email']
        
        text_lower = text.lower()
        found_keywords = []
        for kw in core_keywords:
            if kw in text:
                found_keywords.append(kw)
        for kw in en_keywords:
            if kw in text_lower:
                found_keywords.append(kw)
        
        logger.info(f'_is_garbled_text 关键词检测: 找到{len(found_keywords)}个核心词: {found_keywords[:5]}')
        
        # 如果500字符以上的文本中，一个简历核心词都找不到，说明是乱码
        if total_chars > 200 and len(found_keywords) < 2:
            logger.warning(f'文本{total_chars}字符中仅找到{len(found_keywords)}个简历核心词，判定为乱码')
            return True
        if total_chars > 50 and len(found_keywords) == 0:
            logger.warning(f'文本{total_chars}字符中未找到任何简历核心词，判定为乱码')
            return True
        
        logger.info(f'_is_garbled_text 最终判定: 正常文本')
        return False

    @classmethod
    def _ocr_pdf_with_rapidocr(cls, file_path: str) -> str:
        """
        使用本地 RapidOCR 引擎识别图片型PDF中的文字
        无需调用外部 LLM Vision API，免费且速度快
        """
        try:
            import fitz
            from rapidocr_onnxruntime import RapidOCR
        except ImportError:
            raise ServiceException(message='PDF OCR 功能需要安装依赖：pip install PyMuPDF rapidocr-onnxruntime')

        # 初始化OCR引擎（只初始化一次，加缓存）
        if not hasattr(cls, '_ocr_engine'):
            cls._ocr_engine = RapidOCR()
        engine = cls._ocr_engine

        try:
            doc = fitz.open(file_path)
            all_text = []

            logger.info(f'开始使用 RapidOCR 识别 PDF，共 {doc.page_count} 页')

            # 将每一页转成图片并识别（简历核心信息通常在前5页）
            max_pages = min(doc.page_count, 5)  # 限制最多5页，确保教育背景等信息被完整识别
            for page_num in range(max_pages):
                page = doc[page_num]
                # 适度提高分辨率，平衡清晰度与速度
                mat = fitz.Matrix(1.5, 1.5)
                pix = page.get_pixmap(matrix=mat)
                img_bytes = pix.tobytes('png')

                # RapidOCR 识别
                result, _ = engine(img_bytes)
                if result:
                    # result 格式: [[box, text, confidence], ...]
                    page_text = '\n'.join([item[1] for item in result])
                    all_text.append(page_text)
                    logger.info(f'RapidOCR 识别第{page_num+1}页成功，识别到 {len(page_text)} 字符')
                else:
                    logger.warning(f'RapidOCR 识别第{page_num+1}页未识别到文字')

            doc.close()

            full_text = '\n'.join(all_text)
            if full_text.strip():
                logger.info(f'PDF OCR识别完成，共识别 {len(full_text)} 字符')
                return full_text
            else:
                logger.error('PDF OCR识别结果为空')
                raise ServiceException(message='PDF OCR识别失败：未能从图片中识别出文字')

        except ServiceException:
            raise
        except Exception as e:
            logger.error(f'PDF OCR识别失败: {str(e)}')
            raise ServiceException(message=f'PDF OCR识别失败: {str(e)}')

    @classmethod
    def _parse_document(cls, file_path: str, file_name: str) -> str:
        """
        解析文档为纯文本
        支持PDF（含图片型PDF自动OCR）、Word、纯文本
        """
        ext = file_name.lower().split('.')[-1] if '.' in file_name else ''

        if ext in ['docx', 'doc']:
            # Word 文档
            try:
                result = ResumeParser.parse_word_resume(file_path)
                # ResumeParser.parse_word_resume 返回的是结构化信息字典
                # 需要读取原始文本，这里用 docx 直接读取
                from docx import Document
                doc = Document(file_path)
                text = '\n'.join([para.text for para in doc.paragraphs])
                for table in doc.tables:
                    for row in table.rows:
                        row_text = ' '.join([cell.text for cell in row.cells])
                        text += '\n' + row_text
                return text
            except Exception as e:
                logger.error(f'Word 解析失败: {str(e)}')
                raise ServiceException(message=f'Word 解析失败: {str(e)}')

        elif ext == 'pdf':
            # PDF 文档 - 用 PyMuPDF 提取文本（比 pdfplumber 对中文支持更好）
            try:
                import fitz
                doc = fitz.open(file_path)
                text = ''
                for page_num in range(doc.page_count):
                    text += doc[page_num].get_text()
                doc.close()

                # 检测提取的文本质量
                if cls._is_garbled_text(text):
                    logger.warning(f'PDF文本提取结果质量不佳（长度{len(text)}），自动启用OCR识别')
                    text = cls._ocr_pdf_with_rapidocr(file_path)

                return text
            except ServiceException:
                raise
            except Exception as e:
                logger.error(f'PDF 解析失败: {str(e)}')
                raise ServiceException(message=f'PDF 解析失败: {str(e)}')

        elif ext in ['txt', 'text']:
            # 纯文本
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            except UnicodeDecodeError:
                with open(file_path, 'r', encoding='gbk') as f:
                    return f.read()

        else:
            raise ServiceException(message=f'不支持的文件类型: {ext}')

    # ========== 字段校验规则配置 ==========
    # 校验规则：定义每个关键字段的校验标准
    _FIELD_VALIDATION_RULES = {
        'current_company': {
            'max_length': 40,
            'invalid_prefixes': ['公司介绍', '项目说明', '核心贡献', '基于', '通过', '面向', '负责', '使用', '类似', '实现', '完成', '开发', '设计', '搭建', '优化', '提升'],
            'required_keywords': ['公司', '集团', '科技', '信息', '网络', '有限', '股份', '研究院'],
            'min_length': 2,
        },
        'current_position': {
            'max_length': 30,
            'invalid_prefixes': ['项目说明', '核心贡献', '公司介绍', '工作职责', '工作内容', '基于', '通过', '面向', '负责', '使用', '类似', '实现', '完成', '开发', '设计', '搭建', '优化', '提升', '完成', '参与', '主导', '担任'],
            'required_keywords': ['工程师', '开发', '运维', '测试', '架构', '经理', '主管', '总监', '负责人', '专家', '顾问', '全栈', '前端', '后端', '工程师', '员', '师', '手', '长', 'Java', 'Python', 'Go'],
            'min_length': 2,
        },
        'school': {
            'max_length': 50,
            'invalid_prefixes': ['专业', '学历', '毕业', '本科', '硕士', '博士', '大专'],
            'required_keywords': ['大学', '学院', '学校', '研究院', '研究所', '高专', '职校'],
            'min_length': 2,
        },
        'major': {
            'max_length': 40,
            'invalid_prefixes': ['学校', '院校', '大学', '学院', '毕业'],
            'required_keywords': [],  # 专业名称多样，不强求关键词
            'min_length': 2,
        },
        'name': {
            'max_length': 10,
            'invalid_prefixes': ['姓名', '名字', '候选人', '应聘'],
            'required_keywords': [],
            'min_length': 2,
        },
    }

    @classmethod
    def _validate_parsed_fields(cls, info: dict[str, Any]) -> dict[str, str]:
        """
        规则引擎：校验解析后的关键字段值是否异常
        返回异常字段字典：{字段名: 异常原因}
        不依赖LLM主观判断，完全基于客观规则
        """
        abnormal_fields = {}

        for field, rules in cls._FIELD_VALIDATION_RULES.items():
            value = info.get(field)
            if not value or not isinstance(value, str):
                continue

            value_stripped = value.strip()
            if not value_stripped:
                continue

            # 规则1: 最小长度检查
            if len(value_stripped) < rules.get('min_length', 1):
                abnormal_fields[field] = f'长度过短({len(value_stripped)}字符)'
                continue

            # 规则2: 最大长度检查
            max_len = rules.get('max_length', 100)
            if len(value_stripped) > max_len:
                # 超长文本大概率是描述而非实体名
                abnormal_fields[field] = f'超长文本({len(value_stripped)}字符，超过{max_len})'
                continue

            # 规则3: 黑名单前缀检查
            invalid_prefixes = rules.get('invalid_prefixes', [])
            for prefix in invalid_prefixes:
                if value_stripped.startswith(prefix):
                    abnormal_fields[field] = f'包含描述性前缀"{prefix}"'
                    break
            if field in abnormal_fields:
                continue

            # 规则4: 必填关键词检查
            required_keywords = rules.get('required_keywords', [])
            if required_keywords:
                has_keyword = any(kw in value_stripped for kw in required_keywords)
                if not has_keyword:
                    abnormal_fields[field] = f'不含实体关键词(如{required_keywords[:3]})'
                    continue

            # 规则5: 描述性动词检查（针对职位名）
            if field == 'current_position':
                desc_verbs = ['负责', '基于', '通过', '面向', '使用', '实现', '完成', '开发', '设计', '搭建', '优化', '提升']
                if len(value_stripped) > 15 and any(verb in value_stripped for verb in desc_verbs):
                    # 再检查是否含职位关键词
                    pos_keywords = ['工程师', '开发', '运维', '测试', '架构', '经理', '主管', '总监', '负责人', '专家', '顾问', '全栈', '前端', '后端']
                    if not any(kw in value_stripped for kw in pos_keywords):
                        abnormal_fields[field] = '描述性文本（含动词且无职位关键词）'
                        continue

        if abnormal_fields:
            logger.warning(f'规则校验发现异常字段: {abnormal_fields}')
        else:
            logger.info('规则校验通过，所有关键字段值正常')

        return abnormal_fields

    @classmethod
    async def _llm_recheck_fields(cls, text: str, abnormal_fields: dict[str, str], info: dict[str, Any]) -> None:
        """
        LLM复核：仅对规则标异常的字段重新精确提取
        传入原始简历文本，要求LLM仅从原文中精确提取指定字段
        """
        if not abnormal_fields:
            return

        fields_to_recheck = list(abnormal_fields.keys())

        # 构建复核Prompt
        field_labels = {
            'current_company': '当前/最近公司名称',
            'current_position': '当前/最近职位名称',
            'school': '毕业院校',
            'major': '专业',
            'name': '姓名',
        }

        abnormal_desc = []
        for field, reason in abnormal_fields.items():
            current_value = info.get(field, '')
            label = field_labels.get(field, field)
            abnormal_desc.append(f'- {label}: 当前值"{current_value}"，异常原因：{reason}')

        fields_json = []
        for field in fields_to_recheck:
            fields_json.append(f'"{field}": "正确值或null"')

        prompt = f"""你是一个严谨的简历信息提取助手。以下是从简历中提取的字段值被规则引擎判定为异常，请你从原始简历文本中重新精确提取这些字段的正确值。

【异常字段及原因】
""" + '\n'.join(abnormal_desc) + f"""

【原始简历文本】
{text[:8000]}
【简历文本结束】

【提取要求——严格遵守】
1. 只提取异常字段，其他字段忽略
2. 必须从原始简历文本中找到确切证据，不得猜测或编造
3. 公司名称只要公司名本身，不要带"公司介绍："等前缀，不要带业务描述
4. 职位名称只要职位名本身，不要带"负责""基于""通过"等动词开头
5. 如果简历中确实没有该信息，返回null
6. 只输出JSON，不要解释：

{{""" + ',\n'.join(fields_json) + """}}"""

        try:
            api_key = AiResumeParserConfig.ai_resume_parser_api_key
            base_url = AiResumeParserConfig.ai_resume_parser_base_url
            model = AiResumeParserConfig.ai_resume_parser_model
            temperature = AiResumeParserConfig.ai_resume_parser_temperature or 0.1
            max_tokens = AiResumeParserConfig.ai_resume_parser_max_tokens or 4096

            if '/' in model and not model.startswith(('openai/', 'anthropic/', 'azure/', 'cohere/', 'groq/', 'ollama/', 'vertex_ai/')):
                model = f'openai/{model}'

            try:
                response = await litellm.acompletion(
                    model=model,
                    messages=[
                        {'role': 'system', 'content': '你是一个严谨的简历信息提取助手。只从原文提取，不猜测，不编造。'},
                        {'role': 'user', 'content': prompt},
                    ],
                    api_key=api_key,
                    base_url=base_url,
                    temperature=temperature,
                    max_tokens=max_tokens,
                    timeout=25,
                )
            except Exception as e:
                logger.error(f'LLM复核失败: {str(e)}')
                # LLM复核失败时，清空异常字段值，避免错误数据入库
                for field in abnormal_fields:
                    info[field] = ''
                return

            content = response.choices[0].message.content if response and response.choices else None
            if not content:
                logger.warning('LLM复核返回为空')
                return

            result = cls._parse_llm_response(content)
            if result and isinstance(result, dict):
                # 用LLM复核结果覆盖异常值（需要再次校验）
                for field in fields_to_recheck:
                    new_value = result.get(field)
                    if new_value and isinstance(new_value, str) and new_value.strip() and new_value.strip().lower() != 'null':
                        new_value_stripped = new_value.strip()
                        # 简单校验LLM返回的新值是否仍然异常
                        is_still_abnormal = False
                        rules = cls._FIELD_VALIDATION_RULES.get(field, {})
                        # 检查黑名单前缀
                        for prefix in rules.get('invalid_prefixes', []):
                            if new_value_stripped.startswith(prefix):
                                is_still_abnormal = True
                                break
                        # 检查超长
                        if not is_still_abnormal and len(new_value_stripped) > rules.get('max_length', 100):
                            is_still_abnormal = True

                        if not is_still_abnormal:
                            logger.info(f'LLM复核修正 {field}: "{info.get(field)}" → "{new_value_stripped}"')
                            info[field] = new_value_stripped
                        else:
                            logger.warning(f'LLM复核后 {field} 仍异常: "{new_value_stripped}"，清空')
                            info[field] = ''
                    else:
                        # LLM返回null或空，清空原异常值
                        logger.info(f'LLM复核建议清空 {field}: "{info.get(field)}"')
                        info[field] = ''

        except Exception as e:
            logger.error(f'LLM复核失败: {str(e)}')
            # LLM复核失败时，清空异常字段值，避免错误数据入库
            for field in abnormal_fields:
                info[field] = ''

    # ========== 结构化信息提取 ==========

    @classmethod
    def _preprocess_text(cls, text: str) -> str:
        """
        预处理简历文本：过滤噪声字符串、清理描述性前缀
        减少OCR垃圾文本和无关描述对LLM/规则提取的干扰
        """
        lines = text.split('\n')
        cleaned_lines = []

        for line in lines:
            line_stripped = line.strip()
            if not line_stripped:
                continue

            # 1. 过滤明显为噪声的字符串（随机字符序列、太长且无意义的行）
            # 匹配：只包含字母数字且长度>20，无中文字符，无空格的字符串（如1befb752dadcfbb...）
            if re.match(r'^[a-zA-Z0-9~!@#$%^&*_+=\-]{20,}$', line_stripped):
                # 如果是 Base64/UUID/加密哈希特征，跳过
                if re.search(r'^[a-zA-Z0-9]{20,}[~!@#$%^&*_+=\-]*$', line_stripped):
                    continue

            # 2. 过滤只包含"~"符号和字母数字的噪声行
            if re.match(r'^[~\-—\s]*$', line_stripped):
                continue

            # 3. 过滤数字行（如"1.""2."序号前有只含数字/标点的行）
            if re.match(r'^[\d\s\W]+$', line_stripped) and len(line_stripped) <= 5:
                continue

            # 4. 过滤描述性前缀行（这些行容易被LLM误提取为职位/公司名称）
            # 如 "公司介绍：xxx"、"项目说明：xxx"、"核心贡献：xxx"
            # 保留区块标题（如"公司介绍"后面没有冒号），但过滤描述内容
            desc_prefixes = [
                '公司介绍：', '项目说明：', '核心贡献：', '工作职责：',
                '工作内容：', '岗位职责：', '工作描述：', '工作业绩：',
                '项目描述：', '项目职责：', '项目业绩：', '个人优势：',
                '自我评价：', '工作成果：', '业绩描述：', '技术框架：',
            ]
            # 如果整行以这些前缀开头且内容较长（>20字符），说明是描述性内容，过滤掉
            for prefix in desc_prefixes:
                if line_stripped.startswith(prefix) and len(line_stripped) > len(prefix) + 10:
                    # 但保留"公司介绍："后面紧跟公司名的情况（如"公司介绍：杭州橙风信息科技有限公司"）
                    # 如果公司介绍后面只有公司名（不含逗号、不含"基于""通过"等动词），保留
                    after_prefix = line_stripped[len(prefix):].strip()
                    # 检查是否只有公司名（短且不含描述性词汇）
                    desc_indicators = ['基于', '通过', '面向', '负责', '使用', '类似', '实现', '完成', '优化', '提升', '搭建', '设计', '开发']
                    if any(d in after_prefix for d in desc_indicators) or len(after_prefix) > 25:
                        continue  # 是描述性内容，过滤掉
                    break  # 是公司名，保留

            cleaned_lines.append(line)

        return '\n'.join(cleaned_lines)

    @classmethod
    async def _extract_structured_info(cls, text: str) -> dict[str, Any]:
        """
        从简历文本中提取结构化信息
        === 以大模型LLM为主，规则为辅 ===
        1. 首先用 LLM 完整解析所有字段（这是主要方式）
        2. 用学信网/身份证的规则提取结果覆盖 LLM 可能出错的字段
        3. LLM 失败时才降级到规则提取
        """
        # 预处理：过滤噪声文本
        text = cls._preprocess_text(text)

        info = {
            'name': None,
            'gender': None,
            'age': None,
            'birth_date': None,
            'phone': None,
            'email': None,
            'education': None,
            'major': None,
            'school': None,
            'graduation_date': None,
            'work_years': None,
            'current_company': None,
            'current_position': None,
            'degree': None,
            'school_system': None,
            'study_form': None,
            'id_card_number': None,
            'id_card_address': None,
            'work_experiences': [],
            'project_experiences': [],
            'technical_skills': [],
            'certifications': [],
            '_full_text': text,
        }

        llm_success = False

        # === 第一步：优先使用 LLM 完整解析（主要方式） ===
        if AiResumeParserConfig.ai_resume_parser_enabled:
            try:
                llm_result = await cls._extract_by_llm_full(text)
                if llm_result and isinstance(llm_result, dict):
                    cls._sanitize_llm_result(llm_result)
                    # LLM 结果作为主要数据源（使用 _is_valid_value 过滤垃圾值如 "null" 字符串）
                    for key in info.keys():
                        if key in llm_result and llm_result[key] is not None and llm_result[key] != '':
                            llm_val = llm_result[key]
                            # 对简单字段做有效性校验，过滤 "null" "None" "未知" 等垃圾值
                            if isinstance(llm_val, str):
                                val_stripped = llm_val.strip()
                                if val_stripped.lower() in ('null', 'none', 'n/a', '未知', '无', '暂无', '未提供', '未填写', '-', '—'):
                                    continue
                            info[key] = llm_val
                    logger.info(f'LLM 解析完成: 姓名={info.get("name")}, 学历={info.get("education")}, 专业={info.get("major")}')
                    llm_success = True
            except asyncio.TimeoutError:
                logger.warning('LLM 解析超时，降级到规则提取')
            except Exception as e:
                logger.warning(f'LLM 解析失败，降级到规则提取: {str(e)}')

        # === 第二步：学信网/身份证规则提取（补充或覆盖） ===
        # 这些来源的数据更准确，可以覆盖 LLM 结果
        xuexin_info = {}
        cls._extract_from_xuexin(text, xuexin_info)
        id_card_info = {}
        cls._extract_from_id_card(text, id_card_info)

        # 学信网/身份证提取的结果可以覆盖 LLM 结果（这些更准确）
        for key in ['name', 'gender', 'birth_date', 'education', 'major', 'school', 'graduation_date', 'degree', 'school_system', 'study_form']:
            if xuexin_info.get(key):
                info[key] = xuexin_info[key]

        for key in ['name', 'gender', 'birth_date', 'age', 'id_card_number', 'id_card_address']:
            if id_card_info.get(key):
                info[key] = id_card_info[key]

        # 根据身份证的出生日期重新计算年龄（最准确）
        if info.get('birth_date') and not info.get('age'):
            try:
                import time
                birth_year = int(info['birth_date'][:4])
                info['age'] = time.localtime().tm_year - birth_year
            except:
                pass

        # === 第三步：如果 LLM 完全失败，用规则提取补充所有空字段 ===
        if not llm_success:
            cls._extract_by_rules(text, info)
        
        # === 第三步补充：即使LLM成功，也用规则补充phone/email等常缺失的字段 ===
        # LLM经常漏提取phone/email，用规则兜底
        if not info.get('phone'):
            phone_match = re.search(r'(?:联系\s*方式|电话|手机|联系电话|联系方式)\s*[：:\s]*\n?\s*(1[3-9]\d{9})', text)
            if not phone_match:
                phone_match = re.search(r'(?<!\d)(1[3-9]\d{9})(?!\d)', text)
            if phone_match:
                info['phone'] = phone_match.group(1)
                logger.info(f'规则补充提取手机号: {info["phone"]}')
        
        if not info.get('email'):
            email_match = re.search(r'(?:邮\s*箱|电子邮箱|Email|E-mail)\s*[：:\s]*\n?\s*([a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,})', text)
            if not email_match:
                email_match = re.search(r'([a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,})', text)
            if email_match:
                info['email'] = email_match.group(1)
                logger.info(f'规则补充提取邮箱: {info["email"]}')

        # === 后处理：清洗和纠错 ===
        cls._post_process_info(info)

        # === 最终补全：从身份证号反算出生日期和年龄 ===
        cls._fill_birth_age_from_id(info)

        return info

    @classmethod
    async def _extract_by_llm_full(cls, text: str) -> dict:
        """调用 LLM 完整解析简历文本，返回结构化字典"""
        # 检查配置
        if not AiResumeParserConfig.ai_resume_parser_enabled:
            return {}

        # 检查 api_key 和 model 配置
        api_key = getattr(AiResumeParserConfig, "ai_resume_parser_api_key", None)
        base_url = getattr(AiResumeParserConfig, "ai_resume_parser_base_url", None)
        model = getattr(AiResumeParserConfig, "ai_resume_parser_model", None)
        timeout = getattr(AiResumeParserConfig, "ai_resume_parser_timeout", 120)

        if not model or not api_key:
            logger.warning(f"LLM 配置缺失: model={model}, api_key={bool(api_key)}")
            return {}

        # 模型前缀处理（OpenAI 兼容接口需要加 openai/ 前缀，litellm 才能识别）
        if '/' in model and not model.startswith(('openai/', 'anthropic/', 'azure/', 'cohere/', 'groq/', 'ollama/', 'vertex_ai/')):
            model = f'openai/{model}'

        # 构建 Prompt
        prompt = cls._build_full_parse_prompt(text)

        # 调用 LLM
        try:
            import litellm
            # 关闭不必要的日志
            litellm.set_verbose = False
            litellm.suppress_debug_info = True

            # 准备 kwargs
            kwargs = {
                "model": model,
                "api_key": api_key,
                "messages": [
                    {"role": "user", "content": prompt}
                ],
                "temperature": 0.1,
                "max_tokens": 4096,
                "timeout": timeout,
            }

            if base_url:
                kwargs["base_url"] = base_url

            # 调用 LLM (非流式)
            response = await litellm.acompletion(**kwargs)
            raw_response = response.choices[0].message.content

            # 解析 JSON
            parsed = cls._parse_llm_response(raw_response)
            if parsed and isinstance(parsed, dict):
                return parsed
            return {}

        except Exception as e:
            logger.warning(f"LLM 调用失败: {type(e).__name__}: {str(e)[:200]}")
            return {}

    @classmethod
    def _build_full_parse_prompt(cls, text: str) -> str:
        """
        构建 LLM 解析 Prompt - 针对学信网+身份证截图优化
        """
        max_len = 25000
        truncated_text = text[:max_len] if len(text) > max_len else text

        NL = chr(10)

        prompt_lines = [
            '你是一个专业的简历信息提取专家。请从以下OCR识别后的简历文本中，精准提取结构化信息。',
            '',
            '【数据来源优先级】',
            '1. 最高优先级：学信网截图内容（字段最完整、最准确）',
            '2. 次高优先级：身份证截图内容（姓名、性别、身份证号、出生日期最准确）',
            '3. 低优先级：简历正文文字（仅作补充，如与学信网/身份证冲突，以学信网/身份证为准）',
            '',
            '【学信网字段提取指南】',
            '学信网截图经OCR后，字段通常以"字段名 值"的格式出现，例如：',
            '  姓名: 何超',
            '  性别: 男',
            '  出生日期: 1998年04月04日',
            '  学校名称: 扬州大学广陵学院',
            '  专业: 信息与计算科学',
            '  层次: 本科',
            '  学制: 4',
            '  学习形式: 普通全日制',
            '  入学日期: 2016年09月01日',
            '  毕业日期: 2020年06月22日',
            '  证书编号: 139871202005900161',
            '',
            '【身份证字段提取指南】',
            '身份证截图经OCR后，关键字段格式：',
            '  姓名 何超',
            '  性别 男',
            '  出生 1998年4月4日',
            '  公民身份号码 342626199804044912',
            '  住址 南京市鼓楼区金陵新一村23幢三单元301室',
            '',
            '【字段映射规则（严格遵循）】',
            '- name: 姓名，从学信网"姓名"或身份证"姓名"提取。注意：不要将"性别""民族""年龄"等字段名误认为姓名',
            '- gender: 性别，"男"或"女"',
            '- birth_date: 出生日期，格式YYYY-MM-DD。可从学信网"出生日期"或身份证"出生"提取，也可从身份证号第7-14位推导',
            '- age: 年龄数字，当前年份(2026) - 出生年份。如1998年出生则age=28',
            '- phone: 联系电话/手机号，从简历正文中提取11位手机号',
            '- email: 邮箱地址，从简历正文中提取',
            '- education: 学历层次，从学信网"层次"提取。值必须是：博士/硕士/本科/大专/专科/高中/中专',
            '- major: 专业名称，从学信网"专业"提取。如"信息与计算科学"、"计算机科学与技术"。注意：不要提取"学制""学习形式"等字段的值，不要将"四年制本科学习"作为专业',
            '- school: 学校名称，从学信网"学校名称"提取。如"扬州大学广陵学院"',
            '- graduation_date: 毕业时间，格式YYYY-MM。从学信网"毕业日期"提取，如2020年06月→"2020-06"',
            '- school_system: 学制，从学信网"学制"提取。如"4"→"4年"、"四年"',
            '- study_form: 学习形式，从学信网"学习形式"提取。如"普通全日制"、"全日制"',
            '- id_card_number: 身份证号，18位数字。从身份证"公民身份号码"后面提取。注意：OCR可能把0识别为O，需要纠正',
            '- id_card_address: 住址，从身份证"住址"后面提取',
            '',
            '【禁止行为】',
            '- 禁止从工作经历、项目经验、技能描述中提取信息',
            '- 禁止将"四年制""本科学习"等学制描述误认为专业名',
            '- 禁止猜测或编造任何字段值',
            '- work_experiences, project_experiences, technical_skills, certifications 必须为空数组[]',
            '- work_years, current_company, current_position 必须为null',
            '',
            '【简历文本】',
            truncated_text,
            '【简历文本结束】',
            '',
            '【输出JSON格式】',
            '{',
            '    "name": "姓名（2-4个中文字）",',
            '    "gender": "男/女",',
            '    "age": 年龄数字或null,',
            '    "birth_date": "YYYY-MM-DD或null",',
            '    "phone": "手机号或null",',
            '    "email": "邮箱或null",',
            '    "education": "博士/硕士/本科/大专/专科/高中/中专或null",',
            '    "major": "专业全称或null",',
            '    "school": "学校全称或null",',
            '    "graduation_date": "YYYY-MM或null",',
            '    "work_years": null,',
            '    "current_company": null,',
            '    "current_position": null,',
            '    "id_card_number": "18位身份证号或null",',
            '    "id_card_address": "身份证地址或null",',
            '    "degree": "学位或null",',
            '    "school_system": "学制或null",',
            '    "study_form": "学习形式或null",',
            '    "technical_skills": [],',
            '    "certifications": [],',
            '    "work_experiences": [],',
            '    "project_experiences": []',
            '}',
            '',
            '请直接输出纯JSON，不要有任何解释文字，不要加```json标记，只输出JSON对象。',
        ]

        return NL.join(prompt_lines)

    @classmethod
    def _extract_from_id_card(cls, text: str, info: dict[str, Any]) -> None:
        """
        从身份证OCR文本中优先提取信息
        只提取：姓名、性别、出生日期、身份证号、地址
        """
        id_card_keywords = ['中华人民共和国居民身份证', '居民身份证', '公民身份号码', '签发机关']
        has_id_card = any(kw in text for kw in id_card_keywords)
        has_id_number = bool(re.search(r'\b\d{17}[\dXx]\b', text))
        
        if not has_id_card and not has_id_number:
            return

        logger.info('检测到身份证内容，优先提取信息')

        # === 1. 姓名提取（支持"姓名 XXX"格式） ===
        if not info.get('name'):
            name_patterns = [
                r'姓\s*名\s*[：:\s]\s*([\u4e00-\u9fa5·]{2,10})',
                r'姓名\s*[：:\s]\s*([\u4e00-\u9fa5·]{2,10})',
                r'([\u4e00-\u9fa5·]{2,10})\s*\n\s*性\s*别',
            ]
            for pattern in name_patterns:
                match = re.search(pattern, text)
                if match:
                    name = match.group(1).strip()
                    if 2 <= len(name) <= 10:
                        info['name'] = name
                        logger.info(f'身份证提取姓名: {name}')
                        break

        # === 2. 性别提取（支持空格分隔） ===
        if not info.get('gender'):
            gender_patterns = [
                r'性\s*别\s*[：:\s]\s*(男|女)',
                r'性别\s*[：:\s]\s*(男|女)',
                r'(男|女)\s*\n\s*民\s*族',
            ]
            for pattern in gender_patterns:
                match = re.search(pattern, text)
                if match:
                    info['gender'] = match.group(1)
                    logger.info(f'身份证提取性别: {info["gender"]}')
                    break

        # === 3. 从身份证号提取出生日期和年龄 ===
        # 先尝试标准格式（无空格）
        id_number_match = re.search(r'(?<!\d)(\d{6})(\d{8})(\d{3})([\dXx])(?!\d)', text)
        if not id_number_match:
            # 尝试OCR分词格式：身份证号中间有空格或换行（如 "342626 19980404 4912" 或 "342626\n19980404\n4912"）
            # 先清理：把"公民身份号码"后面的文本中的空格和换行去掉再匹配
            id_card_prefix_match = re.search(r'(?:公民身份号\s*码|身份号\s*码|号码|证号)\s*[：:\s]*\n?\s*((?:[\dXx\s\n]{18,25}))', text)
            if id_card_prefix_match:
                raw_id = id_card_prefix_match.group(1)
                # 去除所有空格和换行
                cleaned_id = re.sub(r'[\s\n]', '', raw_id)
                # 验证是否为有效身份证号
                id_check = re.match(r'^(\d{6})(\d{8})(\d{3})([\dXx])$', cleaned_id)
                if id_check:
                    id_number_match = id_check
        
        if id_number_match:
            birth_str = id_number_match.group(2)
            year = birth_str[:4]
            month = birth_str[4:6]
            day = birth_str[6:8]
            
            if not info.get('birth_date'):
                info['birth_date'] = f"{year}-{month}-{day}"
                logger.info(f'身份证提取出生日期: {info["birth_date"]}')
            
            if not info.get('age'):
                try:
                    current_year = time.localtime().tm_year
                    info['age'] = current_year - int(year)
                except:
                    pass
            
            if not info.get('id_card_number'):
                id_num = id_number_match.group(0)
                # OCR纠错
                id_num = id_num.replace('O', '0').replace('o', '0')
                id_num = id_num.replace('I', '1').replace('l', '1')
                info['id_card_number'] = id_num
                logger.info(f'身份证提取身份证号: {id_num}')

        # === 4. 从文本提取出生日期（支持"出生 1994年 7 月22日"格式） ===
        if not info.get('birth_date'):
            birth_matches = re.findall(r'(?:出\s*生|出生)\s*[：:\s]\s*(\d{4})\s*年\s*(\d{1,2})\s*月(?:\s*(\d{1,2})\s*日)?', text)
            if birth_matches:
                year, month, day = birth_matches[0]
                info['birth_date'] = f"{year}-{month.zfill(2)}-{day.zfill(2) if day else '01'}"
                try:
                    current_year = time.localtime().tm_year
                    info['age'] = current_year - int(year)
                except:
                    pass
                logger.info(f'身份证提取出生日期: {info["birth_date"]}')

        # === 5. 身份证号提取（支持空格分隔、换行分隔、多种前缀） ===
        if not info.get('id_card_number'):
            # 先尝试宽松匹配：找18位身份证号格式（第7-14位是日期）
            # 然后再回溯检查前缀
            id_patterns = [
                # 标准格式：公民身份号码：342626199804044912
                r'公民身份号码\s*[：:\s]\s*(\d{17}[\dXx])',
                # 宽松格式：公民身份号 码 342626199804044912（OCR分词错误）
                r'公民身份号\s*码\s*[：:\s]*\s*(\d{17}[\dXx])',
                # 简称格式
                r'身份证号\s*[：:\s]\s*(\d{17}[\dXx])',
                r'证件号码\s*[：:\s]\s*(\d{17}[\dXx])',
                r'号码\s*[：:\s]\s*(\d{17}[\dXx])',
                # 最宽松：只要前面有"号码"相关词，后面跟18位数字
                r'(?:号码|证号|身份)\s*[：:\s]*\n?\s*(\d{17}[\dXx])',
                # OCR空格分词格式：号码 后面18位数字被空格/换行分隔
                r'(?:公民身份号\s*码|身份号\s*码|号码|证号)\s*[：:\s]*\n?\s*((?:[\dXx][\s\n]*){18})',
            ]
            for pattern in id_patterns:
                match = re.search(pattern, text)
                if match:
                    id_num = match.group(1)
                    # 去除OCR空格和换行
                    id_num = re.sub(r'[\s\n]', '', id_num)
                    # 纠正OCR常见错误
                    id_num = id_num.replace('O', '0').replace('o', '0')
                    id_num = id_num.replace('I', '1').replace('l', '1')
                    if len(id_num) == 18:
                        info['id_card_number'] = id_num
                        logger.info(f'身份证提取身份证号: {id_num}')
                    break
            
            # 兜底：如果没有前缀匹配成功，尝试在文本中查找孤立的18位身份证号
            if not info.get('id_card_number'):
                # 查找18位数字（可能带X），且前6位是有效的地区码开头，第7-14位是日期
                loose_match = re.search(r'(?<![\d])([1-9]\d{5}(?:19|20)\d{2}(?:0[1-9]|1[0-2])(?:0[1-9]|[12]\d|3[01])\d{3}[\dXx])(?![\d])', text)
                if loose_match:
                    id_num = loose_match.group(1)
                    id_num = id_num.replace('O', '0').replace('o', '0')
                    id_num = id_num.replace('I', '1').replace('l', '1')
                    # 验证日期部分是否合理
                    year_str = id_num[6:10]
                    month_str = id_num[10:12]
                    day_str = id_num[12:14]
                    try:
                        year = int(year_str)
                        month = int(month_str)
                        day = int(day_str)
                        if 1940 <= year <= 2010 and 1 <= month <= 12 and 1 <= day <= 31:
                            info['id_card_number'] = id_num
                            logger.info(f'身份证兜底提取身份证号: {id_num}')
                    except:
                        pass
            
            # 最终兜底：OCR将身份证号拆成多段（如"342626 1998 0404 4912"），整体清理后匹配
            if not info.get('id_card_number'):
                # 搜索包含数字和空格/换行的20-30字符片段，清理后验证
                candidate_matches = re.findall(r'((?:[\dXx][\s\n]*){18,25})', text)
                for candidate in candidate_matches:
                    cleaned = re.sub(r'[\s\n]', '', candidate)
                    cleaned = cleaned.replace('O', '0').replace('o', '0').replace('I', '1').replace('l', '1')
                    if len(cleaned) == 18:
                        # 验证日期部分
                        try:
                            year = int(cleaned[6:10])
                            month = int(cleaned[10:12])
                            day = int(cleaned[12:14])
                            if 1940 <= year <= 2010 and 1 <= month <= 12 and 1 <= day <= 31:
                                info['id_card_number'] = cleaned
                                logger.info(f'身份证最终兜底提取: {cleaned}')
                                break
                        except:
                            pass

        # === 6. 地址提取（支持空格分隔、跨行格式） ===
        if not info.get('id_card_address'):
            address_patterns = [
                # 单行格式
                r'住\s*址\s*[：:\s]\s*([\u4e00-\u9fa50-9省市区县镇村街道号路幢室组庄屯乡]{5,150})',
                r'住址\s*[：:\s]\s*([\u4e00-\u9fa50-9省市区县镇村街道号路幢室组庄屯乡]{5,150})',
                # 跨行格式：住址\nXXX省XXX市（OCR分词导致）
                r'住\s*址\s*$\s*([\u4e00-\u9fa50-9省市区县镇村街道号路幢室组庄屯乡]{5,150})',
                # OCR可能将地址拆成多行：住址\nXXX省XXX市\nXXX区XXX路
                r'住\s*址\s*$\s*([\u4e00-\u9fa50-9\n省市区县镇村街道号路幢室组庄屯乡]{5,200})',
            ]
            for pattern in address_patterns:
                match = re.search(pattern, text, re.MULTILINE if '$' in pattern else 0)
                if match:
                    addr = match.group(1).strip()
                    # 去除OCR空格和换行（地址可能跨行）
                    addr = re.sub(r'[\s\n]', '', addr)
                    # 验证地址至少包含"省"或"市"或"区"或"县"等行政区划关键字
                    if len(addr) >= 5 and any(kw in addr for kw in ['省', '市', '区', '县', '镇', '乡', '村', '街', '路']):
                        info['id_card_address'] = addr
                        logger.info(f'身份证提取地址: {addr[:30]}...')
                        break

    @classmethod
    def _extract_from_xuexin(cls, text: str, info: dict[str, Any]) -> None:
        """
        从学信网OCR文本中优先提取教育相关信息
        只提取：姓名、性别、出生日期、学历层次、专业、学校、毕业日期、学位、学制、学习形式
        """
        xuexin_keywords = ['学信网', '教育部学历证书电子注册备案表', '高等教育', '学籍', '学历证书查询', 'CHSI', '学历']
        has_xuexin = any(kw in text for kw in xuexin_keywords)
        # 同时检查是否有学历字段（纯OCR可能丢失标题）
        has_edu_field = any(
            kw in text for kw in ['学历层次', '学校名称', '专业名称', '院校名称', '入学日期', '毕业日期', '毕（结）业', '学制', '学习形式', '学位']
        )

        if not has_xuexin and not has_edu_field:
            return

        logger.info('检测到学信网内容，优先提取教育信息')

        # === 1. 姓名提取（学信网格式 "姓名 XXX" 或 "姓名\nXXX"） ===
        if not info.get('name'):
            name_patterns = [
                r'姓\s*名\s*[：:\s]\s*([\u4e00-\u9fa5·]{2,10})',
                r'姓名\s*[：:\s]\s*([\u4e00-\u9fa5·]{2,10})',
            ]
            for pattern in name_patterns:
                match = re.search(pattern, text)
                if match:
                    name = match.group(1).strip()
                    if 2 <= len(name) <= 10:
                        info['name'] = name
                        logger.info(f'学信网提取姓名: {name}')
                        break

        # === 2. 性别提取（学信网通常也有性别字段） ===
        if not info.get('gender'):
            gender_patterns = [
                r'性\s*别\s*[：:\s]\s*(男|女)',
                r'性别\s*[：:\s]\s*(男|女)',
            ]
            for pattern in gender_patterns:
                match = re.search(pattern, text)
                if match:
                    info['gender'] = match.group(1)
                    logger.info(f'学信网提取性别: {info["gender"]}')
                    break

        # === 3. 出生日期提取（学信网可能有出生信息） ===
        if not info.get('birth_date'):
            birth_patterns = [
                r'(?:出生日期|出生年月|出生)\s*[：:\s]\s*(\d{4})\s*[-年.\s]\s*(\d{1,2})\s*[-月.\s]\s*(\d{1,2})',
                r'(?:出生日期|出生年月|出生)\s*[：:\s]\s*(\d{4})\s*[-年.\s]\s*(\d{1,2})',
            ]
            for pattern in birth_patterns:
                match = re.search(pattern, text)
                if match:
                    groups = match.groups()
                    year = groups[0]
                    month = groups[1].zfill(2)
                    day = groups[2].zfill(2) if len(groups) > 2 else '01'
                    info['birth_date'] = f"{year}-{month}-{day}"
                    try:
                        current_year = time.localtime().tm_year
                        info['age'] = current_year - int(year)
                    except:
                        pass
                    logger.info(f'学信网提取出生日期: {info["birth_date"]}')
                    break

        # === 4. 学历层次提取 ===
        if not info.get('education'):
            edu_patterns = [
                r'学历层次\s*[：:\s]\s*([\u4e00-\u9fa5a-zA-Z]{2,30})',
                r'层次\s*[：:\s]\s*([\u4e00-\u9fa5a-zA-Z]{2,30})',
            ]
            for pattern in edu_patterns:
                match = re.search(pattern, text)
                if match:
                    edu = match.group(1).strip()
                    # 标准化：提取核心学历词
                    if '博士' in edu:
                        edu = '博士'
                    elif '硕士' in edu or '研究生' in edu:
                        edu = '硕士'
                    elif '本科' in edu or '学士' in edu:
                        edu = '本科'
                    elif '专科' in edu or '大专' in edu:
                        edu = '大专'
                    info['education'] = edu
                    logger.info(f'学信网提取学历层次: {edu}')
                    break

        # === 5. 专业名称提取（学信网截图格式："专业：信息与计算科学" 或两行格式） ===
        if not info.get('major'):
            # 先尝试单行格式：专业：信息与计算科学
            # 注意：字符类中用 [ \t] 替代 \s，避免跨行匹配到下一行的字段值
            major_patterns = [
                r'专业名称\s*[：:\s]\s*([\u4e00-\u9fa5a-zA-Z0-9·（）()\- \t]{2,60})',
                r'专业\s*[：:\s]\s*([\u4e00-\u9fa5a-zA-Z0-9·（）()\- \t]{2,60})',
            ]
            for pattern in major_patterns:
                match = re.search(pattern, text)
                if match:
                    major = match.group(1).strip()
                    # 去除OCR空格后进行噪声清洗
                    major_no_space = re.sub(r'\s+', '', major)
                    # 清洗常见噪声词（取噪声词之前的部分）
                    noise_words = ['学制', '学习形式', '层次', '毕业', '院校', '学校', '姓名',
                                   '年制本科', '年制专科', '入学日期', '毕业日期', '证书编号',
                                   '学历层次', '学历', '性别', '出生日期']
                    for noise in noise_words:
                        if noise in major_no_space:
                            major_no_space = major_no_space.split(noise)[0].strip()
                            break
                    # 排除明显不是专业名的值
                    invalid_values = ['四年制', '本科', '专科', '全日制', '普通', '学习', '信息',
                                      '四年制本科学习', '年制本科', '年制专科', '四年制本科', '',
                                      'null', 'None', '未知', '无']
                    invalid_patterns = [
                        r'^四年制', r'^\d+年制', r'本科学习', r'专科学习',
                        r'^本科$', r'^专科$', r'^硕士$', r'^博士$',
                        r'^年.*月', r'^\d{1,2}年.*月',  # "年九月至二" 等日期片段
                        r'^至$', r'^\d+至\d+',
                        r'学习形式', r'学制', r'普通全日制',
                    ]
                    is_invalid = major_no_space in invalid_values
                    if not is_invalid:
                        for inv_pat in invalid_patterns:
                            if re.search(inv_pat, major_no_space, re.IGNORECASE):
                                is_invalid = True
                                break
                    # 额外检查：专业名至少含2个中文字符
                    if not is_invalid:
                        chinese_chars = re.findall(r'[\u4e00-\u9fa5]', major_no_space)
                        if len(chinese_chars) < 2:
                            is_invalid = True
                    if not is_invalid and 2 <= len(major_no_space) <= 60:
                        info['major'] = major_no_space
                        logger.info(f'学信网提取专业名称(单行): {info["major"]}')
                        break
            
            # 如果单行没匹配到，尝试两行格式："专业\n信息与计算科学"
            if not info.get('major'):
                two_line_patterns = [
                    r'专业名称\s*$\s*([\u4e00-\u9fa5a-zA-Z·（）()\-]{2,60})',
                    r'专业\s*$\s*([\u4e00-\u9fa5a-zA-Z·（）()\-]{2,60})',
                ]
                for pattern in two_line_patterns:
                    match = re.search(pattern, text, re.MULTILINE)
                    if match:
                        major = match.group(1).strip()
                        # 验证：排除学制、学习形式等误匹配
                        invalid_vals = ['四年制', '本科', '专科', '全日制', '普通', '学习',
                                        'null', 'None', '未知', '无']
                        if major not in invalid_vals and len(major) >= 2:
                            # 进一步验证：专业名通常不含"年""制""形式"等词
                            if not any(bad in major for bad in ['学制', '学习形式', '层次', '年制', '年月']):
                                # 至少含2个中文字符
                                chinese_chars = re.findall(r'[\u4e00-\u9fa5]', major)
                                if len(chinese_chars) >= 2:
                                    info['major'] = major
                                    logger.info(f'学信网提取专业名称(两行): {major}')
                                    break

        # === 6. 学校名称提取 ===
        if not info.get('school'):
            school_patterns = [
                r'学校名称\s*[：:\s]\s*([\u4e00-\u9fa5a-zA-Z0-9（）()\-]{3,80})',
                r'毕业院校\s*[：:\s]\s*([\u4e00-\u9fa5a-zA-Z0-9（）()\-]{3,80})',
                r'院校名称\s*[：:\s]\s*([\u4e00-\u9fa5a-zA-Z0-9（）()\-]{3,80})',
            ]
            for pattern in school_patterns:
                match = re.search(pattern, text)
                if match:
                    school = match.group(1).strip()
                    if len(school) > 80:
                        school = school[:80]
                    info['school'] = school
                    logger.info(f'学信网提取学校名称: {school}')
                    break

        # === 7. 毕业日期提取（支持"毕（结）业日期"等学信网标准格式） ===
        if not info.get('graduation_date'):
            grad_patterns = [
                r'毕[（(]结[)）]业日期\s*[：:\s]\s*(\d{4})\s*[-年.\s]\s*(\d{1,2})',
                r'毕业日期\s*[：:\s]\s*(\d{4})\s*[-年.\s]\s*(\d{1,2})',
                r'毕业时间\s*[：:\s]\s*(\d{4})\s*[-年.\s]\s*(\d{1,2})',
                r'毕\s*业\s*[：:\s]\s*(\d{4})\s*[-年.\s]\s*(\d{1,2})',
            ]
            for pattern in grad_patterns:
                match = re.search(pattern, text)
                if match:
                    year, month = match.group(1), match.group(2).zfill(2)
                    info['graduation_date'] = f"{year}-{month}"
                    logger.info(f'学信网提取毕业日期: {info["graduation_date"]}')
                    break

        # === 8. 学位提取 ===
        if not info.get('degree'):
            degree_patterns = [
                r'学位\s*[：:\s]\s*([\u4e00-\u9fa5a-zA-Z]{2,30})',
                r'学位类别\s*[：:\s]\s*([\u4e00-\u9fa5a-zA-Z]{2,30})',
            ]
            for pattern in degree_patterns:
                match = re.search(pattern, text)
                if match:
                    degree = match.group(1).strip()
                    info['degree'] = degree
                    logger.info(f'学信网提取学位: {degree}')
                    break

        # === 9. 学制提取（支持单数字如"学制：4"，以及跨行OCR格式"学制\n4"） ===
        if not info.get('school_system'):
            # 单行格式
            sys_patterns = [
                r'学制\s*[：:\s]\s*([\u4e00-\u9fa50-9]{1,20})',
                r'学\s*制\s*[：:\s]\s*([\u4e00-\u9fa50-9]{1,20})',
            ]
            for pattern in sys_patterns:
                match = re.search(pattern, text)
                if match:
                    val = match.group(1).strip()
                    if val:
                        # 纯数字自动补"年"字
                        if val.isdigit():
                            info['school_system'] = val + '年'
                        else:
                            info['school_system'] = val
                        logger.info(f'学信网提取学制(单行): {info["school_system"]}')
                        break
            # 跨行格式：学制\n4（OCR分词导致字段名和值分在两行）
            if not info.get('school_system'):
                multi_line_match = re.search(r'学制\s*$\s*(\d+)', text, re.MULTILINE)
                if multi_line_match:
                    val = multi_line_match.group(1).strip()
                    info['school_system'] = val + '年'
                    logger.info(f'学信网提取学制(跨行): {info["school_system"]}')

        # === 10. 学习形式提取（支持单行和跨行OCR格式） ===
        if not info.get('study_form'):
            form_patterns = [
                # 单行格式：学习形式：普通全日制
                r'学习形式\s*[：:\s]\s*([\u4e00-\u9fa5]{2,20})',
                # 跨行格式：学习形式\n普通全日制（OCR分词导致）
                r'学习形式\s*$\s*([\u4e00-\u9fa5]{2,20})',
            ]
            for pattern in form_patterns:
                match = re.search(pattern, text, re.MULTILINE if '$' in pattern else 0)
                if match:
                    val = match.group(1).strip()
                    # 排除误匹配：学制数字、"本科"等不是学习形式
                    if val.isdigit() or val in ['本科', '专科', '硕士', '博士']:
                        continue
                    # 排除"四""年"等残缺值
                    if len(val) >= 2 and '制' not in val and '年' not in val:
                        info['study_form'] = val
                        logger.info(f'学信网提取学习形式: {info["study_form"]}')
                        break

    @classmethod
    def _extract_by_rules(cls, text: str, info: dict[str, Any]) -> None:
        """
        规则提取（确定性高的字段）
        增强中文简历多种格式支持，包括 | 分隔格式
        仅补充info中尚未有值的字段
        """
        lines = text.split('\n')

        # === 0. 预处理：识别 | 分隔格式的行 ===
        # 格式如: 赖斌 | 男 | 1998 | 浙江建德 | 本科 |
        pipe_info = {}
        for line in lines[:15]:  # 只看简历前15行
            line_stripped = line.strip()
            if not line_stripped:
                continue
            parts = [p.strip() for p in line_stripped.split('|') if p.strip()]
            if len(parts) >= 3:
                # 尝试识别 | 分隔行中的字段
                for i, part in enumerate(parts):
                    # 性别
                    if part in ['男', '女'] and not pipe_info.get('gender'):
                        pipe_info['gender'] = part
                    # 年龄/出生年份
                    elif re.match(r'^\d{4}$', part) and not pipe_info.get('birth_year'):
                        birth_year = int(part)
                        if 1940 <= birth_year <= 2010:
                            pipe_info['birth_year'] = birth_year
                    elif re.match(r'^(\d{1,2})岁?$', part) and not pipe_info.get('age'):
                        age_val = int(re.match(r'^(\d{1,2})', part).group(1))
                        if 18 <= age_val <= 65:
                            pipe_info['age'] = age_val
                    # 学历
                    elif any(edu in part for edu in cls.EDUCATION_KEYWORDS) and not pipe_info.get('education'):
                        for edu in cls.EDUCATION_KEYWORDS:
                            if edu in part:
                                pipe_info['education'] = edu
                                break
                    # 手机号
                    elif re.match(r'^1[3-9]\d{9}$', part) and not pipe_info.get('phone'):
                        pipe_info['phone'] = part
                    # 邮箱
                    elif re.match(r'^[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}$', part) and not pipe_info.get('email'):
                        pipe_info['email'] = part
                    # 籍贯/地址（通常在出生年份后面）
                    elif re.match(r'^[\u4e00-\u9fa5]{2,6}(?:省|市|区|县)?$', part) and not pipe_info.get('location'):
                        # 排除学历等关键词
                        if not any(edu in part for edu in cls.EDUCATION_KEYWORDS) and part not in ['男', '女']:
                            pipe_info['location'] = part
                # 如果 | 分隔行第一部分是姓名（2-4个中文字）
                if parts and not pipe_info.get('name'):
                    first = parts[0].strip()
                    if 2 <= len(first) <= 4 and re.match(r'^[\u4e00-\u9fa5]+$', first):
                        # 确保不是已知关键词
                        if first not in ['男', '女'] and not any(edu in first for edu in cls.EDUCATION_KEYWORDS):
                            pipe_info['name'] = first

        # 从 pipe_info 补充到 info（仅补充空值）
        for key in ['name', 'gender', 'age', 'education', 'phone', 'email']:
            if not info.get(key) and pipe_info.get(key):
                info[key] = pipe_info[key]
        if not info.get('birth_date') and pipe_info.get('birth_year'):
            birth_year = pipe_info['birth_year']
            info['birth_date'] = f"{birth_year}-01-01"
            if not info.get('age'):
                current_year = time.localtime().tm_year
                info['age'] = current_year - birth_year

        # === 1. 姓名提取 ===
        if not info.get('name'):
            name_patterns = [
                r'姓\s*名\s*[：:]\s*([^\s\n，,|]{2,4})',
                r'姓名\s*[：:]\s*([^\s\n，,|]{2,4})',
                r'名\s*字\s*[：:]\s*([^\s\n，,|]{2,4})',
            ]
            for pattern in name_patterns:
                match = re.search(pattern, text, re.MULTILINE)
                if match:
                    info['name'] = match.group(1).strip()
                    break
            else:
                # 尝试匹配文档开头2-4个字符作为姓名
                for line in lines:
                    line = line.strip()
                    if not line:
                        continue
                    # 跳过 | 分隔行（已在上面处理）
                    if '|' in line:
                        continue
                    if 2 <= len(line) <= 4 and not re.search(r'[a-zA-Z0-9]', line) and not any(kw in line for kw in ['简历', '求职', '应聘', '个人信息', '基本信息', '姓名']):
                        info['name'] = line
                        break

        # === 2. 年龄提取 ===
        if not info.get('age'):
            age_patterns = [
                r'年\s*龄\s*[：:]\s*(\d{1,2})\s*岁?',
                r'年龄\s*[：:]\s*(\d{1,2})\s*岁?',
                r'岁数\s*[：:]\s*(\d{1,2})\s*岁?',
                r'(\d{1,2})\s*岁[龄]?',
            ]
            for pattern in age_patterns:
                match = re.search(pattern, text)
                if match:
                    try:
                        age = int(match.group(1))
                        if 10 <= age <= 100:
                            info['age'] = age
                            break
                    except:
                        pass

        # 通过出生日期计算年龄
        if not info.get('birth_date'):
            birth_patterns = [
                # 学信网格式: "出生日期 1998-07-15" 或 "出生日期：1998年7月15日"
                r'(?:出生日期?)[：:\s]*\n?\s*(\d{4})\s*[年/\-]\s*(\d{1,2})\s*[月/\-]?\s*(\d{1,2})?\s*日?',
                r'(?:出生日期?)[：:]\s*(\d{4})\s*[年/\-]\s*(\d{1,2})\s*[月/\-]?\s*(\d{1,2})?日?',
                r'(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})?\s*日?\s*(?:出生)',
                # "生日\n1990/11/26" 或 "生日：1990/11/26" 格式
                r'(?:生日|出生日期?)[：:\s]*\n?\s*(\d{4})[/\-.年]\s*(\d{1,2})[/\-.月]?\s*(\d{1,2})?[/\-.日]?',
            ]
            for pattern in birth_patterns:
                match = re.search(pattern, text)
                if match:
                    year = match.group(1)
                    month = match.group(2) if match.lastindex >= 2 else '01'
                    day = match.group(3) if match.lastindex >= 3 else '01'
                    info['birth_date'] = f"{year}-{month.zfill(2)}-{day.zfill(2)}"
                    try:
                        birth_year = int(year)
                        current_year = time.localtime().tm_year
                        if not info.get('age'):
                            info['age'] = current_year - birth_year
                    except:
                        pass
                    break

        # === 3. 学历提取 ===
        if not info.get('education'):
            edu_patterns = [
                r'学\s*历\s*[：:]\s*(博士[研究生]?|[硕博]士研究生?|[硕博]士|[本专]科[生]?|[本专]科学士?|[本专]学?学位?|[专高]科[生]?|[专高]科学士?|[专高]学?学位?|[中高职高]专[生]?|[中高职高]科学士?|[中高职高]学?学位?|[职高]专[生]?|[职高]科学士?|[职高]学?学位?|[职高]高[中]?[生]?|[职高]高[中]?[生]?|[职高]高[中]?[生]?)',
                r'学历\s*[：:]\s*(博士[研究生]?|[硕博]士研究生?|[硕博]士|[本专]科[生]?|[本专]科学士?|[本专]学?学位?|[专高]科[生]?|[专高]科学士?|[专高]学?学位?|[中高职高]专[生]?|[中高职高]科学士?|[中高职高]学?学位?|[职高]专[生]?|[职高]科学士?|[职高]学?学位?|[职高]高[中]?[生]?|[职高]高[中]?[生]?|[职高]高[中]?[生]?)',
                r'学位\s*[：:]\s*(博士[研究生]?|[硕博]士研究生?|[硕博]士|[本专]科[生]?|[本专]科学士?|[本专]学?学位?|[专高]科[生]?|[专高]科学士?|[专高]学?学位?|[中高职高]专[生]?|[中高职高]科学士?|[中高职高]学?学位?|[职高]专[生]?|[职高]科学士?|[职高]学?学位?|[职高]高[中]?[生]?|[职高]高[中]?[生]?|[职高]高[中]?[生]?)',
            ]
            for pattern in edu_patterns:
                match = re.search(pattern, text)
                if match:
                    info['education'] = match.group(1).strip()
                    break
            else:
                for edu in cls.EDUCATION_KEYWORDS:
                    if edu in text:
                        info['education'] = edu
                        break

        # === 4. 教育背景区块智能提取（school + major + graduation_date + education 一起提取） ===
        # 简历中教育背景常见格式：
        #   格式A: "2018-09 ~ 2022-06  郑州工业应用技术学院  计算机科学与技术（本科）"
        #   格式B: "毕业院校：内蒙古科技大学  专业：计算机科学与技术"
        #   格式C: "黄河交通学院\n计算机科学与技术（本科）"
        #   格式D: 智国法简历的键值对格式 "毕业院校  内蒙古科技大学  专业  计算机科学与技术"
        need_school = not info.get('school')
        need_major = not info.get('major')
        need_grad = not info.get('graduation_date')
        need_edu = not info.get('education')

        if need_school or need_major or need_grad or need_edu:
            # 定位教育背景区块（注意：有些PDF的标题在内容之后，需双向搜索）
            edu_section = ''
            edu_markers = ['教育背景', '教育经历', '教育信息', '学历信息', 'Education']
            for marker in edu_markers:
                idx = text.find(marker)
                if idx >= 0:
                    # 取标记前后各1000字符作为教育区块（因为标题可能在内容前或后）
                    start = max(0, idx - 1000)
                    end = min(len(text), idx + 2000)
                    edu_section = text[start:end]
                    break

            # 如果没有找到明确的教育区块标题，尝试在前30%文本中查找
            if not edu_section:
                search_text = text[:len(text) // 3]
            else:
                search_text = edu_section

            edu_parsed_by_variant2 = False  # 标记变体2是否已直接填充info

            # 模式A: 时间范围 + 学校名 + 专业（含学历括号）
            # 匹配: 2018-09 ~ 2022-06  郑州工业应用技术学院  计算机科学与技术（本科）
            # 注意：PDF提取可能将日期跨行拆分，如 "2023-0\n...\n7"，需要宽松匹配
            edu_pattern_a = re.search(
                r'(\d{4})[/\-.年]\s*(\d{1,2})[/\-.月]?\s*[~\-至—–]+\s*(\d{4})[/\-.年]\s*(\d{1,2})?[/\-.月]?\s*\n?(?:.*\n){0,5}?\s*([\u4e00-\u9fa5a-zA-Z（）\(\)]+?(?:大学|学院|学校|研究院|专科学校|职业技术学院))\s*\n?\s*([\u4e00-\u9fa5a-zA-Z（）\(\)]+?)\s*(?:\n|$)',
                search_text
            )
            if not edu_pattern_a:
                # 变体：同一行，空格分隔
                edu_pattern_a = re.search(
                    r'(\d{4})[/\-.年]\s*(\d{1,2})[/\-.月]?\s*[~\-至—–]+\s*(\d{4})[/\-.年]\s*(\d{1,2})[/\-.月]?\s+([\u4e00-\u9fa5a-zA-Z（）\(\)]+?(?:大学|学院|学校|研究院|专科学校|职业技术学院))\s+([\u4e00-\u9fa5a-zA-Z·（）\(\)]{2,30}?)(?:\s|$)',
                    search_text
                )
            if not edu_pattern_a:
                # 变体2: 宽松全文匹配 - 只要找到 "时间范围" 后面跟 "学校名" 和 "专业" 即可
                # 适用于日期跨行、日期和学校之间有其他内容的情况
                date_match = re.search(
                    r'(\d{4})[/\-.年]\s*(\d{1,2})[/\-.月]?\s*[~\-至—–]+\s*(\d{4})[/\-.年]\s*(\d{1,2})?[/\-.月]?',
                    search_text
                )
                if date_match:
                    # 找到时间范围后，在后续文本中找学校和专业
                    remaining = search_text[date_match.end():][:500]
                    school_match = re.search(r'([\u4e00-\u9fa5a-zA-Z]+?(?:大学|学院|学校|研究院|专科学校|职业技术学院))', remaining)
                    if school_match:
                        after_school = remaining[school_match.end():][:200]
                        major_match = re.search(r'([\u4e00-\u9fa5a-zA-Z（）\(\)]{2,30}?)(?:\n|$|\s{2,})', after_school)
                        if major_match and date_match.lastindex >= 3:
                            grad_year = date_match.group(3) if date_match.lastindex >= 3 else None
                            grad_month_raw = date_match.group(4) if date_match.lastindex >= 4 else ''
                            # 如果月份不完整（如"0"，可能是PDF跨行截断），尝试从后续文本中补全
                            if not grad_month_raw or grad_month_raw in ('0', '1'):
                                # 在日期匹配结束后的200字符内查找孤立数字（7→07月）
                                nearby = search_text[date_match.end():date_match.end() + 300]
                                # 查找 "7\n" 或 "\n7\n" 这种独立出现的月份数字
                                month_fix = re.search(r'(?:^|\n)\s*(\d{1,2})\s*\n', nearby)
                                if month_fix:
                                    try:
                                        m = int(month_fix.group(1))
                                        if 1 <= m <= 12:
                                            grad_month_raw = str(m)
                                    except:
                                        pass
                            grad_month = grad_month_raw if grad_month_raw and int(grad_month_raw) > 0 else '07'  # 默认7月（中国毕业季）
                            school_name = school_match.group(1).strip()
                            major_text = major_match.group(1).strip()

                            if need_grad:
                                info['graduation_date'] = f"{grad_year}-{grad_month.zfill(2)}"
                            if need_school and school_name:
                                info['school'] = school_name
                            if need_major and major_text:
                                major_clean = re.sub(r'[（(]\s*(?:博士|硕士|本科|大专|专科|高中|中专|职高|研究生)\s*[）)]', '', major_text).strip()
                                if major_clean and len(major_clean) >= 2 and major_clean not in ['技能', '特长', '专业', '能力']:
                                    info['major'] = major_clean
                            if need_edu:
                                edu_in_bracket = re.search(r'[（(]\s*(博士|硕士|本科|大专|专科|高中|中专|职高|研究生)\s*[）)]', major_text)
                                if edu_in_bracket:
                                    info['education'] = edu_in_bracket.group(1)
                            edu_parsed_by_variant2 = True

            if edu_pattern_a and not edu_parsed_by_variant2:
                # 安全访问正则分组，避免"no such group"错误
                grad_year = edu_pattern_a.group(3) if edu_pattern_a.lastindex >= 3 else None
                grad_month = edu_pattern_a.group(4) if edu_pattern_a.lastindex >= 4 else None
                school_name = edu_pattern_a.group(5).strip() if edu_pattern_a.lastindex >= 5 else ''
                major_text = edu_pattern_a.group(6).strip() if edu_pattern_a.lastindex >= 6 else ''

                if need_grad and grad_year and grad_month:
                    info['graduation_date'] = f"{grad_year}-{grad_month.zfill(2)}"
                if need_school and school_name:
                    info['school'] = school_name
                if need_major and major_text:
                    # 专业文本可能包含（本科）等学历标注，需要分离
                    major_clean = re.sub(r'[（(]\s*(?:博士|硕士|本科|大专|专科|高中|中专|职高|研究生)\s*[）)]', '', major_text).strip()
                    if major_clean and len(major_clean) >= 2 and major_clean not in ['技能', '特长', '专业', '能力']:
                        info['major'] = major_clean
                if need_edu:
                    # 从括号中提取学历
                    edu_in_bracket = re.search(r'[（(]\s*(博士|硕士|本科|大专|专科|高中|中专|职高|研究生)\s*[）)]', major_text)
                    if edu_in_bracket:
                        info['education'] = edu_in_bracket.group(1)

            # 模式B/D: 键值对格式 "毕业院校：XXX  专业：YYY" 或 "毕业院校  XXX  专业  YYY"
            if not edu_pattern_a and not edu_parsed_by_variant2:
                # 毕业院校
                if need_school:
                    school_kv = re.search(
                        r'毕\s*业\s*院\s*校\s*[：:]\s*([\u4e00-\u9fa5a-zA-Z]+?(?:大学|学院|学校|研究院|专科学校|职业技术学院))',
                        search_text
                    )
                    if not school_kv:
                        # 无冒号的键值对格式 "毕业院校  内蒙古科技大学"
                        school_kv = re.search(
                            r'毕\s*业\s*院\s*校\s*[\s：:]*\n?\s*([\u4e00-\u9fa5a-zA-Z]+?(?:大学|学院|学校|研究院|专科学校|职业技术学院))',
                            search_text
                        )
                    if not school_kv:
                        # 宽松匹配：毕业院校后面跟的任意中文名称（3-20个字符），不以常见非学校词开头
                        school_kv = re.search(
                            r'毕\s*业\s*院\s*校\s*[：:\s]*\n?\s*([\u4e00-\u9fa5a-zA-Z]{3,20}?)(?:\s|$|\n|专|学|历|籍|现)',
                            search_text
                        )
                        if school_kv:
                            val = school_kv.group(1).strip()
                            # 排除明显不是学校名的值
                            if val in ('男', '女', '本科', '硕士', '博士', '大专', '高中', '专科'):
                                school_kv = None
                    if school_kv:
                        info['school'] = school_kv.group(1).strip()

                # 专业
                if need_major:
                    major_kv = re.search(
                        r'专\s*业\s*[：:]?\s*([\u4e00-\u9fa5a-zA-Z]{2,20}?)(?:\s|$|\n|籍|现|英|总体)',
                        search_text
                    )
                    if major_kv:
                        major_val = major_kv.group(1).strip()
                        if len(major_val) >= 2 and not re.search(r'^(姓名|[本专]科|[硕博]士|[专高]科)', major_val):
                            info['major'] = major_val

            # 模式C: 只有学校名（大学/学院关键词），后面跟专业
            if not edu_pattern_a and need_school:
                # 在教育区块中查找包含"大学""学院"等关键词的行
                for line in search_text.split('\n'):
                    line = line.strip()
                    if not line:
                        continue
                    # 匹配纯学校名行
                    school_match = re.match(r'^([\u4e00-\u9fa5a-zA-Z]+?(?:大学|学院|学校|研究院|专科学校|职业技术学院))$', line)
                    if school_match:
                        info['school'] = school_match.group(1)
                        break

            # 模式E: 智国法简历格式 "毕业院校\n内蒙古科技大学\n专业\n计算机科学与技术"
            if need_school and not info.get('school'):
                for i, line in enumerate(lines):
                    line_s = line.strip()
                    if re.match(r'^毕\s*业\s*院\s*校\s*$', line_s):
                        # 下一行是学校名
                        if i + 1 < len(lines):
                            next_line = lines[i + 1].strip()
                            school_m = re.match(r'^([\u4e00-\u9fa5a-zA-Z]+)$', next_line)
                            if school_m and len(next_line) >= 4:
                                info['school'] = next_line
                    if re.match(r'^专\s*业\s*$', line_s):
                        if i + 1 < len(lines):
                            next_line = lines[i + 1].strip()
                            if len(next_line) >= 2 and not re.search(r'^(姓名|[本专]科|[硕博]士)', next_line):
                                if need_major and not info.get('major'):
                                    info['major'] = next_line

            # 旧格式的兜底匹配（"专业：XXX" "毕业院校：XXX"）
            if need_major and not info.get('major'):
                major_patterns = [
                    r'专\s*业\s*[：:]\s*([^：:\n，,\s]{2,20})',
                    r'所学专业\s*[：:]\s*([^：:\n，,\s]{2,20})',
                ]
                for pattern in major_patterns:
                    match = re.search(pattern, text)
                    if match:
                        major = match.group(1).strip()
                        if len(major) >= 2 and not re.search(r'^(姓名|[本专]科|[硕博]士|[专高]科)', major):
                            info['major'] = major
                            break

            if need_school and not info.get('school'):
                school_patterns = [
                    r'毕\s*业\s*院\s*校\s*[：:]\s*([\u4e00-\u9fa5a-zA-Z]+?(?:大学|学院|学校|研究院|专科学校|职业技术学院))',
                    r'毕业院校\s*[：:]\s*([\u4e00-\u9fa5a-zA-Z]+?(?:大学|学院|学校|研究院|专科学校|职业技术学院))',
                    r'院\s*校\s*[：:]\s*([\u4e00-\u9fa5a-zA-Z]+?(?:大学|学院|学校|研究院|专科学校|职业技术学院))',
                    r'学\s*校\s*[：:]\s*([\u4e00-\u9fa5a-zA-Z]+?(?:大学|学院|学校|研究院|专科学校|职业技术学院))',
                ]
                for pattern in school_patterns:
                    match = re.search(pattern, text)
                    if match:
                        info['school'] = match.group(1).strip()
                        break

            # 毕业时间兜底
            if need_grad and not info.get('graduation_date'):
                grad_patterns = [
                    r'(?:毕业时间)[：:]\s*(\d{4})[/\-年](\d{1,2})[/\-月]?',
                    r'(\d{4})\s*年\s*毕业',
                ]
                for pattern in grad_patterns:
                    match = re.search(pattern, text)
                    if match:
                        year = match.group(1)
                        month = match.group(2) if match.lastindex >= 2 else '07'  # 默认7月
                        info['graduation_date'] = f"{year}-{month.zfill(2)}"
                        break

        # === 5.5 学信网/毕业证 OCR 专用提取 ===
        # 学信网截图 OCR 文本中通常包含: 院校名称、专业名称、学历层次、学制、学习形式、出生日期等
        # 毕业证 OCR 文本中也包含类似信息
        xuexin_keywords = ['学信网', '学历证书电子注册备案表', '教育部学籍验证', '学籍学历信息管理', '在线验证报告']
        is_xuexin_text = any(kw in text for kw in xuexin_keywords)

        if is_xuexin_text or not info.get('school') or not info.get('major') or not info.get('education'):
            # 学信网格式: "院校名称  XXX大学" / "院校名称：XXX大学"
            if not info.get('school'):
                school_xuexin = re.search(r'院\s*校\s*名\s*称\s*[：:\s]*\n?\s*([\u4e00-\u9fa5a-zA-Z]+?(?:大学|学院|学校|研究院|专科学校|职业技术学院))', text)
                if school_xuexin:
                    info['school'] = school_xuexin.group(1).strip()
                    logger.info(f'学信网规则提取 school: {info["school"]}')

            # 学信网格式: "专业名称  计算机科学与技术" / "专业名称：计算机科学与技术"
            if not info.get('major'):
                major_xuexin = re.search(r'专\s*业\s*名\s*称\s*[：:\s]*\n?\s*([\u4e00-\u9fa5a-zA-Z（）\(\)]{2,30}?)(?:\s|$|\n)', text)
                if major_xuexin:
                    major_val = major_xuexin.group(1).strip()
                    if len(major_val) >= 2 and not re.search(r'^(姓名|[本专]科|[硕博]士)', major_val):
                        info['major'] = major_val
                        logger.info(f'学信网规则提取 major: {info["major"]}')

            # 学信网格式: "学历层次  本科" / "学历：本科"
            if not info.get('education'):
                edu_xuexin = re.search(r'学\s*历\s*[层次]*\s*[：:\s]*\n?\s*(博士|硕士|研究生|本科|大专|专科|高中|中专|职高)', text)
                if edu_xuexin:
                    info['education'] = edu_xuexin.group(1).strip()
                    logger.info(f'学信网规则提取 education: {info["education"]}')

            # 学信网格式: "学制  4" / "学制：4年"
            if not info.get('school_system'):
                sys_xuexin = re.search(r'学\s*制\s*[：:\s]*\n?\s*(\d)\s*年?', text)
                if sys_xuexin:
                    info['school_system'] = f"{sys_xuexin.group(1)}年"
                    logger.info(f'学信网规则提取 school_system: {info["school_system"]}')

            # 学信网格式: "学习形式  普通全日制" / "学习形式：全日制"
            if not info.get('study_form'):
                form_xuexin = re.search(r'学\s*习\s*形\s*式\s*[：:\s]*\n?\s*(普通全日制|全日制|非全日制|函授|自考|网络教育|成人教育|业余)', text)
                if form_xuexin:
                    info['study_form'] = form_xuexin.group(1).strip()
                    logger.info(f'学信网规则提取 study_form: {info["study_form"]}')

            # 学信网格式: "出生日期  1998-07-15"
            if not info.get('birth_date'):
                birth_xuexin = re.search(r'出\s*生\s*日\s*期\s*[：:\s]*\n?\s*(\d{4})[/\-.年]\s*(\d{1,2})[/\-.月]?\s*(\d{1,2})?[/\-.日]?', text)
                if birth_xuexin:
                    year = birth_xuexin.group(1)
                    month = birth_xuexin.group(2) if birth_xuexin.lastindex >= 2 else '01'
                    day = birth_xuexin.group(3) if birth_xuexin.lastindex >= 3 else '01'
                    info['birth_date'] = f"{year}-{month.zfill(2)}-{day.zfill(2)}"
                    if not info.get('age'):
                        try:
                            info['age'] = time.localtime().tm_year - int(year)
                        except:
                            pass
                    logger.info(f'学信网规则提取 birth_date: {info["birth_date"]}, age: {info.get("age")}')

            # 学位: "授予 XX 学位" 或 学信网 "学位  学士"
            if not info.get('degree'):
                degree_xuexin = re.search(r'学\s*位\s*[：:\s]*\n?\s*(博士|硕士|学士)', text)
                if not degree_xuexin:
                    degree_xuexin = re.search(r'授\s*予\s*\S*?\s*(博士|硕士|学士)\s*学\s*位', text)
                if degree_xuexin:
                    info['degree'] = degree_xuexin.group(1).strip()
                    logger.info(f'学信网规则提取 degree: {info["degree"]}')

        # === 6. 工作年限提取 ===
        if not info.get('work_years'):
            work_years_patterns = [
                r'工\s*作\s*年\s*限\s*[：:]\s*(\d+)\s*年?',
                r'工作年限\s*[：:]\s*(\d+)\s*年?',
                r'工\s*作\s*经\s*验\s*[：:]\s*(\d+)\s*年?',
                r'工作经验\s*[：:]\s*(\d+)\s*年?',
                r'从\s*业\s*(\d+)\s*年',
                r'(\d+)\s*年\s*(?:工作)?经验',
                r'经验\s*(\d+)\s*年',
            ]
            for pattern in work_years_patterns:
                match = re.search(pattern, text)
                if match:
                    try:
                        years = int(match.group(1))
                        if 0 <= years <= 60:
                            info['work_years'] = years
                            break
                    except:
                        pass

        # === 7. 当前职位/目标岗位提取（增强版：支持多种格式） ===
        if not info.get('current_position'):
            # 模式1: 标准键值对格式
            position_patterns = [
                r'求\s*职\s*岗\s*位\s*[：:]\s*([^：:\n，,]{2,30})',
                r'求职岗位\s*[：:]\s*([^：:\n，,]{2,30})',
                r'意\s*向\s*岗\s*位\s*[：:]\s*([^：:\n，,]{2,30})',
                r'意向岗位\s*[：:]\s*([^：:\n，,]{2,30})',
                r'应\s*聘\s*职\s*位\s*[：:]\s*([^：:\n，,]{2,30})',
                r'应聘职位\s*[：:]\s*([^：:\n，,]{2,30})',
                r'目\s*标\s*岗\s*位\s*[：:]\s*([^：:\n，,]{2,30})',
                r'目标岗位\s*[：:]\s*([^：:\n，,]{2,30})',
                r'期望岗位\s*[：:]\s*([^：:\n，,]{2,30})',
            ]
            for pattern in position_patterns:
                match = re.search(pattern, text)
                if match:
                    position = match.group(1).strip()
                    if len(position) >= 2 and not re.search(r'^(姓名|[本专]科)', position):
                        info['current_position'] = position
                        break

        # 模式2: 从工作经历区块提取最近一段的职位
        if not info.get('current_position'):
            # 找到"工作经历"区块
            work_section_markers = ['工作经历', '工作经验', '工作履历', '职业经历']
            work_section = ''
            for marker in work_section_markers:
                idx = text.find(marker)
                if idx >= 0:
                    work_section = text[idx:idx + 3000]
                    break

            if work_section:
                # 在"工作经历"区块中查找公司名称+职位
                # 格式A: "公司名称（时间段）职位名称"
                # 格式B: "公司名称\n职位名称\n时间段"
                # 格式C: "公司名称 职位名称（时间段）"
                work_lines = work_section.split('\n')

                # 跳过区块标题行，找到第一段工作经历
                found_first_entry = False
                for i, line in enumerate(work_lines):
                    line_s = line.strip()
                    if not line_s:
                        continue
                    # 跳过区块标题
                    if any(marker in line_s for marker in work_section_markers):
                        continue
                    # 跳过"公司介绍""工作职责""工作业绩"等描述性标题
                    if re.match(r'^(公司介绍|工作职责|工作业绩|工作内容|岗位职责|工作描述|公司技术栈|公司介绍：)', line_s):
                        continue

                    if not found_first_entry:
                        found_first_entry = True
                        # 当前行可能是"公司名称 时间段"或"公司名称"
                        company_line = line_s
                        # 检查下一行是否是职位
                        if i + 1 < len(work_lines):
                            next_line = work_lines[i + 1].strip()
                            # 如果下一行不是描述性标题且不是空行，可能是职位
                            if next_line and not re.match(r'^(公司介绍|工作职责|工作业绩|工作内容|岗位职责|工作描述|公司技术栈|项目介绍|时间：|技术框架：)', next_line):
                                # 进一步检查：下一行是否包含常见职位关键词或不含"公司""项目"
                                position_keywords = ['工程师', '开发', '运维', '测试', '架构', '经理', '主管', '总监', '负责人', '专家', '顾问', '全栈', '前端', '后端']
                                if any(kw in next_line for kw in position_keywords):
                                    # 确认不是公司名（不含公司关键词或长度合适）
                                    if not re.search(r'(公司|集团|科技|信息|网络|有限|股份)', next_line) or len(next_line) <= 15:
                                        info['current_position'] = next_line[:30]
                                        break
                                elif len(next_line) <= 15 and not re.search(r'(公司|集团|科技|信息|网络)', next_line):
                                    # 短行且不含公司关键词，可能是职位
                                    info['current_position'] = next_line[:30]
                                    break

                        # 如果同一条行同时包含公司名和职位（用空格/括号分隔）
                        # 如 "北京数字观星科技有限公司 java开发"
                        company_pos_match = re.search(
                            r'([^\n]{2,40}(?:公司|集团|科技|信息|网络|有限|股份))\s*[（(]\d{4}[）)]\s*([^\n]{2,20}?(?:工程师|开发|运维|测试|架构|经理|主管|总监|负责人|专家|顾问|全栈|前端|后端))',
                            company_line
                        )
                        if company_pos_match and company_pos_match.lastindex >= 2:
                            info['current_position'] = company_pos_match.group(2).strip()
                            break

                        # 如果当前行就是职位（不含公司关键词且含职位关键词）
                        if not re.search(r'(公司|集团|科技|信息|网络|有限|股份)', company_line):
                            position_keywords = ['工程师', '开发', '运维', '测试', '架构', '经理', '主管', '总监', '负责人', '专家', '顾问', '全栈', '前端', '后端']
                            if any(kw in company_line for kw in position_keywords):
                                info['current_position'] = company_line[:30]
                                break
                else:
                    # 如果没找到，尝试更宽松的匹配：在"工作经历"区块中查找第一个含职位关键词的行
                    for line in work_lines:
                        line_s = line.strip()
                        if not line_s or any(marker in line_s for marker in work_section_markers):
                            continue
                        if re.search(r'(公司|集团|科技|信息|网络|有限|股份)', line_s):
                            continue
                        position_keywords = ['工程师', '开发', '运维', '测试', '架构', '经理', '主管', '总监', '负责人', '专家', '顾问', '全栈', '前端', '后端', 'java', 'python', 'go']
                        if any(kw.lower() in line_s.lower() for kw in position_keywords) and len(line_s) <= 25:
                            info['current_position'] = line_s[:30]
                            break

        # === 7.1 当前公司提取（增强版：支持多种格式） ===
        if not info.get('current_company'):
            # 模式1: 标准键值对格式
            company_patterns = [
                r'当\s*前\s*公\s*司\s*[：:]\s*([^：:\n，,]{2,40})',
                r'当前公司\s*[：:]\s*([^：:\n，,]{2,40})',
                r'现\s*公\s*司\s*[：:]\s*([^：:\n，,]{2,40})',
                r'现公司\s*[：:]\s*([^：:\n，,]{2,40})',
                r'最\s*近\s*公\s*司\s*[：:]\s*([^：:\n，,]{2,40})',
                r'最近公司\s*[：:]\s*([^：:\n，,]{2,40})',
            ]
            for pattern in company_patterns:
                match = re.search(pattern, text)
                if match:
                    company = match.group(1).strip()
                    if len(company) >= 2:
                        info['current_company'] = company
                        break

        # 模式2: 从工作经历区块提取最近一家公司
        if not info.get('current_company'):
            work_section_markers = ['工作经历', '工作经验', '工作履历', '职业经历']
            work_section = ''
            for marker in work_section_markers:
                idx = text.find(marker)
                if idx >= 0:
                    work_section = text[idx:idx + 3000]
                    break

            if work_section:
                work_lines = work_section.split('\n')
                found_first_entry = False
                for i, line in enumerate(work_lines):
                    line_s = line.strip()
                    if not line_s:
                        continue
                    # 跳过区块标题
                    if any(marker in line_s for marker in work_section_markers):
                        continue
                    # 跳过描述性标题
                    if re.match(r'^(公司介绍|工作职责|工作业绩|工作内容|岗位职责|工作描述|公司技术栈|项目介绍)', line_s):
                        continue
                    # 跳过"公司介绍："开头的行
                    if line_s.startswith('公司介绍：') or line_s.startswith('公司介绍'):
                        continue

                    if not found_first_entry:
                        found_first_entry = True
                        # 当前行可能是公司名称行
                        # 格式检查：是否包含公司关键词
                        if re.search(r'(公司|集团|科技|信息|网络|有限|股份)', line_s):
                            # 去掉日期部分（如"2025 年9 月-至今"）
                            company_clean = re.sub(r'\d{4}\s*年\s*\d{1,2}\s*月?\s*[-~至—]\s*(?:至今|\d{4}\s*年?\s*\d{1,2}\s*月?)', '', line_s).strip()
                            # 去掉括号中的内容
                            company_clean = re.sub(r'[（(][^）)]+[）)]', '', company_clean).strip()
                            if company_clean and len(company_clean) >= 2:
                                info['current_company'] = company_clean[:50]
                                break
                        # 如果当前行不含公司关键词但包含"公司介绍"的前缀，跳过
                        # 否则如果看起来像公司名（如"杭州微柠科技有限公司"格式）
                        elif re.search(r'[\u4e00-\u9fa5]{2,}(?:公司|集团|科技|信息|网络)', line_s):
                            company_clean = re.sub(r'\d{4}\s*年\s*\d{1,2}\s*月?\s*[-~至—]\s*(?:至今|\d{4}\s*年?\s*\d{1,2}\s*月?)', '', line_s).strip()
                            company_clean = re.sub(r'[（(][^）)]+[）)]', '', company_clean).strip()
                            if company_clean and len(company_clean) >= 2:
                                info['current_company'] = company_clean[:50]
                                break

        # === 8. 性别提取 ===
        if not info.get('gender'):
            gender_patterns = [
                r'性\s*别\s*[：:]\s*([男女])',
                r'性别\s*[：:]\s*([男女])',
                r'性\s*[：:]\s*([男女])',
            ]
            for pattern in gender_patterns:
                match = re.search(pattern, text)
                if match:
                    info['gender'] = match.group(1)
                    break
            else:
                # 回退：检查文本开头
                header = text[:500]
                if '男' in header and '女' not in header:
                    info['gender'] = '男'
                elif '女' in header and '男' not in header:
                    info['gender'] = '女'

        # === 9. 手机号提取 ===
        if not info.get('phone'):
            phone_patterns = [
                r'手\s*机\s*[：:]\s*(1[3-9]\d{9})',
                r'手机\s*[：:]\s*(1[3-9]\d{9})',
                r'电\s*话\s*[：:]\s*(1[3-9]\d{9})',
                r'电话\s*[：:]\s*(1[3-9]\d{9})',
                r'联系方式\s*[：:]\s*(1[3-9]\d{9})',
                r'联\s*系\s*方\s*式\s*[：:]\s*(1[3-9]\d{9})',
                r'手\s*机\s*号\s*码\s*[：:]\s*(1[3-9]\d{9})',
                r'手机号码\s*[：:]\s*(1[3-9]\d{9})',
            ]
            for pattern in phone_patterns:
                match = re.search(pattern, text)
                if match:
                    info['phone'] = match.group(1)
                    break
            else:
                # 回退：全局匹配
                phone_match = re.search(r'1[3-9]\d{9}', text)
                if phone_match:
                    info['phone'] = phone_match.group()

        # === 10. 邮箱提取 ===
        if not info.get('email'):
            email_patterns = [
                r'电\s*子\s*邮\s*件\s*[：:]\s*([a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,})',
                r'电子邮件\s*[：:]\s*([a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,})',
                r'邮\s*箱\s*[：:]\s*([a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,})',
                r'邮箱\s*[：:]\s*([a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,})',
                r'E[-\s]?mail\s*[：:]\s*([a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,})',
            ]
            for pattern in email_patterns:
                match = re.search(pattern, text, re.IGNORECASE)
                if match:
                    info['email'] = match.group(1).strip()
                    break
            else:
                # 回退：全局匹配
                email_match = re.search(r'[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}', text)
                if email_match:
                    info['email'] = email_match.group()

        # === 11. 技能提取（多区块定位 + 全文兜底） ===
        # 阶段1: 从"专业技能/技能特长"区块提取
        skill_section_markers = [
            r'专业技能[特长能]?[：:]?\s*',
            r'技能特?长?[：:]?\s*',
            r'个人技能[：:]?\s*',
            r'技术技能[：:]?\s*',
            r'技能清单[：:]?\s*',
        ]
        skill_text = ''
        for marker in skill_section_markers:
            parts = re.split(marker, text, maxsplit=1)
            if len(parts) > 1:
                # 取到下一个大标题之前的文本（避免超出技能区块）
                skill_text = parts[1][:2500]
                # 在技能区块中查找技术关键词
                for skill in cls.TECH_SKILLS_KEYWORDS:
                    if skill in skill_text and skill not in info['technical_skills']:
                        info['technical_skills'].append(skill)
                # 也尝试从技能区块提取技能词（通过"熟悉""掌握""精通"等动词引导的描述）
                skill_lines = skill_text.split('\n')[:15]
                for line in skill_lines:
                    if any(kw in line for kw in ['技能', '技术', '熟悉', '掌握', '精通', '了解', '使用', '运用', '擅长']):
                        for skill in cls.TECH_SKILLS_KEYWORDS:
                            if skill in line and skill not in info['technical_skills']:
                                info['technical_skills'].append(skill)
                break

        # 阶段2: 从"个人优势/复合型优势"区块提取（很多简历把技术能力写在个人优势里）
        advantage_sections = re.split(r'(?:个人优势[：:]?\s*|核心优势[：:]?\s*|复合型[^\n]*优势[：:]?\s*|核心能力[：:]?\s*)', text, maxsplit=1)
        if len(advantage_sections) > 1:
            advantage_text = advantage_sections[1][:2000]
            for skill in cls.TECH_SKILLS_KEYWORDS:
                if skill in advantage_text and skill not in info['technical_skills']:
                    info['technical_skills'].append(skill)

        # 阶段3: 从项目经验的"技术栈/技术框架"行提取
        tech_stack_matches = re.findall(r'(?:技术栈|技术框架|技术架构)[：:]*\s*([^\n]+)', text)
        for tech_line in tech_stack_matches:
            for skill in cls.TECH_SKILLS_KEYWORDS:
                if skill in tech_line and skill not in info['technical_skills']:
                    info['technical_skills'].append(skill)
            # 也从技术栈行中提取用顿号、逗号分隔的技术词
            # 先去除括号中的说明文字（如 Spring AI Alibaba），避免拆分出子串
            tech_line_clean = re.sub(r'[（(][^）)]*[）)]', '', tech_line)
            tech_items = re.split(r'[、，,\s·]+', tech_line_clean)
            for item in tech_items:
                item = item.strip().rstrip('.')
                if len(item) >= 2 and item not in info['technical_skills']:
                    # 过滤掉明显不是技术名词的词
                    if not re.match(r'^[\d]+$', item) and not re.match(r'^[（(]', item):
                        # 过滤过长的短语（可能是句子而非技术名词）
                        if len(item) <= 25:
                            # 检查是否包含技术相关特征（英文字母、技术中文词等）
                            if re.search(r'[a-zA-Z]', item) or any(kw in item for kw in [
                                '服务', '框架', '数据库', '缓存', '队列', '搜索引擎', '容器', '网关',
                                '微服务', '分布式', '中间件', '负载', '注册', '配置', '调度', '事务',
                            ]):
                                info['technical_skills'].append(item)

        # 阶段4: 全文兜底——对还没有匹配到的技能关键词做全文搜索
        for skill in cls.TECH_SKILLS_KEYWORDS:
            if skill not in info['technical_skills'] and skill in text:
                info['technical_skills'].append(skill)

        # === 12. 工作经历提取（基于区块定位 + 格式识别） ===
        # 仅在LLM没有提取到工作经历时才用规则补充
        if not info.get('work_experiences'):
            work_sections = re.split(r'(?:工作经历[：:]?\s*|工作经验[：:]?\s*|工作履历[：:]?\s*|职业经历[：:]?\s*)', text, maxsplit=1)
            if len(work_sections) > 1:
                work_text = work_sections[1][:3000]
                # 模式1: 公司 | 职位 | 时间段
                company_matches = re.findall(
                    r'([^\n]{2,40}(?:公司[集团]?|[科网信]?技|[信息]?网络))\s*[\|,;/\\]\s*([^\n]{2,20})\s*[\|,;/\\]\s*(\d{4}[./-]\d{1,2}\s*[-~至]\s*(?:\d{4}[./-]\d{1,2}|至今))',
                    work_text
                )
                for company, position, period in company_matches:
                    info['work_experiences'].append({
                        'company': company.strip(),
                        'position': position.strip(),
                        'start_date': period.split('至')[0].strip() if '至' in period else period.split('-')[0].strip() if '-' in period else period.split('~')[0].strip(),
                        'end_date': period.split('至')[1].strip() if '至' in period else period.split('-')[1].strip() if '-' in period else period.split('~')[1].strip(),
                        'description': '',
                    })

                # 模式2: 公司：XXX 职位：XXX 时间：2020-2023
                work_entries = re.findall(
                    r'(?:公\s*司[：:]\s*)?([^\n，,]{2,40}(?:公司[集团]?|[科网信]?技|[信息]?网络))\s*(?:[，,]\s*|\\s+|\\n)\s*(?:职\s*位[：:]\s*|岗\s*位[：:]\s*)?([^\n，,]{2,20})\s*(?:[，,]\s*|\\s+|\\n)\s*(?:(?:时\s*间[：:]\s*)?(?:(\d{4})[./-](\d{1,2})\s*[-~至]\s*(?:至今)?)?)',
                    work_text
                )
                for company, position, start_year, start_month in work_entries:
                    if company.strip() not in [w['company'] for w in info['work_experiences']]:
                        info['work_experiences'].append({
                            'company': company.strip(),
                            'position': position.strip(),
                            'start_date': f"{start_year}-{start_month.zfill(2)}" if start_year else '',
                            'end_date': '至今' if '至今' in work_text else '',
                            'description': '',
                        })

                # 模式3: 日期范围 + 公司行（宽松匹配，适用于OCR结果格式不规律的情况）
                # 如 "2020.07 - 至今\n杭州维科软件工程责任有限公司\n开发负责人"
                if not info.get('work_experiences'):
                    date_company_pattern = re.finditer(
                        r'(\d{4})[/\-.年]\s*(\d{1,2})?[/\-.月]?\s*[-~至—–]+\s*(\d{4})?[/\-.年]?\s*(\d{1,2})?[/\-.月]?|至今',
                        work_text
                    )
                    for dm in date_company_pattern:
                        # 安全访问正则分组，避免"no such group"错误
                        if dm.lastindex < 1 or not dm.group(1):
                            continue
                        after = work_text[dm.end():][:300]
                        lines_after = [l.strip() for l in after.split('\n') if l.strip()][:5]
                        company_name = ''
                        position_name = ''
                        for line in lines_after:
                            # 查找包含公司关键词的行
                            if not company_name and re.search(r'(?:公司|集团|科技|信息|网络|有限|责任|股份)', line):
                                company_name = line.strip()[:50]
                                continue
                            # 查找职位行（2-15个中文字，不包含公司关键词）
                            if company_name and not position_name and 2 <= len(line) <= 15:
                                if not re.search(r'(?:公司|集团|项目|负责|开发)', line):
                                    position_name = line.strip()
                                    break
                        if company_name:
                            group2 = dm.group(2) if dm.lastindex >= 2 else None
                            group3 = dm.group(3) if dm.lastindex >= 3 else None
                            group4 = dm.group(4) if dm.lastindex >= 4 else None
                            start_str = f"{dm.group(1)}-{(group2 or '01').zfill(2)}"
                            end_str = '至今' if '至今' in work_text[dm.start():dm.end()+10] else (
                                f"{group3}-{(group4 or '01').zfill(2)}" if group3 else ''
                            )
                            info['work_experiences'].append({
                                'company': company_name,
                                'position': position_name,
                                'start_date': start_str,
                                'end_date': end_str,
                                'description': '',
                            })

        # 提取当前公司（通常是最近的工作经历中的公司，仅补充空值）
        if not info.get('current_company') and info.get('work_experiences'):
            info['current_company'] = info['work_experiences'][0].get('company', '')
        if not info.get('current_position') and info.get('work_experiences'):
            info['current_position'] = info['work_experiences'][0].get('position', '')

        # 提取毕业时间（如果没有找到）
        if not info.get('graduation_date'):
            grad_patterns = [
                r'(?:毕业时间)[：:]\s*(\d{4})[/\-年](\d{1,2})[/\-月]?',
                r'(\d{4})\s*年\s*毕业',
            ]
            for pattern in grad_patterns:
                match = re.search(pattern, text)
                if match:
                    year = match.group(1)
                    month = match.group(2) if match.lastindex >= 2 else '07'  # 默认7月
                    info['graduation_date'] = f"{year}-{month.zfill(2)}"
                    break

    @classmethod
    def _fill_birth_age_from_id(cls, info: dict[str, Any]) -> None:
        """
        最终补全：从身份证号反算出生日期和年龄
        确保只要有身份证号就能算出年龄
        """
        id_num = info.get('id_card_number')
        if not id_num or not isinstance(id_num, str):
            return
        
        # 清理身份证号（去除可能的空格）
        id_num = re.sub(r'[\s\n]', '', id_num)
        if len(id_num) != 18:
            return
        
        # 身份证号第7-14位是出生日期 YYYYMMDD
        try:
            year = int(id_num[6:10])
            month = int(id_num[10:12])
            day = int(id_num[12:14])
            
            if not (1940 <= year <= 2010 and 1 <= month <= 12 and 1 <= day <= 31):
                return
            
            # 补全出生日期
            if not info.get('birth_date'):
                info['birth_date'] = f"{year}-{month:02d}-{day:02d}"
                logger.info(f'身份证号反算出生日期: {info["birth_date"]}')
            
            # 补全年龄（精确计算：考虑月份和日期）
            # 条件：age为None、0、或与身份证号推算的年龄差距超过2年
            current_age = info.get('age')
            need_recalc = current_age is None or current_age == 0
            if not need_recalc and isinstance(current_age, (int, str)):
                try:
                    need_recalc = abs(int(current_age) - (time.localtime().tm_year - year)) > 2
                except (ValueError, TypeError):
                    need_recalc = True
            if need_recalc:
                now = time.localtime()
                age = now.tm_year - year
                # 如果还没过生日，减1
                if now.tm_mon < month or (now.tm_mon == month and now.tm_mday < day):
                    age -= 1
                if 16 <= age <= 70:
                    info['age'] = age
                    logger.info(f'身份证号反算年龄: {age}')
        except (ValueError, IndexError):
            pass

    @classmethod
    def _post_process_info(cls, info: dict[str, Any]) -> None:
        """
        后处理：清洗和纠错LLM/规则提取的结果
        解决OCR错误、格式不规范、字段值污染等问题
        """
        # === 1. 手机号OCR纠错 ===
        # OCR常见错误：D→0, O→0, l→1, I→1, S→5, B→8
        phone = info.get('phone')
        if phone and isinstance(phone, str):
            # 只保留数字，过滤掉OCR识别错误的字母
            phone_clean = phone
            ocr_replace_map = {'D': '0', 'O': '0', 'o': '0', 'l': '1', 'I': '1', 'S': '5', 's': '5', 'B': '8', 'b': '8'}
            for old, new in ocr_replace_map.items():
                phone_clean = phone_clean.replace(old, new)
            # 验证是否为有效手机号
            if re.match(r'^1[3-9]\d{9}$', phone_clean):
                info['phone'] = phone_clean
            elif re.match(r'^1[3-9]\d{9}', phone_clean):
                # 前缀匹配但长度不对，截取前11位
                info['phone'] = phone_clean[:11]

        # === 2. 公司名称清洗 ===
        company = info.get('current_company')
        if company and isinstance(company, str):
            # 移除常见前缀描述
            prefixes_to_remove = [
                r'^公司介绍[：:]*\s*',
                r'^关于[：:]*\s*',
                r'^公司名称[：:]*\s*',
                r'^就职于\s*',
                r'^所在公司[：:]*\s*',
                r'^任职于\s*',
            ]
            for prefix in prefixes_to_remove:
                company = re.sub(prefix, '', company).strip()
            info['current_company'] = company

        # === 3. 职位名称清洗（增强版） ===
        position = info.get('current_position')
        if position and isinstance(position, str):
            position = position.strip()
            
            # 规则1: 如果职位名包含以下描述性前缀，说明是项目描述而非职位，直接清空
            # 这些前缀表示动作或描述，不是职位名称
            invalid_prefixes = [
                '项目说明', '核心贡献', '公司介绍', '工作职责', '工作内容',
                '基于', '通过', '面向', '负责', '使用', '类似',
                '搭建', '设计', '开发', '实现', '优化', '提升',
                '完成', '参与', '主导', '负责', '担任',
                '负责开发', '负责设计', '负责搭建',
                '技术栈：', '技术栈:', '技术架构',
            ]
            
            # 检查是否以无效前缀开头
            is_invalid = False
            for prefix in invalid_prefixes:
                if position.startswith(prefix):
                    is_invalid = True
                    break
            
            # 规则2: 如果职位名包含以下关键词且长度超过15字符，可能是项目描述
            if not is_invalid:
                desc_keywords = ['基于', '通过', '面向', '负责', '使用', '实现', '完成', '开发', '设计', '搭建']
                if len(position) > 15 and any(kw in position for kw in desc_keywords):
                    # 进一步检查：是否包含常见职位关键词
                    position_keywords = ['工程师', '开发', '运维', '测试', '架构', '经理', '主管', '总监', '负责人', '专家', '顾问', '全栈', '前端', '后端']
                    if not any(kw in position for kw in position_keywords):
                        is_invalid = True
            
            # 规则3: 长度超过30字符且不含职位关键词，大概率是描述
            if not is_invalid and len(position) > 30:
                position_keywords = ['工程师', '开发', '运维', '测试', '架构', '经理', '主管', '总监', '负责人', '专家', '顾问', '全栈', '前端', '后端']
                if not any(kw in position for kw in position_keywords):
                    is_invalid = True
            
            if is_invalid:
                # 如果是无效职位名，清空并尝试从工作经历中重新提取
                info['current_position'] = None
                # 尝试从工作经历中提取
                if info.get('work_experiences'):
                    for work in info['work_experiences']:
                        if isinstance(work, dict) and work.get('position'):
                            work_pos = work['position']
                            if isinstance(work_pos, str) and len(work_pos) <= 20:
                                info['current_position'] = work_pos
                                break
            else:
                # 移除项目描述等无关信息（如果职位名太长，说明可能混入了项目描述）
                if len(position) > 20:
                    # 尝试截取第一个有效职位名
                    for sep in ['，', '、', '。', '\n', '；']:
                        if sep in position:
                            position = position.split(sep)[0].strip()
                            break
                info['current_position'] = position

        # === 4. 学历标准化 ===
        education = info.get('education')
        if education and isinstance(education, str):
            edu_map = {
                '博士研究生': '博士', '硕士研究生': '硕士', '研究生': '硕士',
                '本科毕业生': '本科', '本科生于': '本科', '大专毕业生': '大专',
                '专科毕业生': '大专', '大专于': '大专', '专科于': '大专',
                '高中毕业生': '高中', '中专毕业生': '中专', '职高毕业生': '职高',
            }
            for k, v in edu_map.items():
                if k in education:
                    info['education'] = v
                    break
            # 如果学历值太长（可能是误提取），只取核心词
            edu = info['education']
            if len(edu) > 4:
                for std in ['博士', '硕士', '本科', '大专', '高中', '中专', '职高']:
                    if std in edu:
                        info['education'] = std
                        break

        # === 4.5 专业字段清洗（防止"四年制本科学习""年九月至二"等OCR噪声被误提取为专业） ===
        major = info.get('major')
        if major and isinstance(major, str):
            major_cleaned = re.sub(r'\s+', '', major)  # 去除OCR空格
            # 检查是否为无效专业值
            invalid_major_patterns = [
                r'^四年制', r'^\d+年制', r'本科学习', r'专科学习',
                r'^本科$', r'^专科$', r'^硕士$', r'^博士$',
                r'学习形式', r'学制', r'普通全日制',
                r'^年.*月', r'^\d{1,2}年.*月',  # "年九月至二" 等日期片段
                r'^至$', r'^\d+至\d+',
                r'入学日期', r'毕业日期', r'证书编号',
                r'^学历层次$', r'^层次$', r'^学历$',
                r'^null$', r'^None$', r'^未知$',
            ]
            is_invalid_major = False
            for inv_pat in invalid_major_patterns:
                if re.search(inv_pat, major_cleaned, re.IGNORECASE):
                    is_invalid_major = True
                    break
            # 额外检查：如果专业值不含任何中文实词（全是数字、标点、字母组合），也视为无效
            if not is_invalid_major:
                chinese_chars = re.findall(r'[\u4e00-\u9fa5]', major_cleaned)
                if len(chinese_chars) < 2:
                    is_invalid_major = True
            if is_invalid_major:
                logger.warning(f'专业字段清洗: "{major}" 为无效值，清空')
                info['major'] = None
            else:
                # 清除尾部粘连的字段名（OCR跨行导致的"专业：计算机科学与技术\n学历类别" → "计算机科学与技术学历类别"）
                trailing_noise = ['学历类别', '学习形式', '学制', '层次', '学历',
                                  '毕业日期', '入学日期', '证书编号', '院校名称',
                                  '学校名称', '毕(结)业', '毕（结）业']
                for noise in trailing_noise:
                    if major_cleaned.endswith(noise) and len(major_cleaned) > len(noise):
                        major_cleaned = major_cleaned[:-len(noise)]
                        logger.info(f'专业字段去除尾部粘连: "{noise}" → "{major_cleaned}"')
                        break
                # 清除后再次校验中文字符数
                chinese_chars = re.findall(r'[\u4e00-\u9fa5]', major_cleaned)
                if len(chinese_chars) >= 2:
                    info['major'] = major_cleaned
                else:
                    logger.warning(f'专业字段清洗后中文字符不足2: "{major_cleaned}"，清空')
                    info['major'] = None
        grad_date = info.get('graduation_date')
        if grad_date and isinstance(grad_date, str):
            # 确保 YYYY-MM 格式
            grad_date = grad_date.strip()
            # 处理 YYYY年MM月 格式
            m = re.match(r'(\d{4})\s*年\s*(\d{1,2})\s*月?', grad_date)
            if m:
                grad_date = f"{m.group(1)}-{m.group(2).zfill(2)}"
            # 处理 YYYY.MM 格式
            m = re.match(r'(\d{4})\.(\d{1,2})', grad_date)
            if m:
                grad_date = f"{m.group(1)}-{m.group(2).zfill(2)}"
            # 处理纯数字格式 YYYYMM
            m = re.match(r'^(\d{4})(\d{2})$', grad_date)
            if m:
                grad_date = f"{m.group(1)}-{m.group(2)}"
            info['graduation_date'] = grad_date

        # === 6. 年龄合理性校验 ===
        age = info.get('age')
        if age is not None:
            try:
                age_int = int(age)
                if age_int < 16 or age_int > 70:
                    info['age'] = None
                elif isinstance(age, str):
                    info['age'] = age_int
            except (ValueError, TypeError):
                info['age'] = None

        # === 6.5 年龄与出生日期交叉校验 ===
        # 如果出生日期存在但年龄与出生年份不一致，用出生日期重算年龄
        birth = info.get('birth_date')
        if birth and isinstance(birth, str):
            year_match = re.search(r'(19\d{2}|20\d{2})', birth)
            if year_match:
                birth_year = int(year_match.group(1))
                now = time.localtime()
                correct_age = now.tm_year - birth_year
                # 如果当前月份 < 出生月份（从birth_date提取），减1
                month_match = re.search(r'\d{4}-(\d{1,2})', birth)
                if month_match:
                    birth_month = int(month_match.group(1))
                    if now.tm_mon < birth_month:
                        correct_age -= 1
                if 16 <= correct_age <= 70:
                    current_age = info.get('age')
                    if current_age is None or current_age == 0 or abs(int(current_age) - correct_age) > 2:
                        logger.info(f'年龄交叉校验: birth_date={birth}, 原age={current_age} → 修正为{correct_age}')
                        info['age'] = correct_age

        # === 7. 技能去重和清洗 ===
        skills = info.get('technical_skills', [])
        if isinstance(skills, list):
            seen = set()
            cleaned = []
            for s in skills:
                if not isinstance(s, str):
                    continue
                s = s.strip()
                if not s or len(s) > 50 or s.lower() in ('null', '无', '未知', '暂无'):
                    continue
                # 去除括号中的版本号说明（如 "Spring Boot(2.x)" → "Spring Boot"）
                s_clean = re.sub(r'[（(]\s*[\d.x]+\s*[）)]', '', s).strip()
                if s_clean.lower() not in seen:
                    seen.add(s_clean.lower())
                    cleaned.append(s_clean)
            info['technical_skills'] = cleaned

        # === 8. 工作经历清洗（公司名 + 职位名） ===
        invalid_position_prefixes = [
            '项目说明', '核心贡献', '公司介绍', '工作职责', '工作内容',
            '基于', '通过', '面向', '负责', '使用', '类似',
            '搭建', '设计', '开发', '实现', '优化', '提升',
            '完成', '参与', '主导', '担任',
            '技术栈：', '技术栈:', '技术架构',
        ]
        position_keywords = ['工程师', '开发', '运维', '测试', '架构', '经理', '主管', '总监', '负责人', '专家', '顾问', '全栈', '前端', '后端']
        
        for work in info.get('work_experiences', []):
            if not isinstance(work, dict):
                continue
            
            # 清洗公司名
            company = work.get('company', '')
            if isinstance(company, str):
                for prefix in [r'^公司介绍[：:]*\s*', r'^关于[：:]*\s*', r'^所在公司[：:]*\s*']:
                    company = re.sub(prefix, '', company).strip()
                work['company'] = company
            
            # 清洗职位名
            position = work.get('position', '')
            if isinstance(position, str):
                position = position.strip()
                
                # 检查是否以无效前缀开头
                is_invalid = False
                for prefix in invalid_position_prefixes:
                    if position.startswith(prefix):
                        is_invalid = True
                        break
                
                # 检查是否包含描述性关键词且不含职位关键词
                if not is_invalid:
                    desc_keywords = ['基于', '通过', '面向', '负责', '使用', '实现', '完成', '开发', '设计', '搭建']
                    if len(position) > 15 and any(kw in position for kw in desc_keywords):
                        if not any(kw in position for kw in position_keywords):
                            is_invalid = True
                
                # 长度检查
                if not is_invalid and len(position) > 30 and not any(kw in position for kw in position_keywords):
                    is_invalid = True
                
                if is_invalid:
                    work['position'] = ''

        # === 9. 姓名清洗 ===
        name = info.get('name')
        if name and isinstance(name, str):
            name = name.strip()
            # 去除可能的"简历""求职"等后缀
            name = re.sub(r'(简历|求职|应聘|个人|的).*$', '', name).strip()
            # 去除前缀（如"姓名："）
            name = re.sub(r'^姓\s*名\s*[：:\s]*', '', name).strip()
            # 排除字段名被误识别为姓名
            INVALID_NAMES = {'性别', '民族', '年龄', '学历', '出生', '住址', '姓名', '学制', '学位', '专业', '学校', '毕业', '身份',
                             '男', '女', '联系电话', '联系方式', '邮箱', '地址', '职称', '职务', '电话',
                             '出生日期', '身份证', '公民', '号码', '签发机关', '有效期', '编号',
                             '投标人', '盖章', '说明', '备注', '日期', '签名', '签章',
                             'null', 'None', 'none', '未知', '无', '暂无'}
            if name in INVALID_NAMES:
                logger.warning(f'姓名清洗: "{name}" 为字段名/无效值，清空')
                info['name'] = None
            elif len(name) < 2 or len(name) > 5:
                info['name'] = None
            else:
                info['name'] = name

        # === 10. 全局 "null"/"None" 字符串过滤 ===
        # LLM 经常返回字符串 "null" 或 "None"，这些不是有效值
        null_strings = {'null', 'None', 'none', 'NULL', '未知', '暂无', '无', '-'}
        null_keys = ['phone', 'email', 'major', 'school', 'degree', 'school_system',
                     'study_form', 'current_company', 'current_position', 'birth_date',
                     'graduation_date', 'id_card_number', 'id_card_address']
        for key in null_keys:
            val = info.get(key)
            if isinstance(val, str) and val.strip() in null_strings:
                logger.info(f'全局null过滤: {key}="{val}" → None')
                info[key] = None

        # === 11. 出生日期残缺值清洗 ===
        birth = info.get('birth_date')
        if birth and isinstance(birth, str):
            birth = birth.strip()
            # 匹配 YYYY-MM-DD 格式
            m = re.match(r'^(\d{4})-(\d{1,2})-(\d{1,2})$', birth)
            if not m:
                # 尝试 YYYY-MM 格式
                m2 = re.match(r'^(\d{4})-(\d{1,2})$', birth)
                if not m2:
                    # 尝试从残缺值中提取年份
                    year_match = re.search(r'(19\d{2}|20\d{2})', birth)
                    if year_match:
                        info['birth_date'] = f"{year_match.group(1)}-01-01"
                        logger.info(f'出生日期残缺修复: "{birth}" → "{info["birth_date"]}"')
                    else:
                        logger.warning(f'出生日期无法解析: "{birth}"，清空')
                        info['birth_date'] = None

    @classmethod
    def _parse_llm_response(cls, content: str) -> dict[str, Any] | None:
        """
        解析 LLM 返回的结果，提取 JSON
        处理多种格式：纯 JSON、带 markdown 代码块的 JSON、JSON 片段等
        """
        if not content:
            return None

        NL = chr(10)
        original_content = content.strip()

        # 1. 清理 markdown code block
        if '```json' in original_content:
            json_part = original_content.split('```json', 1)[1]
            if '```' in json_part:
                json_part = json_part.split('```', 1)[0]
            original_content = json_part.strip()
        elif '```' in original_content:
            match = re.search(r'```(?:json)?\s*([\s\S]*?)```', original_content)
            if match:
                original_content = match.group(1).strip()
            else:
                original_content = re.sub(r'^```(?:json)?', '', original_content)
                original_content = re.sub(r'```$', '', original_content)
                original_content = original_content.strip()

        # 2. 尝试直接解析
        try:
            result = json.loads(original_content)
            if isinstance(result, dict):
                return result
            if isinstance(result, list):
                if result and isinstance(result[0], dict):
                    return result[0]
        except:
            pass

        # 3. 尝试找到第一个 { 到最后一个 } 的内容
        brace_start = original_content.find('{')
        brace_end = original_content.rfind('}')
        if brace_start != -1 and brace_end != -1 and brace_end > brace_start:
            json_candidate = original_content[brace_start:brace_end + 1]
            try:
                result = json.loads(json_candidate)
                if isinstance(result, dict):
                    return result
            except:
                pass

            # 4. 尝试修复常见的 JSON 错误
            fixed = re.sub(r',\s*([}\]])', r'\1', json_candidate)
            try:
                result = json.loads(fixed)
                if isinstance(result, dict):
                    return result
            except:
                pass

            # 5. 补充缺失的闭合括号
            open_braces = json_candidate.count('{')
            close_braces = json_candidate.count('}')
            if open_braces > close_braces:
                json_candidate_padded = json_candidate + '}' * (open_braces - close_braces)
                try:
                    result = json.loads(json_candidate_padded)
                    if isinstance(result, dict):
                        return result
                except:
                    pass

            # 6. 尝试逐行截断解析
            lines = json_candidate.split(NL)
            for i in range(len(lines), 2, -1):
                candidate = NL.join(lines[:i])
                open_c = candidate.count('{')
                close_c = candidate.count('}')
                if open_c > close_c:
                    candidate += '}' * (open_c - close_c)
                try:
                    result = json.loads(candidate)
                    if isinstance(result, dict):
                        return result
                except:
                    continue

        # 7. 尝试从文本中提取 key-value 对（作为最后手段）
        extracted = {}
        simple_patterns = [
            (r'"name"\s*:\s*"([^"]+)"', 'name'),
            (r'"gender"\s*:\s*"([^"]+)"', 'gender'),
            (r'"education"\s*:\s*"([^"]+)"', 'education'),
            (r'"major"\s*:\s*"([^"]+)"', 'major'),
            (r'"school"\s*:\s*"([^"]+)"', 'school'),
            (r'"phone"\s*:\s*"([^"]+)"', 'phone'),
            (r'"email"\s*:\s*"([^"]+)"', 'email'),
            (r'"id_card_number"\s*:\s*"([^"]+)"', 'id_card_number'),
            (r'"birth_date"\s*:\s*"([^"]+)"', 'birth_date'),
            (r'"age"\s*:\s*(\d+)', 'age'),
        ]
        for pattern, key in simple_patterns:
            match = re.search(pattern, original_content)
            if match:
                if key == 'age':
                    try:
                        extracted['age'] = int(match.group(1))
                    except:
                        pass
                else:
                    extracted[key] = match.group(1)

        if extracted:
            logger.info(f'从 LLM 响应中提取到部分字段: {list(extracted.keys())}')
            return extracted

        logger.warning('LLM 响应解析失败，无法提取有效 JSON')
        return None

    @classmethod
    def _is_valid_value(cls, value: Any, field: str) -> bool:
        """
        验证LLM返回的字段值是否有效（垃圾值过滤）
        """
        if value is None:
            return False

        # 空字符串视为无效
        if isinstance(value, str) and value.strip() == '':
            return False

        # null字符串视为无效
        if isinstance(value, str) and value.strip().lower() == 'null':
            return False

        # "未知"、"无"、"暂无"等视为无效
        invalid_strings = ['未知', '无', '暂无', '未提供', '未填写', '不适用', '保密', '—', '-', 'none', 'n/a']
        if isinstance(value, str) and value.strip() in invalid_strings:
            return False

        if field == 'name':
            if not isinstance(value, str):
                return False
            name = value.strip()
            # 姓名长度检查
            if len(name) > 30:
                return False
            # 过滤包含"简历"、"个人"、"求职"等垃圾词
            garbage_words = ['简历', '个人', '求职', '应聘', '候选人', '姓名', '信息']
            for gw in garbage_words:
                if gw in name and len(name) > 6:
                    return False
            return len(name) >= 2

        elif field == 'age':
            try:
                age = int(value)
                return 18 <= age <= 65
            except:
                return False

        elif field == 'work_years':
            try:
                years = int(value)
                return 0 <= years <= 60
            except:
                return False

        elif field == 'phone':
            if not isinstance(value, str):
                return False
            return bool(re.match(r'^1[3-9]\d{9}$', value.strip()))

        elif field == 'email':
            if not isinstance(value, str):
                return False
            return bool(re.match(r'^[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}$', value.strip()))

        elif field == 'gender':
            if not isinstance(value, str):
                return False
            return value.strip() in ['男', '女']

        elif field == 'education':
            if not isinstance(value, str):
                return False
            valid_edu = ['博士', '硕士', '研究生', '本科', '学士', '大专', '专科', '高中', '中专', '职高']
            return any(edu in value for edu in valid_edu)

        elif field in ['major', 'school', 'current_company', 'current_position']:
            if not isinstance(value, str):
                return False
            val = value.strip()
            return 2 <= len(val) <= 100

        elif field == 'degree':
            if not isinstance(value, str):
                return False
            return value.strip() in ['学士', '硕士', '博士', '无']

        elif field in ['school_system', 'study_form', 'id_card_address']:
            if not isinstance(value, str):
                return False
            return 2 <= len(value.strip()) <= 100

        elif field == 'id_card_number':
            if not isinstance(value, str):
                return False
            return bool(re.match(r'^\d{17}[\dXx]$', value.strip()))

        return True

    @classmethod
    def _sanitize_llm_result(cls, llm_result: dict[str, Any]) -> None:
        """
        清洗LLM解析结果：过滤描述性/噪声内容
        LLM容易把描述性文本误识别为公司名或职位名，需要在此清洗
        """
        if not llm_result or not isinstance(llm_result, dict):
            return

        # 描述性关键词：如果公司名/职位名以这些词开头，说明是描述而非实体
        desc_prefixes = [
            '基于', '通过', '面向', '负责', '使用', '类似', '实现',
            '完成', '优化', '提升', '搭建', '设计', '开发', '项目',
            '核心', '贡献', '说明', '描述', '工作', '职责',
            '公司介绍', '项目说明', '核心贡献', '工作职责',
        ]

        # 清洗 current_position
        pos = llm_result.get('current_position')
        if isinstance(pos, str) and pos.strip():
            pos_stripped = pos.strip()
            # 如果包含描述性前缀，清空
            for prefix in desc_prefixes:
                if pos_stripped.startswith(prefix):
                    logger.warning(f'LLM position 包含描述前缀"{prefix}"，清洗: {pos_stripped}')
                    llm_result['current_position'] = None
                    break
            # 如果太长（超过25字符）且无常见职位关键词，清空
            if llm_result.get('current_position') and len(pos_stripped) > 25:
                pos_keywords = ['工程师', '开发', '运维', '测试', '架构', '经理', '主管', '总监', '负责人', '专家', '顾问', '全栈', '前端', '后端', 'java', 'python', 'go']
                if not any(kw.lower() in pos_stripped.lower() for kw in pos_keywords):
                    logger.warning(f'LLM position 过长且无职位关键词，清洗: {pos_stripped}')
                    llm_result['current_position'] = None

        # 清洗 current_company
        comp = llm_result.get('current_company')
        if isinstance(comp, str) and comp.strip():
            comp_stripped = comp.strip()
            # 如果包含描述性前缀，清空
            for prefix in desc_prefixes:
                if comp_stripped.startswith(prefix):
                    logger.warning(f'LLM company 包含描述前缀"{prefix}"，清洗: {comp_stripped}')
                    llm_result['current_company'] = None
                    break
            # 如果太长（超过40字符）且无公司关键词，清空
            if llm_result.get('current_company') and len(comp_stripped) > 40:
                company_keywords = ['公司', '集团', '科技', '信息', '网络', '有限', '股份']
                if not any(kw in comp_stripped for kw in company_keywords):
                    logger.warning(f'LLM company 过长且无公司关键词，清洗: {comp_stripped}')
                    llm_result['current_company'] = None

    @classmethod
    def _merge_llm_result(cls, info: dict[str, Any], llm_result: dict[str, Any]) -> None:
        """
        合并LLM结果到info字典（增强版，带垃圾值过滤和空值保护）
        """
        # 简单字段映射
        simple_fields = [
            'name', 'gender', 'age', 'birth_date', 'phone', 'email',
            'education', 'major', 'school', 'graduation_date', 'work_years',
            'current_company', 'current_position',
            'degree', 'school_system', 'study_form',
            'id_card_number', 'id_card_address',
        ]

        for field in simple_fields:
            llm_value = llm_result.get(field)
            # 过滤无效值
            if not cls._is_valid_value(llm_value, field):
                continue

            # 空值/null 不覆盖已有值
            if info.get(field) is not None and info.get(field) != '' and info.get(field) != 0:
                continue

            info[field] = llm_value

        # 列表字段：去重、过滤
        if llm_result.get('technical_skills') and isinstance(llm_result['technical_skills'], list):
            existing_skills = set(info.get('technical_skills', []))
            for skill in llm_result['technical_skills']:
                # 过滤无效技能值
                if not isinstance(skill, str):
                    continue
                skill = skill.strip()
                if not skill or len(skill) > 50:  # 技能名称不超过50字符
                    continue
                if skill.lower() in ['null', '无', '未知', '暂无']:
                    continue
                if skill not in existing_skills:
                    existing_skills.add(skill)
                    info['technical_skills'].append(skill)

        if llm_result.get('certifications') and isinstance(llm_result['certifications'], list):
            if not info.get('certifications'):
                info['certifications'] = []
            existing_certs = set(info.get('certifications', []))
            for cert in llm_result['certifications']:
                if not isinstance(cert, str):
                    continue
                cert = cert.strip()
                if not cert or len(cert) > 50 or cert.lower() == 'null':
                    continue
                if cert not in existing_certs:
                    existing_certs.add(cert)
                    info['certifications'].append(cert)

        # 工作经历：去重、过滤、合并
        if llm_result.get('work_experiences') and isinstance(llm_result['work_experiences'], list):
            if not info.get('work_experiences'):
                info['work_experiences'] = []

            existing_companies = {w.get('company', '') for w in info['work_experiences']}
            for work in llm_result['work_experiences']:
                if not isinstance(work, dict):
                    continue
                company = work.get('company', '')
                if not company or not isinstance(company, str):
                    continue
                if company in existing_companies:
                    continue
                # 过滤垃圾公司名称
                if any(kw in company for kw in ['简历', '个人', '示例', 'test', 'Test']):
                    continue
                existing_companies.add(company)
                # 清理工作描述中的冗余信息
                description = work.get('description', '')
                if isinstance(description, str):
                    description = description.strip()[:500]  # 限制长度
                    work['description'] = description
                info['work_experiences'].append(work)

        # 项目经验
        if llm_result.get('project_experiences') and isinstance(llm_result['project_experiences'], list):
            if not info.get('project_experiences'):
                info['project_experiences'] = []

            existing_projects = {p.get('project_name', '') for p in info['project_experiences']}
            for project in llm_result['project_experiences']:
                if not isinstance(project, dict):
                    continue
                project_name = project.get('project_name', '')
                if not project_name or not isinstance(project_name, str):
                    continue
                if project_name in existing_projects:
                    continue
                if any(kw in project_name for kw in ['示例', 'test', 'Test']):
                    continue
                existing_projects.add(project_name)
                info['project_experiences'].append(project)

        # 如果年龄为空但有出生日期，计算年龄
        if (info.get('age') is None or info['age'] == 0) and info.get('birth_date'):
            try:
                birth_year = int(info['birth_date'][:4])
                current_year = time.localtime().tm_year
                info['age'] = current_year - birth_year
            except:
                pass

        # 如果工作年限为空但有工作经历，估算工作年限
        if (info.get('work_years') is None or info['work_years'] == 0) and info.get('work_experiences'):
            try:
                total_years = 0
                for work in info['work_experiences']:
                    start = work.get('start_date', '')
                    end = work.get('end_date', '')
                    if start and len(start) >= 4:
                        start_year = int(start[:4])
                        if end and len(end) >= 4 and end.lower() != '至今':
                            end_year = int(end[:4])
                        else:
                            end_year = time.localtime().tm_year
                        total_years += max(0, end_year - start_year)
                if total_years > 0:
                    info['work_years'] = total_years
            except:
                pass

        # 如果当前公司为空，从工作经历取第一个
        if not info.get('current_company') and info.get('work_experiences') and len(info['work_experiences']) > 0:
            info['current_company'] = info['work_experiences'][0].get('company', '')

        if not info.get('current_position') and info.get('work_experiences') and len(info['work_experiences']) > 0:
            info['current_position'] = info['work_experiences'][0].get('position', '')

        logger.info(f'LLM解析结果合并完成，共提取 {len(info.get("work_experiences", []))} 条工作经历，{len(info.get("project_experiences", []))} 条项目经验')

    # ========== 标签生成 ==========

    @classmethod
    def _generate_tags(cls, info: dict[str, Any]) -> list[str]:
        """
        根据结构化信息自动生成标签
        """
        tags = []

        # 学历标签
        if info.get('education'):
            tags.append(f"edu:{info['education']}")

        # 专业标签
        if info.get('major'):
            tags.append(f"major:{info['major']}")

        # 年龄标签（分段）
        age = info.get('age')
        if age:
            if age < 25:
                tags.append("age:25岁以下")
            elif age < 30:
                tags.append("age:25-30岁")
            elif age < 35:
                tags.append("age:30-35岁")
            elif age < 40:
                tags.append("age:35-40岁")
            else:
                tags.append("age:40岁以上")

        # 毕业时间标签
        if info.get('graduation_date'):
            tags.append(f"grad:{info['graduation_date']}")

        # 工作年限标签（分段）
        work_years = info.get('work_years')
        if work_years:
            if work_years < 3:
                tags.append("exp:3年以下")
            elif work_years < 5:
                tags.append("exp:3-5年")
            elif work_years < 10:
                tags.append("exp:5-10年")
            else:
                tags.append("exp:10年以上")

        # 技术技能标签
        for skill in info.get('technical_skills', []):
            tags.append(f"skill:{skill}")

        # 公司标签
        for work in info.get('work_experiences', []):
            company = work.get('company', '')
            if company:
                tags.append(f"company:{company}")

        # 当前职位标签
        if info.get('current_position'):
            tags.append(f"position:{info['current_position']}")

        # 性别标签
        if info.get('gender'):
            tags.append(f"gender:{info['gender']}")

        return tags

    @classmethod
    def _mask_phone(cls, phone: str | None) -> str:
        """
        手机号脱敏
        """
        if not phone or len(phone) < 7:
            return phone or ''
        return phone[:3] + '****' + phone[-4:]

    @classmethod
    def _mask_email(cls, email: str | None) -> str:
        """
        邮箱脱敏
        """
        if not email or '@' not in email:
            return email or ''
        at_index = email.index('@')
        if at_index <= 2:
            return email
        return email[:2] + '***' + email[at_index:]