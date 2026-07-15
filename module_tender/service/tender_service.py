"""
招标文件智能生成模块服务层
"""
import json
import os
import re
import time
import uuid
import asyncio
from typing import Any
from decimal import Decimal

from sqlalchemy import select, text
from sqlalchemy.ext.asyncio import AsyncSession

from common.vo import CrudResponseModel, PageModel
from exceptions.exception import ServiceException
from module_tender.dao.tender_dao import TenderDao
from module_tender.entity.do.tender_do import OaTenderDocument, OaTenderRequirement, OaBidPersonnelMapping
from module_tender.entity.vo.tender_vo import (
    TenderDocumentModel,
    TenderDocumentPageQueryModel,
    TenderDocumentUploadResultModel,
    TenderDocumentDetailModel,
    TenderRequirementModel,
    BidPersonnelMappingModel,
    MatchResultModel,
)
from module_resume_kb.entity.do.resume_do import OaResume
from config.database import AsyncSessionLocal, SyncSessionLocal
from utils.log_util import logger
from utils.time_format_util import TimeFormatUtil
from utils.common_util import SqlalchemyUtil
from config.env import AiResumeParserConfig
import httpx
import fitz


class TenderService:
    """
    招标文件智能生成模块服务层
    """

    # ========== 招标文件解析 ==========

    @classmethod
    async def upload_tender_document_services(
        cls,
        query_db: AsyncSession,
        file_path: str,
        file_name: str,
        user_id: int,
    ) -> TenderDocumentUploadResultModel:
        """
        上传并解析招标文件

        流程：
        1. 提取文件文本（PDF用PyMuPDF，Word用python-docx）
        2. LLM提取关键信息（项目信息+人员要求+评标标准）
        3. 转为结构化要求存入oa_tender_requirement
        4. 创建招标文件主记录
        """
        start_time = time.time()

        try:
            # 1. 解析文件文本
            file_ext = os.path.splitext(file_name)[1].lower()
            full_text = cls._parse_tender_document(file_path, file_ext)
            total_pages = cls._get_file_page_count(file_path, file_ext)

            if not full_text.strip():
                raise ServiceException(message='招标文件内容为空，无法解析')

            logger.info(f'招标文件解析: {file_name}，共{total_pages}页，文本长度{len(full_text)}')

            # 2. LLM提取关键信息（失败不中断，fallback规则提取兜底）
            try:
                tender_info = await cls._extract_tender_info_by_llm(full_text)
            except Exception as llm_err:
                logger.warning(f'LLM提取失败，将使用规则提取: {llm_err}')
                tender_info = None

            # 始终运行规则提取作为补充（LLM可能部分提取成功，部分失败）
            fallback_info = cls._fallback_extract_info(full_text, file_path, file_ext)

            # 合并结果：LLM优先，空字段用fallback填充
            if tender_info:
                for key in ('tender_name', 'tender_code', 'company_name'):
                    if not tender_info.get(key) and fallback_info.get(key):
                        logger.info(f'LLM未提取到{key}，使用规则提取结果: {fallback_info[key]}')
                        tender_info[key] = fallback_info[key]
                if not tender_info.get('requirements'):
                    tender_info['requirements'] = fallback_info.get('requirements', [])
            else:
                tender_info = fallback_info

            if not tender_info:
                raise ServiceException(message='招标文件解析失败：无法提取关键信息')

            # 3. 数据库操作（使用同步session在线程池执行，彻底绕开异步session的greenlet上下文问题）
            loop = asyncio.get_event_loop()
            result = await loop.run_in_executor(
                None,
                cls._save_tender_sync,
                file_name,
                file_path,
                total_pages,
                user_id,
                tender_info,
            )

            parse_time = int((time.time() - start_time) * 1000)
            logger.info(
                f'招标文件解析完成: {file_name}，提取{result["requirement_count"]}条要求，耗时{parse_time}ms'
            )

            return TenderDocumentUploadResultModel(
                tender_id=result['tender_id'],
                tender_uuid=result['tender_uuid'],
                tender_name=result['tender_name'],
                tender_code=result['tender_code'],
                requirement_count=result['requirement_count'],
                message=f'解析完成：提取{result["requirement_count"]}条人员配置要求',
            )

        except ServiceException:
            raise
        except Exception as e:
            logger.error(f'招标文件解析入库失败: {str(e)}')
            raise ServiceException(message=f'招标文件解析入库失败: {str(e)}')

    @classmethod
    def _save_tender_sync(
        cls,
        file_name: str,
        file_path: str,
        total_pages: int,
        user_id: int,
        tender_info: dict,
    ) -> dict:
        """同步保存招标文件数据（在线程池中执行，避免greenlet问题）"""
        with SyncSessionLocal() as db:
            try:
                tender_uuid = str(uuid.uuid4())
                requirements_json = json.dumps(tender_info.get('requirements', []) or [], ensure_ascii=False)
                score_standard_json = json.dumps(tender_info.get('score_standards', []) or [], ensure_ascii=False)
                now_ts = TimeFormatUtil.get_current_timestamp()

                tender = OaTenderDocument(
                    tender_uuid=tender_uuid,
                    file_name=file_name,
                    tender_name=tender_info.get('tender_name', file_name),
                    tender_code=tender_info.get('tender_code', ''),
                    company_name=tender_info.get('company_name', ''),
                    total_pages=total_pages,
                    status=0,
                    requirements_json=requirements_json,
                    score_standard_json=score_standard_json,
                    file_path=file_path,
                    generated_file_path='',
                    admin_id=user_id,
                    create_time=now_ts,
                    update_time=now_ts,
                )
                db.add(tender)
                db.flush()
                db.refresh(tender)

                # 批量创建结构化要求
                requirements_data = tender_info.get('requirements') or []
                requirement_count = 0
                for req_data in requirements_data:
                    req_do = OaTenderRequirement(
                        tender_id=tender.id,
                        requirement_type=req_data.get('requirement_type', ''),
                        requirement_key=req_data.get('requirement_key', ''),
                        operator=req_data.get('operator', ''),
                        requirement_value=req_data.get('requirement_value', ''),
                        score_weight=Decimal(str(req_data.get('score_weight', 0))),
                        description=req_data.get('description', ''),
                        create_time=now_ts,
                    )
                    db.add(req_do)
                    requirement_count += 1

                if requirement_count > 0:
                    db.flush()

                db.commit()
                logger.info(f'同步事务提交成功, tender_id={tender.id}')

                return {
                    'tender_id': tender.id,
                    'tender_uuid': tender_uuid,
                    'tender_name': tender.tender_name,
                    'tender_code': tender.tender_code,
                    'requirement_count': requirement_count,
                }
            except Exception:
                db.rollback()
                raise

    @classmethod
    def _parse_tender_document(cls, file_path: str, file_ext: str) -> str:
        """解析招标文件文本（支持PDF和Word）"""
        try:
            if file_ext == '.pdf':
                doc = fitz.open(file_path)
                all_text = []
                for page in doc:
                    text = page.get_text()
                    if text.strip():
                        all_text.append(text)
                doc.close()
                return '\n'.join(all_text)
            elif file_ext in ('.docx', '.doc'):
                from docx import Document
                doc = Document(file_path)
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                # 也提取表格内容
                for table in doc.tables:
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if cells:
                            paragraphs.append(' | '.join(cells))
                return '\n'.join(paragraphs)
            elif file_ext == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            else:
                raise ServiceException(message=f'不支持的文件格式: {file_ext}')
        except ServiceException:
            raise
        except Exception as e:
            logger.error(f'解析招标文件失败: {str(e)}')
            raise ServiceException(message=f'解析招标文件失败: {str(e)}')

    @classmethod
    def _fallback_extract_info(cls, text: str, file_path: str = '', file_ext: str = '') -> dict:
        """
        使用规则提取关键信息（作为LLM的补充或fallback）
        
        策略：
        1. 基本信息提取：处理多行label-value模式（标签和值在不同行）
        2. 人员要求提取：直接解析docx表格
        """
        result = {
            'tender_name': '',
            'tender_code': '',
            'company_name': '',
            'requirements': [],
            'score_standards': []
        }

        lines = text.split('\n')

        # ===== 1. 提取招标编号 =====
        # 模式A: 同一行 "招标编号：BJFLX-2026-112"
        for i, line in enumerate(lines[:200]):
            match = re.search(r'招\s*标\s*编\s*号[：:\s]*([A-Za-z0-9\-]+)', line)
            if match and match.group(1):
                result['tender_code'] = match.group(1).strip()
                break
        # 模式B: 标签行后跟值行（docx常见："招 标 编 号:" → "：" → "BJFLX-2026-112"）
        if not result['tender_code']:
            for i, line in enumerate(lines[:30]):
                if re.search(r'招\s*标\s*编\s*号', line):
                    # 向下搜索最多3行找编号值
                    for j in range(i + 1, min(len(lines), i + 4)):
                        val_match = re.search(r'([A-Z]{2,}[-]?\d{2,}[-]?[A-Za-z0-9\-]+)', lines[j])
                        if val_match:
                            result['tender_code'] = val_match.group(1).strip()
                            break
                    if result['tender_code']:
                        break
        # 模式C: "1、招标编号：BJFLX-2026-112"
        if not result['tender_code']:
            for i, line in enumerate(lines[:200]):
                match = re.search(r'编\s*号[：:\s]*([A-Za-z0-9]{3,}[-\w]+)', line)
                if match:
                    result['tender_code'] = match.group(1).strip()
                    break

        # ===== 2. 提取项目名称 =====
        # 模式A: 同一行 "项目名称：XXX"
        for i, line in enumerate(lines[:200]):
            match = re.search(r'项\s*目\s*名\s*称[：:\s]*(.+)', line)
            if match and match.group(1).strip():
                result['tender_name'] = match.group(1).strip()
                break
        # 模式B: 第一行通常是标题
        if not result['tender_name']:
            for line in lines[:5]:
                stripped = line.strip()
                if stripped and len(stripped) > 4 and len(stripped) < 100:
                    result['tender_name'] = stripped
                    break

        # ===== 3. 提取招标人/招标单位 =====
        # 模式A: 同一行 "招标人：北京XXX公司"
        for i, line in enumerate(lines[:200]):
            match = re.search(r'招\s*标\s*人[：:\s]+([\u4e00-\u9fa5][\u4e00-\u9fa5（）\(\)有限公司股份集团有限责任]+)', line)
            if match and len(match.group(1).strip()) > 4:
                result['company_name'] = match.group(1).strip()
                break
        # 模式B: 标签行后跟值行（docx常见："招 标 人:" → "：" → "北京XXX公司"）
        if not result['company_name']:
            for i, line in enumerate(lines[:30]):
                if re.search(r'招\s*标\s*人', line) and '代理' not in line:
                    # 向下搜索最多3行找公司名
                    for j in range(i + 1, min(len(lines), i + 4)):
                        company_match = re.search(r'([\u4e00-\u9fa5]{4,}(?:有限公司|股份有限公司|集团|有限责任公司))', lines[j])
                        if company_match:
                            result['company_name'] = company_match.group(1).strip()
                            break
                    if result['company_name']:
                        break
        # 模式C: 从招标公告正文提取 "受招标人XXX委托"
        if not result['company_name']:
            for line in lines[:200]:
                match = re.search(r'受\s*招\s*标\s*人\s*([\u4e00-\u9fa5]{4,}(?:有限公司|股份有限公司|集团|有限责任公司))', line)
                if match:
                    result['company_name'] = match.group(1).strip()
                    break

        # ===== 4. 提取人员配置要求 =====
        # 优先从docx表格中直接解析（最准确）
        if file_path and file_ext in ('.docx', '.doc'):
            try:
                table_reqs = cls._extract_requirements_from_docx_tables(file_path)
                if table_reqs:
                    result['requirements'] = table_reqs
                    logger.info(f'从docx表格提取到{len(table_reqs)}条人员要求')
            except Exception as e:
                logger.warning(f'从docx表格提取人员要求失败: {e}')

        # 如果表格提取失败，用正则从文本提取
        if not result['requirements']:
            result['requirements'] = cls._extract_requirements_from_text(lines)

        logger.info(
            f'规则提取结果: tender_name={result["tender_name"]}, '
            f'tender_code={result["tender_code"]}, '
            f'company_name={result["company_name"]}, '
            f'requirements={len(result["requirements"])}'
        )
        return result

    @classmethod
    def _extract_requirements_from_docx_tables(cls, file_path: str) -> list[dict]:
        """从docx表格中提取人员配置要求"""
        from docx import Document

        doc = Document(file_path)
        requirements = []

        for table in doc.tables:
            if len(table.rows) < 2:
                continue

            # 检查表头是否包含人员要求相关关键词
            header_cells = [cell.text.strip() for cell in table.rows[0].cells]
            # 也检查第二行（有些表格有合并单元格的二级表头）
            row1_cells = [cell.text.strip() for cell in table.rows[1].cells] if len(table.rows) > 1 else []

            all_header_text = ' '.join(header_cells + row1_cells)

            # 匹配人员要求表格的特征：表头含"工作年限/任职要求/职责"等
            if not any(kw in all_header_text for kw in ['工作年限', '任职要求', '职责', '名称']):
                continue
            # 同时检查数据行是否含岗位关键词（后端/前端/UI等在数据行而非表头）
            all_data_text = ''
            for ri in range(2, min(len(table.rows), 5)):
                for cell in table.rows[ri].cells:
                    all_data_text += cell.text.strip() + ' '
            if not any(kw in all_data_text for kw in ['后端', '前端', '开发', 'UI', 'AI', '咨询', '方案', '初级', '中级', '高级', '专家']):
                continue

            logger.info(f'找到人员要求表格，共{len(table.rows)}行')

            # 解析表格数据行
            for ri in range(2, len(table.rows)):  # 跳过前两行表头
                row = table.rows[ri]
                cells = [cell.text.strip() for cell in row.cells]

                if len(cells) < 3:
                    continue

                name = cells[0]  # 如 "后端开发\n（初级）"
                work_years_cells = cells[1:4] if len(cells) >= 4 else cells[1:]  # 工作年限列（本科以下/本科/研究生）
                duty = cells[4] if len(cells) > 4 else ''
                req_text = cells[5] if len(cells) > 5 else ''

                if not name or len(name) < 2:
                    continue

                # 解析岗位名称和级别
                name_clean = name.replace('\n', '').replace('\r', '').strip()
                level_match = re.search(r'[（(](初|中|高|专家)[级]?[）)]', name_clean)
                level = level_match.group(1) if level_match else ''
                position = re.sub(r'[（(].*?[）)]', '', name_clean).strip()

                if not position or not level:
                    continue

                # 提取本科工作年限（最常用的学历要求）
                work_years_str = ''
                for wy_cell in work_years_cells:
                    years_match = re.search(r'(\d+)[-~](\d+)', wy_cell)
                    if years_match:
                        work_years_str = f'{years_match.group(1)}-{years_match.group(2)}年'
                        break
                    years_match = re.search(r'(\d+)\s*年以上', wy_cell)
                    if years_match:
                        work_years_str = f'{years_match.group(1)}年以上'
                        break

                if not work_years_str:
                    continue

                # 创建工作年限要求
                min_years_match = re.match(r'(\d+)', work_years_str)
                min_years = min_years_match.group(1) if min_years_match else '1'

                requirements.append({
                    'requirement_type': '工作',
                    'requirement_key': 'work_years',
                    'operator': 'gte',
                    'requirement_value': min_years,
                    'score_weight': 15,
                    'description': f'{position}（{level}）: 工作年限{work_years_str}'
                })

                # 学历要求（该表格结构中包含"本科"和"研究生"列）
                requirements.append({
                    'requirement_type': '学历',
                    'requirement_key': 'education',
                    'operator': 'in',
                    'requirement_value': '本科,研究生',
                    'score_weight': 10,
                    'description': f'{position}（{level}）: 学历要求本科及以上'
                })

                # 技能要求（从任职要求列提取）
                if req_text:
                    # 提取技术关键词
                    tech_keywords = []
                    for kw in ['SpringBoot', 'MyBatis', 'MySQL', 'Redis', 'Vue', 'React', 'Python',
                               'Java', 'JavaScript', 'TypeScript', 'Linux', 'Docker', 'Kubernetes',
                               'Spring Cloud', 'RabbitMQ', 'Kafka', 'ElasticSearch', 'MongoDB',
                               'AI', '大模型', 'Prompt', 'RAG', 'Agent', 'Figma', 'Sketch',
                               'Photoshop', 'Illustrator', 'HTML', 'CSS', 'Node.js']:
                        if kw.lower() in req_text.lower():
                            tech_keywords.append(kw)

                    if tech_keywords:
                        requirements.append({
                            'requirement_type': '技能',
                            'requirement_key': 'skills',
                            'operator': 'contains',
                            'requirement_value': ','.join(tech_keywords[:8]),
                            'score_weight': 20,
                            'description': f'{position}（{level}）: 技能要求 {",".join(tech_keywords[:5])}'
                        })

        return requirements

    @classmethod
    def _extract_requirements_from_text(cls, lines: list[str]) -> list[dict]:
        """从纯文本行中提取人员要求（当表格解析不可用时）"""
        requirements = []
        seen = set()

        for i, line in enumerate(lines):
            # 匹配 "后端开发（初级）" 等模式
            match = re.search(r'([前后]端开发|UI产品|AI应用|解决方案|项目咨询|测试|运维|工程师|项目经理|架构师|设计师)\s*[（(](初|中|高|专家)[级]?[）)]', line)
            if match:
                position = match.group(1)
                level = match.group(2)

                # 向下搜索工作年限
                for j in range(i, min(len(lines), i + 10)):
                    years_match = re.search(r'(\d+)[-~](\d+)\s*年', lines[j])
                    if years_match:
                        min_years = years_match.group(1)
                        max_years = years_match.group(2)
                        key = ('work_years', position, level)
                        if key not in seen:
                            seen.add(key)
                            requirements.append({
                                'requirement_type': '工作',
                                'requirement_key': 'work_years',
                                'operator': 'gte',
                                'requirement_value': min_years,
                                'score_weight': 15,
                                'description': f'{position}（{level}）: 工作年限{min_years}-{max_years}年'
                            })
                        break
                    years_match = re.search(r'(\d+)\s*年以上', lines[j])
                    if years_match:
                        min_years = years_match.group(1)
                        key = ('work_years', position, level)
                        if key not in seen:
                            seen.add(key)
                            requirements.append({
                                'requirement_type': '工作',
                                'requirement_key': 'work_years',
                                'operator': 'gte',
                                'requirement_value': min_years,
                                'score_weight': 15,
                                'description': f'{position}（{level}）: 工作年限{min_years}年以上'
                            })
                        break

        return requirements[:20]  # 最多20条

    @classmethod
    def _get_file_page_count(cls, file_path: str, file_ext: str) -> int:
        """获取文件页数"""
        try:
            if file_ext == '.pdf':
                doc = fitz.open(file_path)
                count = doc.page_count
                doc.close()
                return count
            elif file_ext in ('.docx', '.doc'):
                # Word文档页数无法直接获取，按字数粗略估算（每页约500字）
                text = cls._parse_tender_document(file_path, file_ext)
                char_count = len(text.strip())
                pages = max(1, char_count // 500)
                return pages
            else:
                # TXT等纯文本按字数估算
                text = cls._parse_tender_document(file_path, file_ext)
                char_count = len(text.strip())
                return max(1, char_count // 500)
        except Exception:
            return 1

    @classmethod
    async def _extract_tender_info_by_llm(cls, text: str) -> dict[str, Any]:
        """
        使用LLM从招标文件文本中提取关键信息
        
        采用两阶段策略：
        1. 第一阶段：从文件开头提取基本信息（项目名称、招标编号、招标单位）
        2. 第二阶段：从全文提取人员配置要求
        """
        api_key = AiResumeParserConfig.ai_resume_parser_api_key
        base_url = AiResumeParserConfig.ai_resume_parser_base_url
        model = AiResumeParserConfig.ai_resume_parser_model
        temperature = AiResumeParserConfig.ai_resume_parser_temperature or 0.1
        max_tokens = AiResumeParserConfig.ai_resume_parser_max_tokens or 8192

        # ===== 第一阶段：提取基本信息（文件开头3000字符足够）=====
        basic_text = text[:3000]
        logger.info(f'LLM第一阶段提取基本信息: 文本长度{len(basic_text)}')

        basic_prompt = f"""从招标文件中提取基本信息。只输出JSON，不要任何解释。

查找规则：
- "项目名称"后面的文字就是 tender_name
- "招标编号"后面的字母数字组合就是 tender_code（如BJFLX-2026-112格式）
- "招标人"后面的公司名就是 company_name（注意可能跨行，跳过冒号行）
- 如果文件开头是标题，它也可能是 tender_name

输出格式：
{{"tender_name": "项目名称", "tender_code": "编号", "company_name": "招标单位名称"}}

招标文件开头部分：
{basic_text}"""

        result = {
            'tender_name': '',
            'tender_code': '',
            'company_name': '',
            'requirements': [],
            'score_standards': []
        }

        try:
            basic_result = await cls._call_llm(basic_prompt, api_key, base_url, model, temperature, max_tokens)
            if basic_result:
                result['tender_name'] = basic_result.get('tender_name', '') or ''
                result['tender_code'] = basic_result.get('tender_code', '') or ''
                result['company_name'] = basic_result.get('company_name', '') or ''
                logger.info(f'LLM第一阶段提取成功: name={result["tender_name"]}, code={result["tender_code"]}, company={result["company_name"]}')
        except Exception as e:
            logger.warning(f'LLM提取基本信息失败，将依赖规则提取: {e}')

        # ===== 第二阶段：提取人员要求（全文，最多80000字符）=====
        max_len = 80000
        req_text = text[:max_len] if len(text) > max_len else text
        logger.info(f'LLM第二阶段提取人员要求: 文本长度{len(req_text)}')

        req_prompt = f"""从招标文件中提取人员配置要求。只输出JSON，不要任何解释。

查找规则：
1. 在文件中搜索"服务需求""服务要求""人员配置""岗位要求"等关键词所在的章节
2. 重点关注表格内容，特别是含有"初级""中级""高级""专家级"等职称分级的表格
3. 对每个岗位级别提取：
   - 工作年限（如"2-4年""5年以上"，取最小值作为requirement_value）
   - 学历要求（如"本科""研究生"）
   - 技能要求（从任职要求列提取技术关键词如SpringBoot/Vue/MySQL等）
4. 每个要求单独作为数组元素

输出格式：
{{"requirements": [
    {{"requirement_type": "工作", "requirement_key": "work_years", "operator": "gte", "requirement_value": "2", "score_weight": 15, "description": "后端开发初级：工作年限2-4年"}},
    {{"requirement_type": "学历", "requirement_key": "education", "operator": "in", "requirement_value": "本科,研究生", "score_weight": 10, "description": "学历要求本科及以上"}}
]}}

招标文件全文：
{req_text}"""

        try:
            req_result = await cls._call_llm(req_prompt, api_key, base_url, model, temperature, max_tokens)
            if req_result and req_result.get('requirements'):
                result['requirements'] = req_result['requirements']
                logger.info(f'LLM第二阶段提取到{len(result["requirements"])}条人员要求')
        except Exception as e:
            logger.warning(f'LLM提取人员要求失败，将依赖规则提取: {e}')

        return result

    @classmethod
    async def _call_llm(cls, prompt: str, api_key: str, base_url: str, model: str, temperature: float, max_tokens: int) -> dict:
        """调用LLM API并返回解析后的JSON"""
        headers = {
            'Authorization': f'Bearer {api_key}',
            'Content-Type': 'application/json',
        }
        payload = {
            'model': model,
            'messages': [
                {
                    'role': 'system',
                    'content': '你是一个招标文件信息提取助手。只输出JSON，不输出任何其他文字。',
                },
                {
                    'role': 'user',
                    'content': prompt,
                },
            ],
            'temperature': temperature,
            'max_tokens': max_tokens,
            'response_format': {'type': 'json_object'},
        }
        url = f'{base_url}/chat/completions'
        logger.info(f'调用LLM API: url={url}, model={model}')

        async with httpx.AsyncClient(timeout=httpx.Timeout(120.0)) as client:
            resp = await client.post(url, json=payload, headers=headers)
            resp.raise_for_status()
            resp_json = resp.json()

        content = None
        if resp_json and resp_json.get('choices'):
            content = resp_json['choices'][0].get('message', {}).get('content')

        if not content:
            logger.warning('LLM返回空内容')
            return {}

        logger.info(f'LLM返回内容长度: {len(content)}')
        logger.info(f'LLM返回完整内容: {content}')

        result = cls._parse_llm_json_response(content)
        return result

    @classmethod
    def _parse_llm_json_response(cls, content: str) -> dict[str, Any]:
        """解析LLM返回的JSON内容（增强容错：处理单引号、中文引号、尾逗号等常见格式问题）"""
        import ast

        content = content.strip()

        # 1. 移除markdown代码块标记
        if content.startswith('```json'):
            content = content[7:]
        if content.startswith('```'):
            content = content[3:]
        if content.endswith('```'):
            content = content[:-3]
        content = content.strip()

        # 2. 尝试直接解析
        try:
            result = json.loads(content)
            return cls._normalize_llm_result(result)
        except json.JSONDecodeError:
            pass

        # 3. 清理常见格式问题后再试
        cleaned = content
        # 替换中文引号为英文双引号
        cleaned = cleaned.replace('"', '"').replace('"', '"')
        cleaned = cleaned.replace(''', "'").replace(''', "'")
        # 处理尾逗号（对象或数组最后一个元素后的逗号）
        cleaned = re.sub(r',\s*([}\]])', r'\1', cleaned)

        try:
            result = json.loads(cleaned)
            return cls._normalize_llm_result(result)
        except json.JSONDecodeError:
            pass

        # 4. 尝试用 ast.literal_eval 解析（兼容单引号）
        try:
            result = ast.literal_eval(cleaned)
            if isinstance(result, dict):
                return cls._normalize_llm_result(result)
        except (SyntaxError, ValueError):
            pass

        # 5. 正则提取最外层JSON对象
        json_match = re.search(r'\{[\s\S]*\}', content)
        if json_match:
            candidate = json_match.group()
            # 再做一次清理
            candidate = candidate.replace('"', '"').replace('"', '"')
            candidate = re.sub(r',\s*([}\]])', r'\1', candidate)
            try:
                result = json.loads(candidate)
                return cls._normalize_llm_result(result)
            except json.JSONDecodeError:
                try:
                    result = ast.literal_eval(candidate)
                    if isinstance(result, dict):
                        return cls._normalize_llm_result(result)
                except (SyntaxError, ValueError):
                    pass

        logger.error(f'LLM返回内容无法解析为JSON: {content[:300]}')
        return {'tender_name': '', 'tender_code': '', 'company_name': '', 'requirements': [], 'score_standards': []}

    @classmethod
    def _normalize_llm_result(cls, result: dict) -> dict:
        """规范化LLM返回结果，确保关键字段存在"""
        # 确保基本字段存在
        if not result.get('tender_name'):
            result['tender_name'] = ''
        if not result.get('tender_code'):
            result['tender_code'] = ''
        if not result.get('company_name'):
            result['company_name'] = ''
        if not result.get('requirements'):
            result['requirements'] = []
        if not result.get('score_standards'):
            result['score_standards'] = []
        # 确保 requirements 是列表
        if not isinstance(result['requirements'], list):
            result['requirements'] = []
        if not isinstance(result['score_standards'], list):
            result['score_standards'] = []
        return result

    # ========== 招标文件列表/详情/删除 ==========

    @classmethod
    async def get_tender_document_list_services(
        cls,
        query_db: AsyncSession,
        query_object: TenderDocumentPageQueryModel,
        is_page: bool = True,
    ) -> PageModel | list:
        """获取招标文件列表"""
        return await TenderDao.get_tender_document_list(query_db, query_object, is_page)

    @classmethod
    async def get_tender_detail_services(
        cls,
        query_db: AsyncSession,
        tender_id: int,
    ) -> TenderDocumentDetailModel:
        """获取招标文件详情（含要求列表）"""
        tender = await TenderDao.get_tender_detail_by_id(query_db, tender_id)
        if not tender:
            return TenderDocumentDetailModel()

        # 将 ORM 对象转为 dict，再用 dict 构造 Pydantic model（避免 from_attributes 兼容性问题）
        tender_dict = SqlalchemyUtil.base_to_dict(tender)
        tender_model = TenderDocumentModel(**tender_dict)

        # 获取要求列表
        requirements = await TenderDao.get_requirements_by_tender_id(query_db, tender_id)
        req_models = []
        for r in requirements:
            req_dict = SqlalchemyUtil.base_to_dict(r)
            req_models.append(TenderRequirementModel(**req_dict))

        return TenderDocumentDetailModel(
            tender=tender_model,
            requirements=req_models,
        )

    @classmethod
    async def get_tender_requirements_services(
        cls,
        query_db: AsyncSession,
        tender_id: int,
    ) -> list[TenderRequirementModel]:
        """获取招标文件的结构化要求列表"""
        requirements = await TenderDao.get_requirements_by_tender_id(query_db, tender_id)
        req_models = []
        for r in requirements:
            req_dict = SqlalchemyUtil.base_to_dict(r)
            req_models.append(TenderRequirementModel(**req_dict))
        return req_models

    @classmethod
    async def delete_tender_document_services(
        cls,
        query_db: AsyncSession,
        tender_id: int,
    ) -> CrudResponseModel:
        """删除招标文件"""
        try:
            tender = await TenderDao.get_tender_detail_by_id(query_db, tender_id)
            if not tender:
                raise ServiceException(message='招标文件不存在')

            # 删除关联文件
            if tender.file_path and os.path.exists(tender.file_path):
                try:
                    os.remove(tender.file_path)
                except:
                    pass
            if tender.generated_file_path and os.path.exists(tender.generated_file_path):
                try:
                    os.remove(tender.generated_file_path)
                except:
                    pass

            await TenderDao.delete_tender_document(query_db, tender_id)
            return CrudResponseModel(is_success=True, message='删除成功')
        except ServiceException:
            raise
        except Exception as e:
            logger.error(f'删除招标文件失败: {str(e)}')
            raise ServiceException(message=f'删除招标文件失败: {str(e)}')

    # ========== 人员匹配推荐 ==========

    @classmethod
    async def match_personnel_services(
        cls,
        query_db: AsyncSession,
        tender_id: int,
    ) -> MatchResultModel:
        """
        根据招标文件要求执行人员匹配推荐

        流程：
        1. 获取招标文件结构化要求
        2. 构建SQL查询条件，从简历库筛选候选人员
        3. 按评分权重计算匹配得分
        4. 按得分排序，保存到oa_bid_personnel_mapping
        """
        try:
            # 1. 获取要求
            requirements = await TenderDao.get_requirements_by_tender_id(query_db, tender_id)
            if not requirements:
                raise ServiceException(message='招标文件未提取到人员要求，无法匹配')

            # 更新状态为匹配中
            await TenderDao.update_tender_document(query_db, tender_id, status=1)

            # 2. 构建查询条件，从简历库筛选候选人员
            candidates = await cls._query_candidates(query_db, requirements)

            if not candidates:
                await TenderDao.update_tender_document(query_db, tender_id, status=0)
                return MatchResultModel(
                    tender_id=tender_id,
                    total_candidates=0,
                    matched_count=0,
                    selected_count=0,
                    match_list=[],
                )

            # 3. 计算匹配得分
            scored_candidates = []
            for resume in candidates:
                score, reasons = cls._calculate_match_score(resume, requirements)
                scored_candidates.append((resume, score, reasons))

            # 4. 按得分排序
            scored_candidates.sort(key=lambda x: x[1], reverse=True)

            # 5. 清除旧匹配记录，保存新匹配
            await TenderDao.delete_mappings_by_tender_id(query_db, tender_id)

            mappings = []
            for idx, (resume, score, reasons) in enumerate(scored_candidates):
                mapping = OaBidPersonnelMapping(
                    tender_id=tender_id,
                    resume_id=resume.id,
                    match_score=Decimal(str(round(score, 2))),
                    match_reason='；'.join(reasons),
                    sort_order=idx + 1,
                    is_selected=0,
                    create_time=TimeFormatUtil.get_current_timestamp(),
                )
                mappings.append(mapping)

            await TenderDao.batch_create_mappings(query_db, mappings)

            # 更新状态为已完成
            await TenderDao.update_tender_document(query_db, tender_id, status=2)

            # 构建返回结果
            match_list = []
            saved_mappings = await TenderDao.get_mappings_by_tender_id(query_db, tender_id)
            for row in saved_mappings:
                mapping = row[0]
                model = BidPersonnelMappingModel(
                    id=mapping.id,
                    tender_id=mapping.tender_id,
                    resume_id=mapping.resume_id,
                    match_score=float(mapping.match_score),
                    match_reason=mapping.match_reason,
                    sort_order=mapping.sort_order,
                    is_selected=mapping.is_selected,
                    create_time=mapping.create_time,
                    resume_name=row[1] or '',
                    resume_gender=row[2] or '',
                    resume_age=row[3] or 0,
                    resume_education=row[4] or '',
                    resume_work_years=row[5] or 0,
                    resume_current_company=row[6] or '',
                    resume_current_position=row[7] or '',
                    resume_skills=row[8] or '',
                    resume_certificates=row[9] or '',
                )
                match_list.append(model)

            return MatchResultModel(
                tender_id=tender_id,
                total_candidates=len(candidates),
                matched_count=len(match_list),
                selected_count=0,
                match_list=match_list,
            )

        except ServiceException:
            raise
        except Exception as e:
            logger.error(f'人员匹配失败: {str(e)}')
            raise ServiceException(message=f'人员匹配失败: {str(e)}')

    @classmethod
    async def _query_candidates(
        cls,
        query_db: AsyncSession,
        requirements: list[OaTenderRequirement],
    ) -> list[OaResume]:
        """
        根据要求构建SQL查询，从简历库筛选候选人员

        策略：先取满足硬性条件（学历、工作年限）的候选人，
        再在内存中按技能/证书做软匹配评分
        """
        query = select(OaResume).where(OaResume.status == 1)

        for req in requirements:
            key = req.requirement_key
            op = req.operator
            val = req.requirement_value

            if key == 'education':
                if op == 'eq':
                    query = query.where(OaResume.education == val)
                elif op == 'in':
                    edu_list = [v.strip() for v in val.split(',')]
                    query = query.where(OaResume.education.in_(edu_list))

            elif key == 'work_years':
                try:
                    years = int(val)
                except:
                    continue
                if op == 'gte':
                    query = query.where(OaResume.work_years >= years)
                elif op == 'lte':
                    query = query.where(OaResume.work_years <= years)
                elif op == 'eq':
                    query = query.where(OaResume.work_years == years)

            elif key == 'skills':
                # 技能用LIKE过滤（至少匹配一个）
                if op == 'contains':
                    skills = [s.strip() for s in val.split(',') if s.strip()]
                    if skills:
                        conditions = []
                        for skill in skills:
                            conditions.append(OaResume.technical_skills.like(f'%{skill}%'))
                        # 用OR连接，至少匹配一个技能
                        from sqlalchemy import or_
                        query = query.where(or_(*conditions))

        result = await query_db.execute(query)
        return result.scalars().all()

    @classmethod
    def _calculate_match_score(
        cls,
        resume: OaResume,
        requirements: list[OaTenderRequirement],
    ) -> tuple[float, list[str]]:
        """
        计算候选人匹配得分

        评分规则：
        - 学历匹配: +20分
        - 工作年限匹配: +20分
        - 技能匹配（每匹配一项）: +10分
        - 证书匹配（每匹配一项）: +15分
        """
        score = 0.0
        reasons = []

        for req in requirements:
            key = req.requirement_key
            op = req.operator
            val = req.requirement_value
            weight = float(req.score_weight) if req.score_weight else 0

            if key == 'education':
                if op == 'eq' and resume.education == val:
                    pts = weight if weight > 0 else 20
                    score += pts
                    reasons.append(f'学历匹配({val})+{pts}分')
                elif op == 'in':
                    edu_list = [v.strip() for v in val.split(',')]
                    if resume.education in edu_list:
                        pts = weight if weight > 0 else 20
                        score += pts
                        reasons.append(f'学历匹配({resume.education})+{pts}分')

            elif key == 'work_years':
                try:
                    years = int(val)
                except:
                    continue
                resume_years = resume.work_years or 0
                matched = False
                if op == 'gte' and resume_years >= years:
                    matched = True
                elif op == 'lte' and resume_years <= years:
                    matched = True
                elif op == 'eq' and resume_years == years:
                    matched = True
                if matched:
                    pts = weight if weight > 0 else 20
                    score += pts
                    reasons.append(f'工作年限匹配({resume_years}年)+{pts}分')

            elif key == 'skills':
                if op == 'contains':
                    skills = [s.strip() for s in val.split(',') if s.strip()]
                    resume_skills_str = resume.technical_skills or ''
                    for skill in skills:
                        if skill.lower() in resume_skills_str.lower():
                            pts = weight if weight > 0 else 10
                            score += pts
                            reasons.append(f'技能匹配({skill})+{pts}分')

            elif key == 'certificates':
                if op == 'contains':
                    certs = [c.strip() for c in val.split(',') if c.strip()]
                    resume_certs_str = resume.certifications or ''
                    for cert in certs:
                        if cert.lower() in resume_certs_str.lower():
                            pts = weight if weight > 0 else 15
                            score += pts
                            reasons.append(f'证书匹配({cert})+{pts}分')

        return score, reasons

    # ========== 获取匹配结果 ==========

    @classmethod
    async def get_match_result_services(
        cls,
        query_db: AsyncSession,
        tender_id: int,
    ) -> MatchResultModel:
        """获取匹配结果列表"""
        mappings = await TenderDao.get_mappings_by_tender_id(query_db, tender_id)
        selected_count = await TenderDao.count_selected(query_db, tender_id)

        match_list = []
        for row in mappings:
            mapping = row[0]
            model = BidPersonnelMappingModel(
                id=mapping.id,
                tender_id=mapping.tender_id,
                resume_id=mapping.resume_id,
                match_score=float(mapping.match_score),
                match_reason=mapping.match_reason,
                sort_order=mapping.sort_order,
                is_selected=mapping.is_selected,
                create_time=mapping.create_time,
                resume_name=row[1] or '',
                resume_gender=row[2] or '',
                resume_age=row[3] or 0,
                resume_education=row[4] or '',
                resume_work_years=row[5] or 0,
                resume_current_company=row[6] or '',
                resume_current_position=row[7] or '',
                resume_skills=row[8] or '',
                resume_certificates=row[9] or '',
            )
            match_list.append(model)

        return MatchResultModel(
            tender_id=tender_id,
            total_candidates=len(match_list),
            matched_count=len(match_list),
            selected_count=selected_count,
            match_list=match_list,
        )

    # ========== 选择/取消人员 ==========

    @classmethod
    async def select_personnel_services(
        cls,
        query_db: AsyncSession,
        mapping_id: int,
        is_selected: int,
    ) -> CrudResponseModel:
        """选择/取消选择人员"""
        try:
            await TenderDao.update_mapping_selection(query_db, mapping_id, is_selected)
            action = '选中' if is_selected == 1 else '取消选中'
            return CrudResponseModel(is_success=True, message=f'{action}成功')
        except Exception as e:
            logger.error(f'选择人员失败: {str(e)}')
            raise ServiceException(message=f'选择人员失败: {str(e)}')

    # ========== 生成投标文件 ==========

    @classmethod
    async def generate_bid_file_services(
        cls,
        query_db: AsyncSession,
        tender_id: int,
        output_format: str = 'docx',
    ) -> str:
        """
        根据选中人员生成标准格式投标文件

        流程：
        1. 获取选中的候选人列表
        2. 按人员一览表格式组装数据
        3. 使用python-docx生成Word文档
        4. 更新招标文件的generated_file_path
        """
        try:
            # 1. 获取选中的候选人
            selected = await TenderDao.get_selected_mappings(query_db, tender_id)
            if not selected:
                raise ServiceException(message='未选择任何人员，无法生成投标文件')

            # 获取招标文件信息
            tender = await TenderDao.get_tender_detail_by_id(query_db, tender_id)
            if not tender:
                raise ServiceException(message='招标文件不存在')

            # 2. 生成Word文档
            output_dir = './uploads/tender_generated'
            os.makedirs(output_dir, exist_ok=True)

            file_name = f"投标文件_{tender.tender_name}_{int(time.time())}.docx"
            output_path = os.path.join(output_dir, file_name)

            cls._generate_docx(tender, selected, output_path)

            # 3. 更新招标文件路径
            await TenderDao.update_tender_document(
                query_db, tender_id, generated_file_path=output_path
            )

            logger.info(f'投标文件生成成功: {output_path}，共{len(selected)}人')
            return output_path

        except ServiceException:
            raise
        except Exception as e:
            logger.error(f'生成投标文件失败: {str(e)}')
            raise ServiceException(message=f'生成投标文件失败: {str(e)}')

    @classmethod
    def _generate_docx(
        cls,
        tender: OaTenderDocument,
        selected: list,
        output_path: str,
    ) -> None:
        """使用python-docx生成投标文件Word文档"""
        from docx import Document
        from docx.shared import Pt, Cm, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn

        doc = Document()

        # 设置默认字体
        style = doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(12)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 标题
        title = doc.add_heading('', level=0)
        title_run = title.add_run(f'{tender.tender_name}投标文件')
        title_run.font.name = '黑体'
        title_run.font.size = Pt(22)
        title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 项目信息
        doc.add_heading('一、项目基本信息', level=1)
        info_table = doc.add_table(rows=4, cols=2, style='Table Grid')
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        info_data = [
            ('项目名称', tender.tender_name),
            ('项目编号', tender.tender_code),
            ('招标单位', tender.company_name),
            ('投标人数', f'{len(selected)} 人'),
        ]
        for i, (label, value) in enumerate(info_data):
            cell_label = info_table.cell(i, 0)
            cell_value = info_table.cell(i, 1)
            cell_label.text = label
            cell_value.text = str(value)
            # 加粗标签列
            for paragraph in cell_label.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        # 人员一览表
        doc.add_heading('二、拟投入人员一览表', level=1)

        headers = ['序号', '姓名', '性别', '年龄', '学历', '工作年限', '当前公司', '当前职位', '专业技能']
        person_table = doc.add_table(rows=len(selected) + 1, cols=len(headers), style='Table Grid')
        person_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 表头
        for i, header in enumerate(headers):
            cell = person_table.cell(0, i)
            cell.text = header
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)

        # 数据行
        for idx, row_data in enumerate(selected, 1):
            mapping = row_data[0]
            name = row_data[1] or ''
            gender = row_data[2] or ''
            age = row_data[3] or 0
            education = row_data[4] or ''
            work_years = row_data[5] or 0
            company = row_data[6] or ''
            position = row_data[7] or ''
            skills = row_data[8] or ''
            # 解析技能JSON
            try:
                skills_list = json.loads(skills)
                skills_text = '、'.join(skills_list[:10]) if isinstance(skills_list, list) else skills
            except:
                skills_text = skills[:100] if skills else ''

            row_values = [str(idx), name, gender, str(age), education, f'{work_years}年', company, position, skills_text]
            for col, val in enumerate(row_values):
                cell = person_table.cell(idx, col)
                cell.text = val
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(10)

        # 人员详细信息
        doc.add_heading('三、人员详细信息', level=1)

        for idx, row_data in enumerate(selected, 1):
            mapping = row_data[0]
            name = row_data[1] or ''
            gender = row_data[2] or ''
            age = row_data[3] or 0
            education = row_data[4] or ''
            work_years = row_data[5] or 0
            company = row_data[6] or ''
            position = row_data[7] or ''
            skills = row_data[8] or ''
            certs = row_data[9] or ''
            major = row_data[10] or ''
            school = row_data[11] or ''

            doc.add_heading(f'{idx}. {name}（{gender}，{age}岁，{education}）', level=2)

            detail_table = doc.add_table(rows=6, cols=2, style='Table Grid')
            detail_data = [
                ('学历/专业', f'{education}，{major}，{school}'),
                ('工作年限', f'{work_years}年'),
                ('当前职务', f'{company} - {position}'),
                ('专业技能', cls._format_json_skills(skills)),
                ('持有证书', cls._format_json_skills(certs)),
                ('匹配得分', f'{float(mapping.match_score):.1f}分（{mapping.match_reason}）'),
            ]
            for i, (label, value) in enumerate(detail_data):
                detail_table.cell(i, 0).text = label
                detail_table.cell(i, 1).text = value
                for paragraph in detail_table.cell(i, 0).paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

            doc.add_paragraph('')  # 空行分隔

        # 保存文档
        doc.save(output_path)

    @classmethod
    def _format_json_skills(cls, json_str: str) -> str:
        """格式化JSON技能/证书列表为可读文本"""
        if not json_str:
            return '无'
        try:
            items = json.loads(json_str)
            if isinstance(items, list):
                return '、'.join(items) if items else '无'
            return str(items)
        except:
            return json_str[:200] if json_str else '无'

    @classmethod
    async def download_bid_file_services(
        cls,
        query_db: AsyncSession,
        tender_id: int,
    ) -> str:
        """获取生成的投标文件路径"""
        tender = await TenderDao.get_tender_detail_by_id(query_db, tender_id)
        if not tender:
            raise ServiceException(message='招标文件不存在')
        if not tender.generated_file_path or not os.path.exists(tender.generated_file_path):
            raise ServiceException(message='投标文件尚未生成，请先执行生成操作')
        return tender.generated_file_path
